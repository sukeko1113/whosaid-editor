"""Gemini API への文字起こしリクエストと、Word ファイル生成"""
from __future__ import annotations

import re
import time
from pathlib import Path

from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from . import lang

from .segments import Utterance, utterances_to_segments


# ======================================================================
# プロンプト部品
#
# v1.3.0: プロンプトを固定文字列から「部品の組み立て」に変更した。
#   - 整形ルール(従来) / 逐語ルール(新設) を切り替え可能
#   - 参加者名簿(roster)があれば話者ラベルを実名にする
# ======================================================================

_RULES_CLEANUP = """- 内容は正確に一言一句書き起こす(情報の欠落・改変なし)
- フィラー(「えー」「あのー」「えっと」「まあ」「そのー」など)は適宜削除
- 言い淀み・不要な繰り返しは整理し、読みやすい日本語に整える
- 話者の意図・固有名詞・数字は正確に保つ
- 聞き取れなかった箇所は [不明] と記す"""

_RULES_VERBATIM = """【逐語(一言一句)ルール・最重要】
- 発言は一切要約・整文しない。話されたとおりに書き起こす
- 「えー」「あのー」「えっと」等の言いよどみ(フィラー)、言い直し、繰り返しも省略せずそのまま残す
- 聞き取れない箇所は推測で補完せず、必ず「(聴取不能)」と表記する
- 音声に存在しない語句を付け加えない。文法的に不自然でも直さない
- 固有名詞が不明瞭な場合は、聞こえたとおりに表記する(勝手に「正しい」名前に直さない)
- 実際に発話された内容だけを書く。音声上で繰り返されていない限り、出力で同じ言葉を繰り返さない"""

_RULES_TS = """- **段落の先頭に必ず [MM:SS] 形式のタイムスタンプを入れる**
  (音声内のその段落が始まる時刻、ゼロ埋め2桁、例: [00:00], [03:45])
- 段落は話題のまとまり・明確な区切り・または30秒~2分ごとに分ける"""

_RULES_DIAR = """- **話者が変わるごとに新しい行として書き出す**
- **各行の冒頭に必ず以下の形式を入れる**:
  [MM:SS] 【話者ラベル】 本文...
  (時刻はゼロ埋め2桁の分:秒。ミリ秒や1/100秒は付けない)
- 短い相づち(「はい」「うん」程度)は前の発言と同じ行に含めてよい"""

# 逐語モードでは相づちを吸収させない。**吸収は _RULES_VERBATIM の
# 「一切要約・整文しない」と正面から矛盾する。**しかも消えるのではなく、
# 別の人の「はい」が前の話者の行に混ざるので、「誰が言ったか」の記録としては
# 消えるより悪い(誤って別人の発言として残る)。
# CLAUDE.md の原則:「短い相づちの自動削除・自動重複除去はしない。
# 同意の意思表示が記録から消えるため」。
#
# 併せて「話者が変わるごと」だけを区切りにするのをやめる。同じ人が話し続けると
# 1 区間が 112 秒に達し(実測)、**その区間には話者を割り当てられない**——
# 間に何人も話しているため。逐語の正解と突き合わせる測定も、区間が窓より
# 長いと成立しなくなる(段階 2 で判明)。
_RULES_DIAR_VERBATIM = """- **話者が変わるごとに新しい行として書き出す**
- **各行の冒頭に必ず以下の形式を入れる**:
  [MM:SS] 【話者ラベル】 本文...
  (時刻はゼロ埋め2桁の分:秒。ミリ秒や1/100秒は付けない)
- **短い相づち(「はい」「うん」「なるほど」等)も必ず独立した行にする。**
  前の発言の行に含めてはならない。誰がいつ同意したかが記録として要る
- **同じ話者が続く場合も、文の切れ目や間(ま)で行を分ける。
  1 行が 20 秒を超えないようにする**"""

_TAIL = "- 説明や前置き、Markdown 装飾は不要。書き起こし本文のみを出力する。"

# v2.0.0: 手動割当モード用。名簿は渡さず、声質だけで区間に切り分けさせる。
# 実名の推定をやめたので、プロンプトが短くなり応答も安定する。
_RULES_CLUSTER = """- **行を分けるのは「話す人が変わったとき」だけ**
- **各行は必ず次の形式で始める**:
  [MM:SS] 【A】 本文...
  (時刻はこの音声の先頭からの 分:秒。ゼロ埋め2桁。ミリ秒は付けない)
- 話者ラベルは A, B, C, ... のアルファベット1文字のみ。**名前や役職は絶対に書かない**
- 同じ声の人物には常に同じラベルを使う。声が変われば必ず別のラベルにする
- **同じ人が話し続けている間は、絶対に行を分けない**。
  息継ぎ・間・読点・言いよどみで区切ってはいけない。
  文がいくつ続いても、次の人が話し始めるまでは 1 行に書く。
  (悪い例: 「工程表」「えー、財源計画」「年度別資金繰り」と 3 行に分ける
   → 正しくは 1 行に「工程表、えー、財源計画、年度別資金繰り…」と続ける)
- 同じ人が 3 分以上話し続けた場合のみ、話題の切れ目で行を分けてよい
- 短い相づち(「はい」「うん」程度)も、別の人物の声なら独立した行にする
- 誰の声か判別できない場合は 【?】、複数人が同時に話している場合は 【*】 とする"""

_EXAMPLE_CLUSTER = """出力例:
[00:00] 【A】 本日はお忙しい中お集まりいただきありがとうございます。それでは議事を始めます。まず、お手元の資料をご覧ください。えー、本日の議題は、予算と、あの、人事の二点になります。
[00:25] 【B】 すみません、確認ですが、前回の議事録は配布済みですか?
[00:32] 【A】 はい、配布済みです。修正点も反映されています。
[00:40] 【C】 その件、私からも補足させてください。"""


def _roster_block(roster: str) -> str:
    return f"""この音声の参加者は以下のとおり事前に判明しています。話者の判別には、声質に加えて、
発言内容(名乗り・指名・役職に応じた発言内容)も手がかりにしてください。

【参加者名簿】
{roster.strip()}

- 話者ラベルは名簿にある呼称を【】で囲んで用いる(例: 【議長(理事長)】【佐藤理事】)
- どうしても特定できない場合のみ【発言者不明】、複数人の発話が重なる場合は【発言者複数・重複】とする。
  無理に名簿の誰かに割り当てるより、不明と明示するほうが望ましい"""


_LABEL_ABC = """- 話者ラベル(発言者A, 発言者B, 発言者C...)は声の特徴で識別し、同じ人物には常に同じラベルを使う
- 話者を特定できない場合は【発言者不明】とする"""

_EXAMPLE_DIAR = """出力例:
[00:00] 【発言者A】 本日はお忙しい中お集まりいただきありがとうございます。それでは議事を始めます。
[00:25] 【発言者B】 すみません、確認ですが、前回の議事録は配布済みですか?
[00:32] 【発言者A】 はい、配布済みです。修正点も反映されています。"""

# **例は指示より強く効く。**逐語では相づちが独立した行になっている様子と、
# 同じ話者でも 20 秒以内で分かれている様子を実演する。従来の例は 25 秒・32 秒
# 間隔で相づちも出てこないため、粗い区切りを手本として示していた。
_EXAMPLE_DIAR_VERBATIM = """出力例:
[00:00] 【発言者A】 えー、本日はお忙しい中お集まりいただき、あの、ありがとうございます。
[00:08] 【発言者B】 はい。
[00:09] 【発言者A】 それでは、えーと、議事を始めます。
[00:14] 【発言者C】 うん。
[00:15] 【発言者B】 すみません、確認ですが、前回の議事録は配布済みですか?
[00:20] 【発言者A】 はい、配布済みです。えー、修正点も反映されています。"""

_EXAMPLE_TS = """出力例:
[00:00] 本日の会議を開始します。まず議題ですが、予算と人事の二点になります。
[00:42] それでは一つ目、来期予算について議論を始めます。
[03:15] 続いて人事についての検討に移ります。"""


def build_prompt(
    with_timestamps: bool,
    with_diarization: bool,
    roster: str = "",
    verbatim: bool = False,
    cluster_only: bool = False,
) -> str:
    """設定に応じて文字起こしプロンプトを組み立てる。

    cluster_only=True のときは v2.0.0 の手動割当モード用。
    名簿を渡さず、声質だけで A/B/C… に切り分けた区間リストを作らせる。
    """
    parts: list[str] = ["この音声を日本語で書き起こしてください。", ""]

    if cluster_only:
        parts.append("ルール:")
        parts.append(_RULES_CLUSTER)
        parts.append(_RULES_VERBATIM if verbatim else _RULES_CLEANUP)
        if verbatim:
            parts.append("- 句読点は聞こえたとおりの区切りで付けてよい(内容の変更は不可)")
        parts.append(_TAIL)
        parts += ["", _EXAMPLE_CLUSTER]
        return "\n".join(parts)

    if with_diarization and roster.strip():
        parts += [_roster_block(roster), ""]

    parts.append("ルール:")
    if with_diarization:
        parts.append(_RULES_DIAR_VERBATIM if verbatim else _RULES_DIAR)
        if not roster.strip():
            parts.append(_LABEL_ABC)
    elif with_timestamps:
        parts.append(_RULES_TS)

    parts.append(_RULES_VERBATIM if verbatim else _RULES_CLEANUP)
    if verbatim:
        # 逐語モードでも段落・句読点の最低限の整形は許可する
        parts.append("- 句読点は聞こえたとおりの区切りで付けてよい(内容の変更は不可)")
    parts.append(_TAIL)

    if with_diarization and not roster.strip():
        parts += ["", _EXAMPLE_DIAR_VERBATIM if verbatim else _EXAMPLE_DIAR]
    elif with_timestamps and not with_diarization:
        parts += ["", _EXAMPLE_TS]

    return "\n".join(parts)


# 後方互換(旧コードが参照していた固定プロンプト)
PROMPT_PLAIN = build_prompt(False, False)
PROMPT_TIMESTAMP = build_prompt(True, False)
PROMPT_DIARIZATION = build_prompt(True, True)


DIARIZATION_NOTE = (
    "※ 本文中の話者ラベル(発言者A/B/C 等)は分割チャンクごとに声で識別しているため、"
    "チャンク境界をまたぐと同一人物が別ラベルになる可能性があります。"
    "長い音声では発言者の対応関係を読みながらご確認ください。"
)

ROSTER_NOTE = (
    "※ 本文中の話者ラベルは、入力された参加者名簿に基づきAIが推定したものです。"
    "特定できなかった発言は【発言者不明】等と表記されます。"
    "最終利用の前に、必ず原音声と突き合わせてご確認ください。"
)


# ======================================================================
# 暴走ループ(縮退出力)の検出
#
# v1.3.0: temperature 0.0 相当の生成で、同じパターンを無限に繰り返す
# 「暴走ループ」が発生する事例が確認された(検証実験 2026-07)。
# 検出したら temperature を段階的に上げて再生成する。
# ======================================================================

_LOOP_THRESHOLD = 15
_INLINE_LOOP = re.compile(r"(.{1,12}?)\1{%d,}" % (_LOOP_THRESHOLD - 1))

# 暴走時に段階的に試す temperature
_TEMPERATURES = (0.0, 0.3, 0.7)


def is_degenerate(text: str) -> bool:
    """暴走ループの検出(3段構え)

    (a) 行内の短いパターンの大量繰り返し(例: 「お、お、お、…」が1行に続く)
    (b) 行の種類が極端に少ない(例: 2行が交互に数千回続く)
    (c) 同一行の連続
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return True  # 空出力も失敗扱い

    for ln in lines:
        if len(ln) >= _LOOP_THRESHOLD * 2 and _INLINE_LOOP.search(ln):
            return True

    if len(lines) >= 40 and len(set(lines)) / len(lines) < 0.2:
        return True

    run = 1
    for prev, cur in zip(lines, lines[1:]):
        if cur == prev:
            run += 1
            if run >= _LOOP_THRESHOLD:
                return True
        else:
            run = 1
    return False


# ======================================================================
# タイムスタンプ処理(v1.2 から変更なし)
# ======================================================================

# Gemini が返す [MM:SS] または [M:SS](チャンク内相対時刻)を捕捉
_TS_REL_PATTERN = re.compile(r"\[(\d{1,2}):(\d{2})\]")

# 絶対時刻 + 任意の話者ラベルを段落冒頭で検出(太字化用)
_TS_LEAD_PATTERN = re.compile(
    r"^(\[\d{1,2}:\d{2}(?::\d{2})?\])"   # group(1): 時刻
    r"\s*"
    r"(【[^】]+】)?"                       # group(2): 話者ラベル(任意)
    r"\s*"
    r"(.*)"                                # group(3): 本文
)


def shift_timestamps(text: str, offset_seconds: int) -> str:
    """テキスト中の [MM:SS] (チャンク内相対時刻) を offset_seconds 加算して
    [HH:MM:SS] (元音声内の絶対時刻) に変換する。
    """
    def repl(m: "re.Match[str]") -> str:
        total = int(m.group(1)) * 60 + int(m.group(2)) + offset_seconds
        h = total // 3600
        mm = (total % 3600) // 60
        ss = total % 60
        return f"[{h:02d}:{mm:02d}:{ss:02d}]"
    return _TS_REL_PATTERN.sub(repl, text)


# ======================================================================
# API エラーの分類
#
# 残高切れ・キー不正は何度やり直しても直らない。再試行で時間を浪費した
# うえに、エラーの生 JSON を本文として書き込んでしまうと、割当画面には
# 「文字起こし失敗: ... RESOURCE_EXHAUSTED ...」という区間だけが並ぶ。
# 直せないエラーは即座に、何をすればよいかを添えて中断する。
# ======================================================================

class FatalTranscriptionError(RuntimeError):
    """再試行しても解消しないエラー(残高切れ・APIキー不正など)"""


class CancelledError(Exception):
    """利用者の操作による中止。

    transcribe_audio の中では、汎用の `except Exception` より前に
    `except CancelledError: raise` で抜けている。この順序を崩すと
    中止が「通信エラー」と誤認されて再試行に回るので注意。
    """


_MSG_DEPLETED = (
    "Gemini API の残高が尽きています。\n\n"
    "Google AI Studio (https://aistudio.google.com/) を開き、"
    "課金設定と残高を確認して補充してください。\n\n"
    "補充後に同じ設定で実行し直せば続きから再開できます"
    "(失敗したチャンクはキャッシュされていません)。"
)

_MSG_QUOTA = (
    "Gemini API の利用上限に達しています。\n\n"
    "しばらく待ってから実行し直すか、Google AI Studio "
    "(https://aistudio.google.com/) で上限と課金設定を確認してください。\n\n"
    "実行し直せば続きから再開できます。"
)

_MSG_AUTH = (
    "Gemini API キーが正しくないか、このキーでは使えないモデルです。\n\n"
    "https://aistudio.google.com/apikey で発行し直し、"
    "詳細設定の「API キー」に貼り直して「保存」を押してください。"
)


def classify_api_error(exc: Exception) -> str | None:
    """再試行しても無駄なエラーなら、利用者向けの説明文を返す。

    再試行する価値があるエラー(一時的な通信断・分あたりのレート上限など)
    では None を返す。
    """
    text = str(exc).lower()

    if any(k in text for k in ("credits are depleted", "prepayment", "billing")):
        return _MSG_DEPLETED
    if "exceeded your current quota" in text or "quota_exceeded" in text:
        return _MSG_QUOTA
    if any(k in text for k in (
        "api key not valid", "api_key_invalid", "invalid api key",
        "unauthenticated", "permission_denied", "permission denied",
    )):
        return _MSG_AUTH
    return None


def _is_rate_limited(exc: Exception) -> bool:
    """分あたりのレート上限。待てば通るので、長めに待って再試行する。"""
    text = str(exc).lower()
    return "429" in text or "resource_exhausted" in text or "rate limit" in text


# ======================================================================
# 文字起こし本体
# ======================================================================

def transcribe_audio(
    client: genai.Client,
    audio_path: Path,
    model: str,
    with_timestamps: bool = False,
    with_diarization: bool = False,
    roster: str = "",
    verbatim: bool = False,
    max_retries: int = 3,
    on_log=None,
    cluster_only: bool = False,
    is_cancelled=None,
) -> str:
    """1チャンクをGeminiで文字起こし。

    - 通信エラー等は指数バックオフで再試行(従来どおり)
    - 暴走ループを検出したら temperature を上げて再生成(v1.3.0)
    - is_cancelled が True を返したら CancelledError を送出する。
      チャンク1つが10分近くかかるので、待ち時間の途中でも中止できる
      ようにしてある(呼び出し側の「チャンクとチャンクの間」の判定だけ
      だと、押してから止まるまで数分待たされる)。
    """
    last_error: Exception | None = None
    prompt = build_prompt(with_timestamps, with_diarization, roster, verbatim, cluster_only)

    def log(msg: str) -> None:
        if on_log:
            on_log(msg)

    def cancelled() -> bool:
        return bool(is_cancelled and is_cancelled())

    def check_cancel() -> None:
        if cancelled():
            raise CancelledError()

    def sleep_cancellable(seconds: float) -> None:
        """1秒ごとに中止を確認しながら待つ。

        time.sleep(20) をそのまま使うと、その20秒間は中止ボタンが効かない。
        """
        for _ in range(int(seconds)):
            check_cancel()
            time.sleep(1)
        rest = seconds - int(seconds)
        if rest > 0:
            time.sleep(rest)

    for attempt in range(max_retries):
        uploaded = None
        try:
            check_cancel()
            # 再試行時はアップロードからやり直すので、何回目かを添えないと
            # 「なぜまた最初に戻ったのか」が利用者に分からない。
            suffix = "" if attempt == 0 else f" ({attempt + 1}/{max_retries}回目)"
            log(f"  音声をアップロード中...{suffix}")
            uploaded = client.files.upload(file=str(audio_path))

            waited = 0
            while uploaded.state.name == "PROCESSING":
                check_cancel()
                time.sleep(1)
                waited += 1
                if waited % 15 == 0:
                    log(f"  Gemini 側でファイル処理待ち... ({waited}秒経過)")
                if waited > 300:  # 5分タイムアウト
                    raise RuntimeError("ファイル処理がタイムアウトしました。")
                uploaded = client.files.get(name=uploaded.name)

            if uploaded.state.name == "FAILED":
                raise RuntimeError(f"アップロード失敗: {audio_path.name}")

            text = ""
            for temp in _TEMPERATURES:
                check_cancel()
                log("  文字起こしを生成中...(長いチャンクでは数分かかります)")
                response = client.models.generate_content(
                    model=model,
                    contents=[uploaded, prompt],
                    config=types.GenerateContentConfig(temperature=temp),
                )
                text = (response.text or "").strip()
                if not is_degenerate(text):
                    break
                log(f"  ※ 出力の暴走を検出 (temp={temp})。設定を変えて再生成します...")
                sleep_cancellable(2)
            else:
                # 全温度で暴走。最後の結果に警告を付けて返す(処理は止めない)
                log("  ※ 再生成でも暴走が解消しませんでした。このチャンクは要確認です。")
                text = f"【警告: このチャンクの出力は不安定でした。原音声を確認してください】\n{text}"

            return text

        except CancelledError:
            # 汎用ハンドラより前に置くこと。後ろだと中止が「通信エラー」
            # と誤認されて再試行に回り、中止が効かなくなる。
            raise
        except Exception as e:
            fatal = classify_api_error(e)
            if fatal:
                # 直せないエラー。再試行せず、原因と対処を添えて中断する。
                log(f"  中断: {fatal.splitlines()[0]}")
                raise FatalTranscriptionError(fatal) from e

            last_error = e
            if attempt < max_retries - 1:
                # レート上限は数秒では抜けない。長めに待つ。
                # それ以外の通信エラーも 1→2秒 では短すぎて再試行が空振りする
                # ことが多かったため 5→10秒 にしてある(2026-07 実戦投入で確認)。
                wait = (20 * (attempt + 1)) if _is_rate_limited(e) else (5 * (2 ** attempt))
                # 例外の文字列は非常に長いことがあるので頭だけ出す。
                log(f"  通信エラー: {type(e).__name__}: {str(e)[:200]}")
                log(f"  {wait}秒後に再試行します ({attempt + 2}/{max_retries}回目)...")
                sleep_cancellable(wait)

        finally:
            # **失敗しても、上げた音声を必ず消す。**正常終了のときだけ
            # 消していたので、途中でエラーになると Gemini 側に残っていた
            # (48 時間後に自動削除されるが、録音は機微情報なので待たない)。
            if uploaded is not None:
                try:
                    client.files.delete(name=uploaded.name)
                except Exception:
                    pass

    assert last_error is not None
    raise last_error


# ======================================================================
# v2.0.0: 出力テキスト → セグメント(発言区間)への変換
# ======================================================================

# チャンク内相対時刻 + 話者ラベルの行頭パターン
_SEG_LINE = re.compile(
    r"^\s*[\[［(（]\s*"
    r"(?:(\d{1,2})\s*[:：]\s*)?"        # 任意の「時」
    r"(\d{1,3})\s*[:：]\s*(\d{1,2})"     # 分:秒
    r"(?:[.,]\d+)?"                      # ミリ秒(あれば捨てる)
    r"\s*[\]］)）]\s*"
    r"(?:[【\[]\s*(?P<label>[^】\]]{1,24}?)\s*[】\]])?"   # 任意の話者ラベル
    r"\s*(?P<body>.*)$"
)

# ラベルの正規化: 「発言者A」「話者A」「Speaker A」→ "A"
_LABEL_STRIP = re.compile(r"^\s*(?:発言者|話者|speaker)\s*[::\-]?\s*", re.IGNORECASE)

CLUSTER_UNKNOWN = "?"
CLUSTER_MULTI = "*"


def normalize_cluster_label(label: str | None) -> str:
    """Gemini が返した話者ラベルを 1〜3 文字のクラスタ記号に正規化する。"""
    if not label:
        return CLUSTER_UNKNOWN
    s = _LABEL_STRIP.sub("", label.strip())
    if not s:
        return CLUSTER_UNKNOWN
    # **擬似ラベルの語彙は言語で変わる。**足りないと 【Multiple】 が下の
    # 「先頭の英字 1 文字」に落ち、クラスタ "M" という実在しない話者が立つ。
    # しかも merge_consecutive がその連続を 1 人の発言として連結する。
    #
    # 日本語の語彙も常に見る。名簿や設定が日本語のまま英語音声を流すことが
    # あり、どちらが来ても取りこぼさないほうが安全(誤判定の危険は無い——
    # 「複数」が英語の話者ラベルとして現れることはない)。
    labels = lang.current().labels
    low = s.lower()
    multi = tuple(labels.multi_words) + lang.JA.labels.multi_words
    unknown = tuple(labels.unknown_words) + lang.JA.labels.unknown_words
    if (any(k in s or k in low for k in multi)) or s in ("*", "＊"):
        return CLUSTER_MULTI
    if (any(k in s or k in low for k in unknown)) or s in ("?", "？"):
        return CLUSTER_UNKNOWN
    # 先頭の英字 1 文字を採用(「A」「A(男性)」「A さん」など)
    m = re.match(r"[A-Za-z]", s)
    if m:
        return m.group(0).upper()
    return s[:3]


# 同じ話者の連続をまとめるときの上限
MERGE_MAX_SECONDS = 180.0    # これを超えたら区切る(聴き直しが辛くなるため)
MERGE_MAX_GAP = 20.0         # 無音が長い場合は別の発言とみなす

# 逐語モードの連結の上限。**180 秒のままだと、同じ人が話し続ける限り連結が
# 進み、1 区間が 112 秒に達する(実測)。その区間には話者を割り当てられない**
# ——間に何人も話しているため。Gemini 自体は 5〜8 秒ごとに出しており、
# 潰していたのはこちらの後処理だった。
# 20 秒はプロンプトの指示(「1 行が 20 秒を超えないように」)と揃えた値。
# **連結を止めるのではない。**文の途中で割れた断片は従来どおりまとまる。
VERBATIM_MERGE_MAX_SECONDS = 20.0

# ----------------------------------------------------------------------
# 発言の長さの見積もり
#
# 区間の終わりは「次の区間の始まり」で決めているが、会話では相づちが
# 発言に重なるため、これだけだと長い発言の再生窓が 1 秒に潰れる。
#
#   [07:03] 【A】 彼が要するに院外理事で、あの、理事で、あとはみんな…
#   [07:04] 【C】 うん。      ← A がまだ話している最中の相づち
#
# 実際に確認された事例では、37 文字の発言が 1 秒で切られ、
# 再生すると「彼」だけ聞こえて終わっていた。
# 文字数から必要な秒数を見積もり、足りなければ終わりを延ばす。
# ----------------------------------------------------------------------

# 話速の目安。**言語で変わる**ので lang.py が持つ。
#   日本語 4.5 … 逐語では言いよどみが多く実測 4〜6 文字/秒。短く見積もると
#                 発言が途中で切れるので、範囲の下端を採る
#   英語  11.0 … 約 132 wpm。同じ理由で 12.5(150 wpm)より低く採る
#
# 下の定数は互換のために残してある(日本語の値に固定)。
# 切り替えを効かせるのは estimate_speech_seconds() の側。
SPEECH_CHARS_PER_SECOND = lang.JA.speech.chars_per_second
MIN_SEGMENT_SECONDS = 1.0
MAX_ESTIMATED_SECONDS = 120.0


def estimate_speech_seconds(text: str) -> float:
    """本文の文字数から、その発言に必要な秒数を見積もる。"""
    n = len(text.strip())
    if not n:
        return MIN_SEGMENT_SECONDS
    return min(MAX_ESTIMATED_SECONDS,
               max(MIN_SEGMENT_SECONDS,
                   n / lang.current().speech.chars_per_second))


def redistribute_times(raw: list[dict], chunk_len: float) -> list[dict]:
    """時刻が詰まりすぎている区間を、本文の長さで按分し直す。

    Gemini は発言が立て込むと 1 秒刻みで機械的に時刻を振ってくる。

        [06:47] 【B】 それができればいいんだけど。   ← 13文字
        [06:48] 【A】 理事がもう。                   ←  6文字
        [06:49] 【B】 はい。
        [06:50] 【A】 ま、梅田さんとあともう1人の県議以外は全部内部理事。 ← 25文字
        [06:51] 【C】 うん。
        [06:56] 【B】 県議で。

    4 秒に 5 発言。25 文字の発言が 1 秒で終わるはずがないので、この範囲の
    時刻は信用できない。一方 [06:51]→[06:56] のように次まで余裕がある区間は
    信用してよい。

    そこで「次まで余裕がある区間」を目印(アンカー)とみなし、詰まっている
    部分はアンカーまでの時間を本文の長さで按分する。

    単純に終わりを延ばすだけだと(v2.0.4)、隣り合う区間の再生窓が重なって
    同じ音声が聞こえてしまう。按分なら重ならない。
    """
    n = len(raw)
    if n == 0:
        return raw

    needs = [estimate_speech_seconds(r["text"]) for r in raw]
    starts = [float(r["rel"]) for r in raw]

    def next_start(k: int) -> float:
        return starts[k + 1] if k + 1 < n else chunk_len

    i = 0
    while i < n:
        if next_start(i) - starts[i] >= needs[i] - 0.01:
            # 余裕がある。従来どおり「次の始まり」まで
            raw[i]["rel_end"] = next_start(i)
            i += 1
            continue

        # 詰まっている。余裕のある区間(アンカー)までをひとまとまりにする。
        # アンカー自身も含めることで、詰まった区間が時間を借りられる。
        j = i
        while j < n - 1 and next_start(j) - starts[j] < needs[j] - 0.01:
            j += 1

        span_start = starts[i]
        span_end = next_start(j)
        span = max(0.3, span_end - span_start)
        total = sum(needs[i:j + 1]) or 1.0

        # 余裕がある場合に必要以上へ引き伸ばさない。
        # 足りない場合だけ、全体を同じ割合で縮める。
        scale = min(1.0, span / total)

        cursor = span_start
        for k in range(i, j + 1):
            share = needs[k] * scale
            raw[k]["rel"] = round(cursor, 3)
            raw[k]["rel_end"] = round(cursor + share, 3)
            cursor += share
        # 余った時間は最後の区間に付ける(どの区間にも属さない音声を残さない)
        if raw[j]["rel_end"] < span_end:
            raw[j]["rel_end"] = round(span_end, 3)
        i = j + 1

    return raw

# 連結時、直前がこれらで終わっていれば読点を足さない
# 互換のために残してある(日本語の値)。切り替えは _join_fragments() の側。
_SENTENCE_ENDS = lang.JA.text.sentence_ends


def _join_fragments(left: str, right: str) -> str:
    """断片を連結する。**連結の仕方は言語で変わる**ので lang.py が持つ。

    Gemini が行を分けた位置は、話者が息を継いだ位置に対応している。
    そのまま連結すると読みにくいので、直前が句読点で終わっていなければ
    区切りの記号を補う。
    (逐語での「間」の表現であり、発話内容そのものは変えない)

        日本語: 語の間に何も入れず、読点「、」を補う
                「けども」+「えー」→「けども、えー」
        英語:   語の間に空白が要る。コンマ + 空白を補う
                "Chairperson" + "we stand" → "Chairperson, we stand"

    英語には「右が句読点で始まるときは前の語に直付けする」分岐が要る
    (", however," のような断片)。日本語では句読点が必ず前の語に付くので、
    この区別が要らなかった。
    """
    t = lang.current().text
    left = left.rstrip()
    right = right.lstrip()
    if not left:
        return right
    if not right:
        return left
    if t.clinging and right[0] in t.clinging:
        return left + right          # ", and" のような句読点始まりは直付け
    if left[-1] in t.sentence_ends or right[0] in t.sentence_ends:
        return left + t.word_separator + right
    return left + t.fragment_comma + right


def merge_consecutive(
    raw: list[dict],
    max_seconds: float = MERGE_MAX_SECONDS,
    max_gap: float = MERGE_MAX_GAP,
) -> list[dict]:
    """同じ話者ラベルが連続する行をひとつの発言にまとめる。

    Gemini は指示しても、息継ぎや読点のたびに行を分けてくることがある。
    (「工程表」「えー、財源計画」「年度別資金繰り」…と 1 文が 7 行に割れる)
    そのままだと 52 分の音声で 600 区間を超え、割当作業が現実的でなくなる。
    話者が変わらない限り 1 件にまとめる。

    まとめない例外:
      - 【?】【*】(判別不能・複数人同時) … 中身がばらばらなので連結しない
      - 直前の終わりから max_gap 秒以上空いている … 別の発言とみなす
      - まとめると max_seconds を超える … 聴き直せる長さに保つ
    """
    out: list[dict] = []
    for r in raw:
        if out:
            prev = out[-1]
            same = prev["cluster"] == r["cluster"]
            pseudo = r["cluster"] in (CLUSTER_UNKNOWN, CLUSTER_MULTI)
            gap = r["rel"] - prev["rel_end"]
            span = r["rel_end"] - prev["rel"]
            if same and not pseudo and gap <= max_gap and span <= max_seconds:
                prev["text"] = _join_fragments(prev["text"], r["text"])
                prev["rel_end"] = r["rel_end"]
                continue
        out.append(dict(r))
    return out


def parse_utterances(
    text: str,
    chunk_seconds: float | None = None,
    merge_same_speaker: bool = True,
    verbatim: bool = False,
) -> list[Utterance]:
    """1 チャンクの出力テキストを Utterance の並びに変換する。

    - 時刻はチャンク内の相対秒のまま返す(絶対秒にするのは共通の後段の仕事)
    - 時刻が巻き戻っている行は直前の時刻に合わせる(Gemini が稀に乱す)
    - 行頭が時刻でない行は、直前の区間の続きとして連結する
    - 区間の終了時刻は「次の区間の開始」、最後だけチャンク末尾

    ここまでが「クラウド経路がやること」。ローカル経路(local_asr.py)は
    同じ形の Utterance を faster-whisper から直接作る(設計書 §4.2)。
    """
    raw: list[dict] = []
    prev_rel = 0.0

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        m = _SEG_LINE.match(line)
        if not m:
            if raw:
                raw[-1]["text"] = (raw[-1]["text"] + " " + line).strip()
            else:
                raw.append({
                    "rel": 0.0,
                    "cluster": CLUSTER_UNKNOWN,
                    "text": line,
                })
            continue

        h, mm, ss = m.group(1), m.group(2), m.group(3)
        rel = (int(h) * 3600 if h else 0) + int(mm) * 60 + int(ss)
        rel = max(0.0, float(rel))
        if chunk_seconds:
            rel = min(rel, float(chunk_seconds))
        if rel < prev_rel:
            rel = prev_rel
        prev_rel = rel

        body = (m.group("body") or "").strip()
        cluster = normalize_cluster_label(m.group("label"))

        # 話者ラベルも本文も無い行(区切りだけ)はスキップ
        if not body and m.group("label") is None:
            continue
        raw.append({"rel": rel, "cluster": cluster, "text": body})

    # 本文が空の区間は落とす(直前に吸収されなかったもの)
    raw = [r for r in raw if r["text"]]
    if not raw:
        return []

    # 各行の終わりを「次の行の開始」として先に決める。
    # まとめる判定に終了時刻が要るため、連結より前に確定させる。
    #
    # 最後の行だけは暫定値を入れておく。ここでチャンク末尾まで伸ばすと、
    # 末尾の無音ぶんまで長さに数えてしまい、最後の一続きが
    # 「長すぎる」と判定されて連結されなくなる。
    chunk_len = float(chunk_seconds) if chunk_seconds else raw[-1]["rel"]
    for i, r in enumerate(raw):
        if i + 1 < len(raw):
            r["rel_end"] = raw[i + 1]["rel"]
        else:
            r["rel_end"] = min(chunk_len, r["rel"] + 2.0)
        if r["rel_end"] <= r["rel"]:
            r["rel_end"] = r["rel"] + 1.0

    if merge_same_speaker:
        # **関数の中身は触らない**(CLAUDE.md の注意点)。既にある引数で
        # 上限だけを変える。逐語では 20 秒、整文では従来どおり 180 秒。
        raw = merge_consecutive(
            raw,
            max_seconds=(VERBATIM_MERGE_MAX_SECONDS if verbatim
                         else MERGE_MAX_SECONDS))

    # 連結が済んでから、時刻を割り振り直す。
    # (詰まっている箇所を本文の長さで按分。区間どうしは重ならない)
    raw = redistribute_times(raw, chunk_len)

    # 末尾の音声がどの区間にも属さない状態を作らない
    if raw and chunk_len > raw[-1]["rel_end"]:
        raw[-1]["rel_end"] = chunk_len

    return [
        Utterance(rel_start=r["rel"], rel_end=r["rel_end"],
                  text=r["text"], cluster=r["cluster"])
        for r in raw
    ]


def parse_segments(
    text: str,
    chunk_index: int = 0,
    offset_seconds: float = 0.0,
    chunk_seconds: float | None = None,
    start_index: int = 0,
    merge_same_speaker: bool = True,
    verbatim: bool = False,
) -> list[dict]:
    """1 チャンクの出力テキストを発言区間の辞書リストに変換する。

    parse_utterances(経路ごと)＋ utterances_to_segments(共通の後段)の
    組み合わせに委ねている。戻り値の各要素は Segment(...) にそのまま渡せる
    キーを持つ(従来どおり)。
    """
    segments = utterances_to_segments(
        parse_utterances(text, chunk_seconds, merge_same_speaker, verbatim),
        chunk_index=chunk_index,
        offset_seconds=offset_seconds,
        start_index=start_index,
    )
    return [{
        "index": s.index,
        "start": s.start,
        "end": s.end,
        "text": s.text,
        "cluster": s.cluster,
        "chunk": s.chunk,
    } for s in segments]


# ======================================================================
# Word 出力
# ======================================================================

def _ensure_japanese_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "游明朝"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "游明朝")


def write_docx(
    transcripts: list[str],
    output_path: Path,
    title: str,
    note: str | None = None,
) -> None:
    """チャンクごとの文字起こし文字列リストを1つの .docx にまとめる。

    段落冒頭が [HH:MM:SS] や 【発言者X】 の形式なら、その部分を太字にする。
    note が指定されていれば、本文先頭に注意書き(斜体・灰色)を入れる。
    """
    doc = Document()
    _ensure_japanese_font(doc)
    doc.add_heading(title, level=1)

    if note:
        p = doc.add_paragraph()
        run = p.add_run(note)
        run.italic = True
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
        doc.add_paragraph()  # 空行で区切る

    for chunk_text in transcripts:
        for para in chunk_text.split("\n"):
            para = para.strip()
            if not para:
                continue
            m = _TS_LEAD_PATTERN.match(para)
            if m:
                ts, speaker, body = m.group(1), m.group(2), m.group(3)
                p = doc.add_paragraph()
                ts_run = p.add_run(ts + " ")
                ts_run.bold = True
                if speaker:
                    sp_run = p.add_run(speaker + " ")
                    sp_run.bold = True
                if body:
                    p.add_run(body)
            else:
                doc.add_paragraph(para)

    doc.save(output_path)
