"""セグメント(発言区間)のデータモデルと永続化・Word 出力。

v2.0.0 で追加。従来の「参加者名簿を渡して AI に話者を推定させる」方式に代わり、
    1) AI は声質だけで話者クラスタ(発言者A/B/C…)に分けた区間リストを作る
    2) ユーザーがタイムラインを区間単位でたどりながら、音声を聴いて話者を確定する
という流れに変更したため、その中間成果物を保持する構造が必要になった。

中間成果物は `<出力フォルダ>/<音声名>.speakers.json` に保存する。
アプリを閉じても、あとから同じファイルを開いて割当作業を再開できる。
"""
from __future__ import annotations

import json
import re
from dataclasses import dataclass, field, asdict, replace
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Optional, Sequence


SCHEMA_VERSION = 5

UNKNOWN_LABEL = "発言者不明"
MULTI_LABEL = "発言者複数・重複"

# 人が足した相づちを出力にどう書くか(設計書 §11)。データは 1 つのまま、
# 出し方だけを選ぶ。標準の表記法を調べた結果、BTSJ(基本的な文字化の原則・
# 宇佐美まゆみ 2019)にちょうどこの規定があったので、それに従う。
INSERT_STYLE_LINE = "line"       # 行を分ける。続く行は末尾 ,, でつなぐ(BTSJ 2.3.2)
INSERT_STYLE_INLINE = "inline"   # 行に埋め込む。(山本：はい) の形(BTSJ 3.2.3)
INSERT_STYLES = (INSERT_STYLE_LINE, INSERT_STYLE_INLINE)

# BTSJ 2.3.2 の「発話文が終わっていない」印。独自記号を作ると外部に説明できない
CONTINUE_MARK = ",,"

# 記号だけ出しても受け取った人が読めないので、差し込みがあるときだけ添える。
# 見出しを「凡例」にしない——検証要約にすでに同名の項目があり、1 つの文書に
# 同じ見出しが 2 つ並ぶ(検査を書いていて気付いた・2026-08-18)。
INSERT_LEGEND = {
    INSERT_STYLE_LINE:
        f"表記について: 発言の末尾の「{CONTINUE_MARK}」は、その発言が下の行に"
        "続いていることを示します(BTSJ 2.3.2 に準拠)。"
        "間の行は、そこに重なって入った別の人の発言です。",
    INSERT_STYLE_INLINE:
        "表記について: 本文中の「(氏名：…)」は、その位置に重なって入った"
        "別の人の発言です(BTSJ 3.2.3 に準拠)。"
        "氏名は冒頭の出席者一覧の略記です。",
}

# 特別扱いの話者 ID(名簿の人物ではない)
SPECIAL_UNKNOWN = "__unknown__"
SPECIAL_MULTI = "__multi__"
SPECIAL_NOISE = "__noise__"

SPECIAL_SPEAKERS: dict[str, str] = {
    SPECIAL_UNKNOWN: UNKNOWN_LABEL,
    SPECIAL_MULTI: MULTI_LABEL,
    SPECIAL_NOISE: "発言なし・雑音",
}

# 擬似クラスタの記号(transcribe.normalize_cluster_label の出力と対応)
PSEUDO_UNKNOWN = "?"
PSEUDO_MULTI = "*"

# 処理経路(Project.engine["mode"])。作業ファイルにそのまま入る値なので、
# 画面・Word の検証要約・パイプラインで同じ文字列を使う。
ENGINE_CLOUD = "cloud"
ENGINE_LOCAL = "local"
ENGINE_LABELS = {ENGINE_CLOUD: "クラウド", ENGINE_LOCAL: "ローカル"}

# 人が時刻を直したり区間を分けたりするときに許す最短の長さ。
# 0 にすると start == end の区間ができて、再生も出力も意味を失う。
MIN_SEGMENT_SECONDS = 0.1


def segment_key(seg: "Segment") -> tuple[float, float]:
    """区間を一意に指す鍵 —— **(orig_start, start) の組**。

    **`orig_start` だけでは足りない。**`split_segment` は「元は 1 つだった」と
    分かるよう、分割した両方に親の `orig_start` をそのまま与える（再実行の
    引き継ぎのためで、これ自体は正しい）。その結果 `orig_start` は一意でなく
    なり、実データで 2 組の重複が出た（2026-08-19 に実測。候補の一覧が両方に
    同じものを出し、片方で × を押すともう片方からも消えていた）。

    `start` は分割の境界で必ず変わるので、組にすれば区別できる。
    再実行の引き継ぎは従来どおり `orig_start` だけを見るので壊れない。

    **`index` は鍵にしてはいけない**——分割・結合・再実行で振り直る。
    """
    orig = seg.orig_start if seg.orig_start is not None else seg.start
    return (round(float(orig), 3), round(float(seg.start), 3))


def key_text(key: tuple[float, float]) -> str:
    """鍵を文字列に（sidecar に書くとき）。"""
    return f"{key[0]:.3f}+{key[1]:.3f}"


def added_signature(seg: "Segment") -> tuple[float, float, str]:
    """人が足した区間かを見分けるための身元 —— (orig_start, 終わり, まとまり)。

    `segment_key` とは別物。あちらは「どの区間か」を指す鍵で、こちらは
    「これは人が足したものか」を `edit_log` の記録と突き合わせるためのもの。

    **開始時刻だけでは足りない。**転写の区間と開始が一致することがある
    （`Project.added_utterance_signatures` を見よ）。
    """
    orig = seg.orig_start if seg.orig_start is not None else seg.start
    return (round(float(orig), 3), round(float(seg.end), 3),
            str(seg.cluster or ""))


def audio_span(seg: "Segment", time_offset: float) -> tuple[float, float]:
    """その区間が実音声のどこで鳴っているか(開始, 終了)。

    再生の規約: 実音声の位置 = 保存時刻 + ずれ補正。ただし一度直した区間の
    start/end は実音声の時刻そのものなので、補正を足さない。
    画面の再生・時刻編集の初期値・自動点検の照合窓が、すべてこの 1 つの
    規約を見るようにしてある(散らばると必ず食い違う)。
    """
    if seg.time_edited:
        return seg.start, seg.end
    return max(0.0, seg.start + time_offset), max(0.0, seg.end + time_offset)


def fmt_hms(seconds: float) -> str:
    """秒 → [HH:MM:SS] 用の 'HH:MM:SS' 文字列"""
    total = int(round(seconds))
    return f"{total // 3600:02d}:{(total % 3600) // 60:02d}:{total % 60:02d}"


def fmt_ms(seconds: float) -> str:
    """秒 → 'MM:SS'(1時間未満のときの簡易表示用)"""
    total = int(round(seconds))
    if total >= 3600:
        return fmt_hms(total)
    return f"{total // 60:02d}:{total % 60:02d}"


def fmt_hms_frac(seconds: float) -> str:
    """秒 → 'HH:MM:SS.s'(0.1 秒精度)。時刻を人が直接編集するときの表示形式。

    一覧の表示に使う fmt_hms() は 1 秒に丸めるが、区間の境目を耳で合わせる
    作業では 0.1 秒が要る。負の秒は 0 として扱う(時刻に負は無い)。
    """
    tenths = max(0, int(round(seconds * 10)))
    total, frac = divmod(tenths, 10)
    return f"{total // 3600:02d}:{(total % 3600) // 60:02d}:{total % 60:02d}.{frac}"


# 全角のまま打たれても受ける(日本語入力の途中で切り替え忘れが起きやすい)
_ZEN_TO_HAN = str.maketrans("０１２３４５６７８９：．", "0123456789:.")
_INT_FIELD = re.compile(r"^\d+$", re.ASCII)
_SEC_FIELD = re.compile(r"^\d+(?:\.\d*)?$", re.ASCII)


def parse_hms(text: str) -> float:
    """'HH:MM:SS.s' / 'MM:SS.s' / 'SS.s' のいずれかを秒に直す(0.1 秒精度)。

    時刻入力欄から受ける。読めない文字列は ValueError にして、
    呼び出し側が「編集前の値に戻す」判断をできるようにする。
    最上位の桁だけは 60 以上を許す('90' = 1分30秒、'75:00' = 75分)。
    """
    s = str(text).strip().translate(_ZEN_TO_HAN).replace(" ", "").replace("　", "")
    if not s:
        raise ValueError("時刻が空です。")
    parts = s.split(":")
    if len(parts) > 3:
        raise ValueError(f"時刻の形式が読めません: {text!r}")
    if not all(_INT_FIELD.match(p) for p in parts[:-1]) or not _SEC_FIELD.match(parts[-1]):
        raise ValueError(f"時刻の形式が読めません: {text!r}")
    values = [float(p) for p in parts]
    for v in values[1:]:
        if v >= 60.0:
            raise ValueError(f"分・秒は 60 未満で指定してください: {text!r}")
    total = 0.0
    for v in values:
        total = total * 60.0 + v
    return round(total, 1)


# ----------------------------------------------------------------------
# 話者
# ----------------------------------------------------------------------

@dataclass
class Speaker:
    """出席者(候補者リストに並ぶ 1 人)"""

    id: str
    name: str
    note: str = ""          # 役職・特徴などの補足
    order: int = 0          # 名簿上の並び順(初期の候補順に使う)

    @property
    def display(self) -> str:
        return f"{self.name}({self.note})" if self.note else self.name

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    @classmethod
    def from_dict(cls, d: dict[str, Any]) -> "Speaker":
        return cls(
            id=str(d["id"]),
            name=str(d.get("name", "")),
            note=str(d.get("note", "")),
            order=int(d.get("order", 0)),
        )


# 括弧は**最初の開きから最後の閉じまで**を取る。役職に括弧が入れ子で出る
# (「企画官（命）学校法人経営指導室長」)ため、`[^)）]*` だと途中で切れて
# 行全体が名前になってしまう(実データで発生・2026-08-18)。
_ROSTER_LINE = re.compile(
    r"^\s*(?:[-*・]\s*)?"          # 行頭の箇条書き記号は無視
    r"(?P<name>[^(（:：]+?)"        # 名前
    r"(?:[(（](?P<note1>.*)[)）])?"        # (役職) — 入れ子を許す
    r"\s*(?:[:：]\s*(?P<note2>.*))?$"      # : 補足
)

# 役職・所属を表す語。**2 文字以上のものだけ**を使う。「市」「村」「部」の
# ような 1 文字を入れると「田村」さんが「田」になる(実際に起きる)。
_ROLE_WORDS = re.compile(
    "衆議院|参議院|議員|知事|市長|町長|村長|区長|副会長|会長|理事長|理事|監事|"
    "代表取締役|取締役|代表|社長|専務|常務|部長|課長|係長|室長|局長|次長|所長|"
    "園長|校長|学長|教授|准教授|講師|教諭|秘書|専門官|参事官|調査官|企画官|"
    "事務官|技官|主査|主任|委員長|委員|顧問|相談役|幹事|事務局|同窓会|"
    "学園|学校|高校|中学校|小学校|大学|短大|会社|法人|協会|組合|事務所|"
    "センター|文科省|厚労省|経産省|国交省|財務省|総務省|防衛省"
)

# 氏名として残す最低の長さ。これを割ってまで削らない
_MIN_NAME_CHARS = 2


def suggest_split(line: str, others: Sequence[str] = ()) -> tuple[str, str]:
    """1 行の肩書つき氏名を (名前, 企業・役職) に分ける**提案**を返す。

    「三ツ林衆議院議員」→ ("三ツ林", "衆議院議員")
    「山本学　文科省 高等教育局…」→ ("山本学", "文科省 高等教育局…")

    規則は 2 つ。
      1. 空白があればそこで切る
      2. 役職・所属の語(_ROLE_WORDS)が出てきたら、その手前まで
      3. **他の人の行にも出てくる 3 字の並びは、個人名ではなく所属**
         (「加茂暁星」が別の人の行にも出るなら、その人の名前ではない)

    **当てにいきすぎない。**「山口京子蓮田市長」の「蓮田」は地名だが、
    規則で取ろうとすると本物の姓を削る事故が出る(「山田町長」の「山田」が
    姓なのか地名なのかは字面では決まらない)。**提案であって確定ではない**
    ので、外れたぶんは人が直す。設計書 §11.8。

    others には他の出席者の行(分ける前の全体)を渡す。
    """
    s = (line or "").strip()
    if not s:
        return "", ""
    # 1. 空白
    head = re.split(r"\s", s, maxsplit=1)[0].strip()
    cut = len(head) if head else len(s)
    # 2. 役職・所属の語
    m = _ROLE_WORDS.search(s[:cut])
    if m and m.start() >= _MIN_NAME_CHARS:
        cut = m.start()
    # 3. 他の人にも出てくる並び
    for i in range(_MIN_NAME_CHARS, max(_MIN_NAME_CHARS, cut - 2)):
        if any(s[i:i + 3] in o for o in others if o):
            cut = i
            break
    name = s[:cut].strip()
    note = s[cut:].strip()
    if len(name) < _MIN_NAME_CHARS:      # 削りすぎたら分けない
        return s, ""
    return name, note


def suggest_roster_rows(text: str) -> list[tuple[str, str]]:
    """名簿テキスト(1 行 1 人)を [(名前, 企業・役職), …] の提案にする。

    すでに「名前(役職)」の形で書かれている行はそのまま尊重し、
    分かれていない行だけ suggest_split() にかける。
    """
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    rows: list[tuple[str, str]] = []
    for sp in parse_roster(text):
        if sp.note:                       # すでに分かれている
            rows.append((sp.name, sp.note))
        else:
            others = [ln for ln in lines if sp.name not in ln]
            rows.append(suggest_split(sp.name, others))
    return rows


def parse_roster(text: str) -> list[Speaker]:
    """名簿テキスト(1行1人)を Speaker のリストにする。

    受け付ける形:
        佐藤
        佐藤(理事)
        佐藤(理事): 議長役。名乗ることが多い
        - 佐藤理事：会計担当
    """
    speakers: list[Speaker] = []
    for i, raw in enumerate(text.splitlines()):
        line = raw.strip()
        if not line:
            continue
        m = _ROSTER_LINE.match(line)
        if not m:
            name, note = line, ""
        else:
            name = (m.group("name") or "").strip()
            parts = [p.strip() for p in (m.group("note1"), m.group("note2")) if p and p.strip()]
            note = " / ".join(parts)
        if not name:
            continue
        speakers.append(Speaker(id=f"sp{i + 1:02d}", name=name, note=note, order=len(speakers)))
    return speakers


def roster_to_text(speakers: Iterable[Speaker]) -> str:
    lines = []
    for sp in speakers:
        lines.append(f"{sp.name}({sp.note})" if sp.note else sp.name)
    return "\n".join(lines)


# ----------------------------------------------------------------------
# セグメント
# ----------------------------------------------------------------------

@dataclass(frozen=True)
class Utterance:
    """転写エンジンが返す 1 発言。Segment を組み立てる前の中間の形。

    クラウド(Gemini)とローカル(faster-whisper)で、ここまでは同じ形に揃える。
    どちらの経路も「チャンク → Utterance の並び」を返し、オフセットの足し込みと
    通し番号の付与から先は共通の後段が引き受ける
    (claude/claude_ローカル転写_設計書.md §4.2)。

    時刻はチャンクの先頭からの相対秒。絶対秒にするのは後段の仕事。
    """

    rel_start: float
    rel_end: float
    text: str
    # 声のまとまりの記号。"A"/"B"… のほか、判別不能は "?"、複数人同時は "*"。
    # チャンク番号を頭に付けた "0:A" の形にするのは後段(Segment を作るとき)。
    cluster: str = PSEUDO_UNKNOWN


@dataclass
class Segment:
    """1 つの発言区間"""

    index: int
    start: float                    # 元音声内の絶対秒
    end: float
    text: str
    cluster: str                    # 例 "0:A"(チャンク0の発言者A)。声質ベースの仮ラベル
    chunk: int = 0
    speaker_id: Optional[str] = None    # 確定した話者(None = 未確定)
    reviewed: bool = False              # この区間の音声を実際に聴いて確定したか
    note: str = ""
    text_edited: bool = False           # ユーザーが本文を手直ししたか(再実行時に保護)
    time_edited: bool = False           # ユーザーが時刻を直したか
    # 時刻を「自分の耳で確かめた」か。機械が出した時刻を当てただけの区間と区別する。
    # 話者の reviewed と同じ考え方で、あとから未確認だけを拾い直せるようにする。
    time_reviewed: bool = False
    # パイプライン(AI)が出した元の時刻。start/end をユーザーが直しても動かさない。
    # 再実行したときに新旧の区間を突き合わせる鍵と、「元に戻す」の戻り先に使う。
    # None を渡すと start/end で埋める(新規生成時と、これを持たない旧ファイル)。
    orig_start: Optional[float] = None
    orig_end: Optional[float] = None

    def __post_init__(self) -> None:
        if self.orig_start is None:
            self.orig_start = self.start
        if self.orig_end is None:
            self.orig_end = self.end

    @property
    def duration(self) -> float:
        return max(0.0, self.end - self.start)

    @property
    def cluster_tail(self) -> str:
        """クラスタ記号だけ('0:A' → 'A')"""
        return self.cluster.partition(":")[2] or self.cluster

    @property
    def is_pseudo_cluster(self) -> bool:
        """『誰か判別できない』『複数人が重なっている』の擬似クラスタか。

        これらは「同じ声のまとまり」ではなく雑多な寄せ集めなので、
        学習にも一括適用にも使ってはいけない。
        """
        return self.cluster_tail in (PSEUDO_UNKNOWN, PSEUDO_MULTI)

    @property
    def cluster_label(self) -> str:
        """UI 表示用の短いクラスタ名。

        チャンク内で閉じたクラスタ(Gemini)は 'C1-A'。
        全長で分けたクラスタ(話者分離)は '声A' —— チャンク番号を出しても
        意味が無く、むしろ「C1-A と C2-A は別」という誤解を招く。
        """
        chunk, _, tail = self.cluster.partition(":")
        if chunk == "g":
            return f"声{tail}"
        try:
            return f"C{int(chunk) + 1}-{tail}"
        except ValueError:
            return self.cluster

    def preview(self, width: int = 60) -> str:
        t = " ".join(self.text.split())
        return t if len(t) <= width else t[: width - 1] + "…"

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    @classmethod
    def from_dict(cls, d: dict[str, Any]) -> "Segment":
        # orig_start / orig_end は schema 3 から。旧ファイルには無いので None を
        # 渡し、__post_init__ に start / end を入れさせる(移行処理は不要)。
        orig_start = d.get("orig_start")
        orig_end = d.get("orig_end")
        # schema 3 までは「時刻を直した = 自分の耳で合わせた」しかなかった。
        # 機械が出した時刻を当てただけの区間と区別する印は 4 で足したので、
        # 古いファイルの time_edited は確認済みとして読む。
        time_edited = bool(d.get("time_edited", False))
        return cls(
            index=int(d["index"]),
            start=float(d.get("start", 0.0)),
            end=float(d.get("end", 0.0)),
            text=str(d.get("text", "")),
            cluster=str(d.get("cluster", "")),
            chunk=int(d.get("chunk", 0)),
            speaker_id=d.get("speaker_id") or None,
            reviewed=bool(d.get("reviewed", False)),
            note=str(d.get("note", "")),
            text_edited=bool(d.get("text_edited", False)),
            time_edited=time_edited,
            time_reviewed=bool(d.get("time_reviewed", time_edited)),
            orig_start=float(orig_start) if orig_start is not None else None,
            orig_end=float(orig_end) if orig_end is not None else None,
        )


def utterances_to_segments(
    utterances: Iterable[Utterance],
    *,
    chunk_index: int = 0,
    offset_seconds: float = 0.0,
    start_index: int = 0,
) -> list[Segment]:
    """Utterance の並びを Segment に組み立てる(両経路の共通の後段)。

    クラウド(Gemini)もローカル(faster-whisper)も、チャンク単位で Utterance を
    作るところまでが経路ごとの仕事で、その先——チャンク先頭からの相対秒に
    オフセットを足して絶対秒にする / クラスタ記号にチャンク番号の名前空間を
    付ける / 通し番号を振る——は同じ。1 か所に集めておかないと、経路を足す
    たびに同じ処理が増える(設計書 §4.2)。

    ここで時刻をいじるのはオフセットの足し込みだけ。文字数按分
    (redistribute_times)は通さない。按分は Gemini のタイムスタンプが
    ドリフトする既知バグへの対策であって、実測時刻にかけるものではない。
    クラウド経路では parse_utterances の中で既に済ませてある。
    """
    out: list[Segment] = []
    for i, u in enumerate(utterances):
        start = offset_seconds + u.rel_start
        end = offset_seconds + u.rel_end
        # 長さ 0 の区間は聴き直せない(再生しても何も鳴らない)ので、
        # 最低限の長さを与えて操作できる状態にする。
        if end <= start:
            end = start + 1.0
        out.append(Segment(
            index=start_index + i,
            start=round(start, 2),
            end=round(end, 2),
            text=u.text,
            cluster=f"{chunk_index}:{u.cluster}",
            chunk=chunk_index,
        ))
    return out


# ----------------------------------------------------------------------
# 話者分離の結果を区間へ落とす
# (claude/claude_話者分離_設計書.md §4・§5・§8)
# ----------------------------------------------------------------------

# 区間の何割が話者区間と重なれば、その話者とみなすか。
# 自動クラスタリングが作る「1 区間だけの幽霊話者」は、この閾値でどの区間も
# 取れずに自然に消える(PoC で確認)。
MIN_SPEAKER_OVERLAP = 0.5

# 連結してよい間隔。これ以上空いていれば別の発言とみなす。
SPEAKER_MERGE_MAX_GAP = 0.3

# 全長の名前空間。チャンク内で閉じる "0:A" と区別する(設計書 §5)。
GLOBAL_NAMESPACE = "g"

_SENTENCE_ENDS = "。．！？!?」』…"


def _speaker_letters(turns: Iterable[Any]) -> dict[int, str]:
    """話者番号を A/B/C… に写す。**先に話した人から順**に振る。

    分離器が返す番号は連番とは限らない(実測で 0/1/3/4/7/22/33… のように飛ぶ)。
    番号をそのまま見せると意味の無い数字が並ぶので、出てきた順に文字を当てる。
    """
    order: list[int] = []
    for t in sorted(turns, key=lambda x: (x.start, x.end)):
        if t.speaker not in order:
            order.append(t.speaker)
    letters: dict[int, str] = {}
    for i, spk in enumerate(order):
        letters[spk] = chr(ord("A") + i) if i < 26 else f"S{i + 1}"
    return letters


def assign_speaker_clusters(
    segments: Sequence[Segment],
    turns: Sequence[Any],
    min_overlap: float = MIN_SPEAKER_OVERLAP,
) -> list[str]:
    """区間ごとの `cluster` 文字列を作る(設計書 §4・§5)。

    区間の時間範囲と**最も重なる話者区間**の話者を採る。重なりが区間の長さの
    `min_overlap` に満たなければ `?`(判別不能)にする。

    返すのは文字列の並びだけで、区間には書き込まない。書き込む場所を 1 つに
    保つため(呼び出し側で `seg.cluster = ...` する)。

    話者区間は重なりうるので、「最も重なる 1 つ」を選ぶ。同時に話している
    区間には、より長く重なっていたほうが入る。**それを `*` に直すのは人**
    ——重なりから機械的に判定しようとすると適合率 50% にしかならないことを
    実測した(設計書 §6)。
    """
    letters = _speaker_letters(turns)
    out: list[str] = []
    for seg in segments:
        best_spk, best_ov = None, 0.0
        for t in turns:
            ov = min(seg.end, t.end) - max(seg.start, t.start)
            if ov > best_ov:
                best_spk, best_ov = t.speaker, ov
        span = max(0.01, seg.end - seg.start)
        if best_spk is None or best_ov / span < min_overlap:
            out.append(f"{GLOBAL_NAMESPACE}:{PSEUDO_UNKNOWN}")
        else:
            out.append(f"{GLOBAL_NAMESPACE}:{letters[best_spk]}")
    return out


def merge_same_speaker(
    segments: Sequence[Segment],
    max_gap: float = SPEAKER_MERGE_MAX_GAP,
) -> list[Segment]:
    """同じ話者の、文の途中で切れた区間を連結する(設計書 §8)。

    ローカル転写は 74% の区間が文の途中で終わる(実測)。話者ラベルが付けば
    安全に連結できる。条件は 3 つとも満たすこと:

      1. 同じ話者(擬似クラスタ `?` `*` は連結しない)
      2. 前の区間が句点等で終わっていない(文が続いている)
      3. 間隔が max_gap 未満

    **人が手を付けた区間は連結しない。**割当・本文の手直し・時刻の修正が
    入っているものを勝手にまとめると、その作業が消える。

    `orig_start` / `orig_end` は前後の端を保つ。再実行時に新旧の区間を
    突き合わせる鍵なので、連結で失うと引き継ぎが壊れる。
    """
    out: list[Segment] = []
    for seg in segments:
        if out:
            prev = out[-1]
            touched = any((
                prev.speaker_id, seg.speaker_id,
                prev.text_edited, seg.text_edited,
                prev.time_edited, seg.time_edited,
            ))
            joinable = (
                not touched
                and prev.cluster == seg.cluster
                and not prev.is_pseudo_cluster
                and prev.text
                and prev.text[-1] not in _SENTENCE_ENDS
                and seg.start - prev.end < max_gap
            )
            if joinable:
                prev.text = prev.text + seg.text
                prev.end = seg.end
                prev.orig_end = seg.orig_end
                continue
        out.append(seg)
    for i, seg in enumerate(out):
        seg.index = i
    return out


# 語句を探すときに前後を何文字見せるか。**文脈が無いと判断できない。**
# 「資格」は 10 回出るが、うち 1 回は本物の資格(防災士の資格)で、直しては
# いけない。時刻と語だけを並べても○×は付けられない(設計書 §16.3)。
CONTEXT_CHARS = 22


@dataclass(frozen=True)
class TextHit:
    """本文に語句が出てくる 1 箇所。**判定は出さない。**

    `listen_order` / `candidates` と同じ線で、これは「見る場所を並べる」道具。
    直すかどうかは人が前後を読んで決める(設計書 §16.3)。
    """

    key: tuple[float, float]    # 区間を指す鍵(segment_key)
    nth: int                    # その区間の中で何番目に出たか(0 から)
    at: float                   # 区間の開始秒(表示用)
    index: int                  # 区間番号。**表示だけ。**鍵には使わない
    before: str                 # 語句の手前(CONTEXT_CHARS 文字まで)
    term: str                   # 語句そのもの
    after: str                  # 語句の後ろ(CONTEXT_CHARS 文字まで)
    head: bool                  # 手前が切り詰められているか(… を出すため)
    tail: bool                  # 後ろが切り詰められているか

    @property
    def target(self) -> tuple[tuple[float, float], int]:
        """`replace_text` に渡す指定。"""
        return (self.key, self.nth)


# ----------------------------------------------------------------------
# プロジェクト(1 音声ファイル分の作業状態)
# ----------------------------------------------------------------------

@dataclass
class Project:
    audio_path: str
    duration: float = 0.0
    chunk_seconds: int = 600
    model: str = ""
    verbatim: bool = False
    # 音声の中身から作った指紋。ファイル名が同じでも中身が変われば別物として
    # 扱うために使う(空文字なら未記録 = 古い形式の作業ファイル)。
    audio_fingerprint: str = ""
    # 再生位置のずれ補正(秒)。Gemini の時刻推定が実音声より早い/遅いときに
    # 使う。録音ごとに傾向が違うので、設定ではなく作業ファイルに持たせる。
    time_offset: float = 0.0
    # ---- ここから v5(検証履歴)。追加は文書レベルのみで、区間の形は変えない ----
    # 元音声の SHA-256。「この書面はこの録音から作った」を第三者が
    # Get-FileHash / certutil / sha256sum で検算するための値。
    # 指紋(audio_fingerprint)とは役割が違う: あちらはキャッシュの同一性判定。
    source_sha256: str = ""
    # 処理経路の記録(自由形式)。慣例のキー: mode("cloud"/"local")・model・
    # app_version・at(UTC の ISO8601)。将来のモデル出所記録(Model BOM)にも
    # ここを拡張して使う。
    engine: dict[str, Any] = field(default_factory=dict)
    # Word を出力するたびに +1 し、書面に「版」として併記する。
    # ファイルを開いただけでは進めない(開いた事実は版ではない)。
    doc_revision: int = 0
    # 追記型の編集履歴。慣例のキー: at(UTC)・actor("user"/"inspect")・
    # kind(time/text/speaker/…)・target(orig_start)・before/after・batch_id。
    # v5 では器だけを定義し、記録の書き込みは編集履歴の実装(Day 45)で行う。
    edit_log: list[dict[str, Any]] = field(default_factory=list)
    speakers: list[Speaker] = field(default_factory=list)
    segments: list[Segment] = field(default_factory=list)
    json_path: Optional[str] = None      # 保存先(load 時に設定)

    # -------------------------------------------------- 話者アクセス
    def speaker(self, speaker_id: Optional[str]) -> Optional[Speaker]:
        if not speaker_id:
            return None
        if speaker_id in SPECIAL_SPEAKERS:
            return Speaker(id=speaker_id, name=SPECIAL_SPEAKERS[speaker_id], order=999)
        for sp in self.speakers:
            if sp.id == speaker_id:
                return sp
        return None

    def speaker_name(self, speaker_id: Optional[str]) -> str:
        sp = self.speaker(speaker_id)
        return sp.name if sp else ""

    def add_speaker(self, name: str, note: str = "") -> Speaker:
        existing = {sp.id for sp in self.speakers}
        i = len(self.speakers) + 1
        while f"sp{i:02d}" in existing:
            i += 1
        sp = Speaker(id=f"sp{i:02d}", name=name, note=note, order=len(self.speakers))
        self.speakers.append(sp)
        return sp

    def remove_speaker(self, speaker_id: str) -> None:
        self.speakers = [sp for sp in self.speakers if sp.id != speaker_id]
        for seg in self.segments:
            if seg.speaker_id == speaker_id:
                seg.speaker_id = None
                seg.reviewed = False
        for i, sp in enumerate(self.speakers):
            sp.order = i

    # -------------------------------------------------- 区間の分割・結合
    def renumber(self) -> None:
        """index を 0..n-1 に振り直す。区間を増減させたら必ず呼ぶ。"""
        for i, seg in enumerate(self.segments):
            seg.index = i

    def split_segment(self, index: int, boundary: float, cut: int) -> tuple[Segment, Segment]:
        """1 つの区間を、境界時刻と本文の位置で 2 つに分ける。

        1 区間に 2 人の発言が(重なりなしで)順に混ざっているときに使う。
        転写から丸ごと落ちた発言も、隣の区間を分割して空いた側に本文を
        書き足せば復元できる。

        boundary: 境界の秒(実音声の時刻)。区間の内側に収める。
        cut:      本文を切る位置(文字数。ダイアログのカーソル位置)。

        後半のクラスタを擬似不明にするのは、元の声質ラベルが区間全体に付いた
        もので、後半の声には保証がないため。擬似クラスタは候補学習からも一括
        適用からも自動的に外れる(Segment.is_pseudo_cluster)ので、分割で
        生まれた不確かな区間が学習を汚したり誤って伝播したりするのを防げる。
        """
        seg = self.segments[index]
        if seg.duration < MIN_SEGMENT_SECONDS * 2:
            raise ValueError("この区間は短すぎて分割できません。")
        lo = seg.start + MIN_SEGMENT_SECONDS
        hi = seg.end - MIN_SEGMENT_SECONDS
        boundary = min(max(round(float(boundary), 1), lo), hi)
        cut = max(0, min(int(cut), len(seg.text)))

        head = Segment(
            index=index,
            start=seg.start,
            end=boundary,
            text=seg.text[:cut],
            cluster=seg.cluster,            # 前半は元の声のまとまりのまま
            chunk=seg.chunk,
            speaker_id=seg.speaker_id,
            reviewed=False,                 # 範囲が変わったので聴き直し対象
            note=seg.note,
            text_edited=seg.text_edited,
            time_edited=True,
            # 境界は人が決めるが、外側の端は元の区間のまま。親が未確認なら
            # 子も未確認にする(分割したというだけで確認済みには昇格させない)。
            time_reviewed=seg.time_reviewed,
            # 再実行時に「元は 1 つだった」と分かるよう、親の値を両方が共有する
            orig_start=seg.orig_start,
            orig_end=seg.orig_end,
        )
        tail = Segment(
            index=index + 1,
            start=boundary,
            end=seg.end,
            text=seg.text[cut:],
            cluster=f"{seg.chunk}:{PSEUDO_UNKNOWN}",
            chunk=seg.chunk,
            speaker_id=None,                # 後半の声は別人かもしれない
            reviewed=False,
            note="",
            text_edited=seg.text_edited,
            time_edited=True,
            time_reviewed=seg.time_reviewed,
            orig_start=seg.orig_start,
            orig_end=seg.orig_end,
        )
        self.segments[index:index + 1] = [head, tail]
        self.renumber()
        return head, tail

    # -------------------------------------------------- 相づちを足す
    def added_utterance_keys(self) -> set[float]:
        """人が足した区間の orig_start（丸め済み）。

        **再実行の引き継ぎがこれを見て、突き合わせから外す。**外さないと、
        足した区間の独自の時刻が近くの無関係な区間に誤って照合され、
        その区間を置き換えて消す（設計書 §4）。

        識別に区間のフラグを使わないのは、スキーマを増やさないため
        （v3 の一括移行まで待つ）。edit_log は既にあり、再実行でも
        引き継がれるので、ここに置くのが自然。
        """
        keys: set[float] = set()
        for rec in self.edit_log:
            op = rec.get("op")
            k = rec.get("orig_start")
            if k is None:
                continue
            if op == "add_utterance":
                keys.add(round(float(k), 3))
            elif op == "remove_added_utterance":
                keys.discard(round(float(k), 3))
        return keys

    def added_utterance_signatures(self) -> set[tuple[float, float, str]]:
        """人が足した区間の身元 `(orig_start, 終わり, まとまり)`。

        **開始時刻だけでは身元にならない。**転写の区間とたまたま開始が
        一致することがあり、実データで転写の区間が「人が足した」と誤判定
        されていた（42:15.9 で発生・2026-08-20）。誤判定された区間には
        ［この区間を消す］が押せてしまう。これは
        「音声認識が出した区間には削除の入口を作らない」（CLAUDE.md）に反する。

        **消した記録は見ない。**消えた区間はもう存在しないので、記録から
        引く必要がない。引くと、過去に消し損ねて残った区間まで
        「足したものではない」ことになり、**二度と消せなくなる**
        （実データで 42:12.7 に残骸が 1 件できていた）。
        """
        out: set[tuple[float, float, str]] = set()
        for rec in self.edit_log:
            if rec.get("op") != "add_utterance":
                continue
            k, e = rec.get("orig_start"), rec.get("end")
            if k is None or e is None:
                continue
            out.add((round(float(k), 3), round(float(e), 3),
                     str(rec.get("cluster") or "")))
        return out

    def is_added_utterance(self, seg: Segment) -> bool:
        """この区間は人が足したものか（消してよいか）。"""
        return added_signature(seg) in self.added_utterance_signatures()

    def add_utterance(self, start: float, end: float, text: str,
                      cluster: str = "", cut: Optional[int] = None,
                      parent_orig: Optional[float] = None,
                      parent_start: Optional[float] = None) -> Segment:
        """聞こえたのに本文に無い発話を、区間として足す（設計書 §2）。

        時刻と声のまとまりは機械（話者分離の turn）が用意し、本文は人が打つ。
        **重なりを禁止しない。**相づちは主発言と重なるのが本性なので、
        既存区間と時間的に重なってよい。

        話者は付けずに返す。付けるのは呼び出し側の通常の割当操作で、
        そのときの ✓/△ は既存の意味論に従う（機械が ✓ を立てる経路は無い）。
        """
        start = round(float(start), 3)
        end = round(max(float(end), start + MIN_SEGMENT_SECONDS), 3)
        text = (text or "").strip()
        if not text:
            raise ValueError("本文が空です。")
        # 前後の区間と同じチャンクに属させる（チャンク番号は再生や
        # クラスタ記号の表示に使われる）
        near = min(self.segments, key=lambda s: abs(s.start - start),
                   default=None)
        seg = Segment(
            index=0,                        # renumber で振り直す
            start=start,
            end=end,
            text=text,
            cluster=cluster or f"{near.chunk if near else 0}:{PSEUDO_UNKNOWN}",
            chunk=near.chunk if near else 0,
            speaker_id=None,
            reviewed=False,
            text_edited=True,               # 人が打った本文
            time_edited=False,              # 時刻は turn 由来の機械値
        )
        # **同じ時刻に既にあるものの「後ろ」に入れる。**前に入れていたため
        # 同時刻に 2 件足すと並びが作った順と逆になり、割り込み位置(cut)と
        # 本文の対応が入れ替わっていた(2026-08-20)。
        pos = len([s for s in self.segments if s.start <= start])
        self.segments.insert(pos, seg)
        self.renumber()
        self.edit_log.append({
            "op": "add_utterance",
            "at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
            "actor": "user",
            "orig_start": float(seg.orig_start),
            "start": start,
            "end": end,
            "cluster": seg.cluster,
            # **どの区間の本文の、どこに割り込んだか。**
            # 区間は割らない(割ると直すときに元の本文を復元できない)。
            # 代わりにここを覚えておき、Word に出すときだけ差し込む。
            "cut": None if cut is None else int(cut),
            # **差し込み先は (orig_start, start) の組で指す。**orig_start だけ
            # では分割した 2 つを区別できない（segment_key を見よ）。
            # parent_start が無い古い記録は orig_start だけで突き合わせる。
            "parent_orig": None if parent_orig is None else float(parent_orig),
            "parent_start": None if parent_start is None else float(parent_start),
        })
        return seg

    # ------------------------------------------------------------------
    # 編集履歴（設計書 §12）
    #
    # **記録するのは「人がした判断」であって「区間の書き換え」ではない。**
    # 一括適用は 50 区間を変えても人の判断は 1 回なので、記録も 1 件で
    # 対象を列挙する。区間ごとに 1 件ずつ残すと、実データで 95KB になり
    # (実測 2026-08-19)、しかも「何回判断したか」が読み取れなくなる。
    #
    # **before を残す。**「機械が何と言い、人が何に直したか」が無ければ
    # 検証履歴の意味がない。本文を丸ごと持っても、実データで作業ファイルは
    # 1.35 倍（280KB→379KB）、起こりうる最大でも 2.4 倍で収まる。
    #
    # **書き換えの経路をここに集める。**画面が seg.speaker_id を直接
    # 書き換えると記録を素通りするので、必ずこのメソッドを通す。
    # ------------------------------------------------------------------
    def _log(self, op: str, **kw: Any) -> None:
        self.edit_log.append({
            "op": op,
            "at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
            "actor": "user",
            **kw,
        })

    def _key(self, seg: Segment) -> list[float]:
        """記録が指す鍵 —— **(orig_start, start) の組**。

        `index` は振り直るので使わない。**`orig_start` だけでも足りない**
        ——分割した 2 つが同じ値を持つため（`segment_key` を見よ）。
        JSON に入れるので list で返す。
        """
        return list(segment_key(seg))

    def assign_speaker(self, index: int, speaker_id: Optional[str],
                       reviewed: bool = True) -> None:
        """1 区間に話者を割り当てる（人が聴いて決めた 1 回の判断）。

        speaker_id=None なら未確定に戻す（reviewed も落ちる）。
        """
        seg = self.segments[index]
        before, before_rev = seg.speaker_id, seg.reviewed
        seg.speaker_id = speaker_id
        seg.reviewed = bool(speaker_id) and reviewed
        if (before, before_rev) == (seg.speaker_id, seg.reviewed):
            return                      # 変わっていないなら記録しない
        self._log("assign", segment=self._key(seg),
                  before=before, after=speaker_id, reviewed=seg.reviewed)

    def apply_speaker_to(self, indexes: Sequence[int],
                         speaker_id: Optional[str],
                         heard_index: Optional[int] = None) -> None:
        """まとめて適用（1 回の判断で複数区間）。

        heard_index の区間だけ `reviewed=True`（その 1 つを聴いて決めた）。
        残りは `reviewed=False`——**機械の結果をまとめて当てただけ**で、
        個別には聴いていない。この区別が製品価値そのもの（CLAUDE.md）。
        """
        heard: list[list[float]] = []
        bulk: list[list[float]] = []
        for i in indexes:
            seg = self.segments[i]
            was = (seg.speaker_id, seg.reviewed)
            seg.speaker_id = speaker_id
            seg.reviewed = bool(speaker_id) and (i == heard_index)
            if was == (seg.speaker_id, seg.reviewed):
                continue
            (heard if i == heard_index else bulk).append(self._key(seg))
        if not (heard or bulk):
            return
        self._log("assign_bulk", after=speaker_id,
                  heard=heard, bulk=bulk, count=len(heard) + len(bulk))

    def replace_speaker(self, before_id: Optional[str], after_id: Optional[str],
                        keys: Sequence[tuple[float, float]]) -> int:
        """選ばれた区間の話者を、別の人に付け替える。直した数を返す。

        使いどころは**途中退席**——「32:17 に吉沢さんが帰ったので、それ以降の
        吉沢忠一は全部西村香介」。実データで 108 区間あった（2026-08-20）。

        **すべて △（まとめて適用）になる。**元が ✓ でも △ に落ちる。
        人が確かめたのは「その人はもう居なかった」という事実であって、
        区間ごとの声ではない。✓ のまま残すと「1 区間ずつ聴いて確かめた」
        という意味になり、`reviewed` の意味論が壊れる（CLAUDE.md）。

        **記録は 1 件。**人の判断は「吉沢は 32:17 に帰った」の 1 回。
        `before` も残す（何を何に直したかが無ければ検証履歴にならない
        ——編集履歴設計書 §1.2）。
        """
        if before_id == after_id:
            return 0
        want = {(round(float(k[0]), 3), round(float(k[1]), 3)) for k in keys}
        if not want:
            return 0
        targets: list[list[float]] = []
        undone: list[list[float]] = []      # ✓ から △ に落ちたもの
        for seg in self.segments:
            if segment_key(seg) not in want or seg.speaker_id != before_id:
                continue
            if seg.reviewed:
                undone.append(self._key(seg))
            seg.speaker_id = after_id
            seg.reviewed = False            # 機械が ✓ を立てる経路は作らない
            targets.append(self._key(seg))
        if not targets:
            return 0
        self._log("replace_speaker_bulk", before=before_id, after=after_id,
                  targets=targets, count=len(targets),
                  was_reviewed=undone)
        return len(targets)

    def carry_speakers(
        self,
        assignments: Sequence[tuple[tuple[float, float], Optional[str]]],
        source: str = "",
    ) -> int:
        """別の転写から話者の割当を写す。写した数を返す。

        使いどころは**転写し直したとき**。モデルを変えると区間の切れ目が
        変わるので、`_carry_over_assignments`（`orig_start` で突き合わせる）
        では引き継げない。時間の重なりで対応づけて写す。

        **すべて △（まとめて適用）になる。**元が ✓ でも △ に落ちる。
        区間の切れ目が変わっている以上、「その区間を聴いて確定した」とは
        言えない。✓ のまま写すと `reviewed` の意味論が壊れる（CLAUDE.md）。

        **記録は 1 件。**どこから何件写したかを残す——あとから
        「これは前の転写から写したものだ」と読めるようにするため。
        """
        want = {(round(float(k[0]), 3), round(float(k[1]), 3)): sid
                for k, sid in assignments if sid}
        if not want:
            return 0
        done: list[list[float]] = []
        for seg in self.segments:
            sid = want.get(segment_key(seg))
            if not sid or seg.speaker_id == sid:
                continue
            seg.speaker_id = sid
            seg.reviewed = False        # 機械が ✓ を立てる経路は作らない
            done.append(self._key(seg))
        if not done:
            return 0
        self._log("carry_speakers", targets=done, count=len(done),
                  source=source)
        return len(done)

    def restore_assignments(
        self, snapshot: Sequence[tuple[int, Optional[str], bool]]
    ) -> None:
        """取り消し。**取り消したことも履歴に残す。**

        消してしまうと「一度は当てたが戻した」経緯が読めなくなる。
        検証履歴は結果ではなく経緯の記録なので、打ち消しも記録に残す。
        """
        keys: list[list[float]] = []
        for index, sid, reviewed in snapshot:
            if not (0 <= index < len(self.segments)):
                continue
            seg = self.segments[index]
            if (seg.speaker_id, seg.reviewed) == (sid, reviewed):
                continue
            seg.speaker_id, seg.reviewed = sid, reviewed
            keys.append(self._key(seg))
        if keys:
            self._log("undo_assign", targets=keys, count=len(keys))

    def edit_text(self, index: int, new_text: str) -> bool:
        """本文を人が直す。変わったときだけ記録して True を返す。"""
        seg = self.segments[index]
        new_text = new_text if new_text is not None else ""
        if new_text == seg.text:
            return False
        before = seg.text
        seg.text = new_text
        seg.text_edited = True          # 再実行で上書きされないよう印を付ける
        self._log("edit_text", segment=self._key(seg),
                  before=before, after=new_text)
        return True

    def find_text(self, term: str) -> list[TextHit]:
        """本文に `term` が出てくる箇所を、前後の文脈つきで全部並べる。

        **1 回の出現につき 1 件返す。**同じ区間に 2 回出ることがある
        (実データ「同層会ですね同層会が」)。区間単位にまとめると、片方だけ
        直したい場合に手が出せない。
        """
        term = term or ""
        if not term:
            return []
        out: list[TextHit] = []
        for seg in self.segments:
            text = seg.text or ""
            pos = text.find(term)
            nth = 0
            while pos >= 0:
                lo = max(0, pos - CONTEXT_CHARS)
                hi = min(len(text), pos + len(term) + CONTEXT_CHARS)
                out.append(TextHit(
                    key=segment_key(seg), nth=nth, at=float(seg.start),
                    index=seg.index,
                    before=text[lo:pos], term=term,
                    after=text[pos + len(term):hi],
                    head=lo > 0, tail=hi < len(text)))
                nth += 1
                pos = text.find(term, pos + len(term))
        return out

    def replace_text(self, before: str, after: str,
                     targets: Sequence[tuple[tuple[float, float], int]]) -> int:
        """選ばれた箇所だけ語句を置き換える。直した箇所数を返す。

        **一括で置き換えない。**「資格」は 10 回出るが 1 回は本物なので、
        全部を機械的に直すと本文が壊れる(設計書 §16.3)。どこを直すかは
        呼び出し側(画面)で人が選ぶ。

        **記録は 1 件。**人の判断は「吉田は吉沢の聞き違いだ」の 1 回であって、
        区間の数だけ判断したわけではない(編集履歴設計書 §1.1)。

        **`reviewed` の ✓ は立てない。**人が音声を聴いて確かめたわけでは
        ないため(CLAUDE.md)。`text_edited` は立てる——人が直した本文なので、
        再実行で機械の出力に戻されては困る。
        """
        before = before or ""
        if not before or before == after:
            return 0
        by_key: dict[tuple[float, float], set[int]] = {}
        for key, nth in targets:
            by_key.setdefault((round(float(key[0]), 3),
                               round(float(key[1]), 3)), set()).add(int(nth))
        done = 0
        touched: list[list[float]] = []
        for seg in self.segments:
            want = by_key.get(segment_key(seg))
            if not want:
                continue
            text, out, pos, nth = seg.text or "", [], 0, 0
            at = text.find(before)
            while at >= 0:
                out.append(text[pos:at])
                out.append(after if nth in want else before)
                done += nth in want
                pos = at + len(before)
                nth += 1
                at = text.find(before, pos)
            out.append(text[pos:])
            new_text = "".join(out)
            if new_text == seg.text:
                continue                    # 指定が古い(本文が変わっていた)
            seg.text = new_text
            seg.text_edited = True
            touched.append(self._key(seg))
        if not done:
            return 0
        self._log("replace_text_bulk", before=before, after=after,
                  targets=touched, count=done, segments=len(touched))
        return done

    def edit_time(self, index: int, start: float, end: float,
                  reviewed: bool, _log: bool = True) -> bool:
        """区間の時刻を人が直す／確かめる。変わったときだけ記録。

        値が同じでも `reviewed` が上がる（✎△ → ✎）ことがあるので、
        そこも変化として記録する。**値の書き込みもここで行う**——
        画面側で書くと記録を素通りする経路が生まれる（設計書 §1.3）。

        _log=False は「まとめて適用」用。呼び出し側が 1 件の記録に
        まとめるので、ここでは書き込みだけ行う。
        """
        seg = self.segments[index]
        was = (seg.start, seg.end, seg.time_edited, seg.time_reviewed)
        changed = was != (start, end, True, bool(reviewed))
        seg.start, seg.end = start, end
        seg.time_edited = True          # 以後この区間にずれ補正を足さない
        seg.time_reviewed = bool(reviewed)
        if changed and _log:
            self._log("edit_time", segment=self._key(seg),
                      before=[round(was[0], 3), round(was[1], 3)],
                      after=[round(start, 3), round(end, 3)],
                      reviewed=bool(reviewed))
        return changed

    def apply_times_to(
        self, items: Sequence[tuple[int, float, float]], reviewed: bool = False
    ) -> int:
        """時刻をまとめて当てる（1 回の判断で複数区間・記録は 1 件）。

        点検の提案の一括適用。**すべて ✎△**（機械が当てただけ）。
        人の耳の確認はあとから 1 件ずつ ✎ に上げる。
        """
        keys: list[list[float]] = []
        for index, start, end in items:
            if self.edit_time(index, start, end, reviewed, _log=False):
                keys.append(self._key(self.segments[index]))
        if keys:
            self._log("apply_times_bulk", targets=keys, count=len(keys),
                      reviewed=bool(reviewed))
        return len(keys)

    def revert_time(self, index: int) -> bool:
        """パイプラインが出した元の時刻に戻す（以後はまたずれ補正が効く）。"""
        seg = self.segments[index]
        if not seg.time_edited:
            return False
        before = [round(seg.start, 3), round(seg.end, 3)]
        seg.start = float(seg.orig_start)
        seg.end = float(seg.orig_end)
        seg.time_edited = False
        seg.time_reviewed = False
        self._log("revert_time", segment=self._key(seg),
                  before=before, after=[round(seg.start, 3), round(seg.end, 3)])
        return True

    def restore_times(
        self, items: Sequence[tuple[int, float, float, bool, bool]]
    ) -> int:
        """時刻の一括適用を丸ごと元に戻す。**戻したことも記録に残す。**"""
        keys: list[list[float]] = []
        for index, start, end, edited, reviewed in items:
            if not (0 <= index < len(self.segments)):
                continue
            seg = self.segments[index]
            if (seg.start, seg.end, seg.time_edited, seg.time_reviewed) == (
                    start, end, edited, reviewed):
                continue
            seg.start, seg.end = start, end
            seg.time_edited, seg.time_reviewed = edited, reviewed
            keys.append(self._key(seg))
        if keys:
            self._log("undo_times", targets=keys, count=len(keys))
        return len(keys)

    def clear_speakers(self, speaker_ids: Sequence[str]) -> int:
        """名簿から消えた人の割当を外す。**外したことを記録に残す。**

        「一度はこの人に当てていたが、名簿から消したので外れた」という
        経緯は、あとから読む人にとって重要（誰の発言か分からなくなった
        理由がこれ）。
        """
        gone = set(speaker_ids)
        keys: list[list[float]] = []
        for seg in self.segments:
            if seg.speaker_id in gone:
                seg.speaker_id = None
                seg.reviewed = False
                keys.append(self._key(seg))
        if keys:
            self._log("clear_speakers", removed=sorted(gone),
                      targets=keys, count=len(keys))
        return len(keys)

    def set_added_speaker(self, index: int, speaker_id: str) -> None:
        """足した発話に、その場で選んだ話者を入れる。

        **いま聴いた直後に人が選んだので ✓。**機械が立てる経路ではない。
        """
        seg = self.segments[index]
        seg.speaker_id = speaker_id
        seg.reviewed = True
        self._log("assign", segment=self._key(seg),
                  before=None, after=speaker_id, reviewed=True)

    def log_counts(self) -> dict[str, int]:
        """検証要約に出す内訳。**op ごとの件数だけ**（明細は Day 75）。"""
        out: dict[str, int] = {}
        for rec in self.edit_log:
            op = str(rec.get("op") or "?")
            out[op] = out.get(op, 0) + 1
        return out

    def log_last_at(self) -> str:
        """最後に手を入れた時刻（ISO・UTC）。無ければ空。"""
        times = [str(r.get("at") or "") for r in self.edit_log]
        return max(times) if times else ""

    def remove_added_utterance(self, index: int) -> None:
        """区間を消す。**人が明示的に消すときだけ。**

        短い相づちの**自動**削除をしない原則（CLAUDE.md）はそのまま。
        禁じているのは自動処理で、**音声を聴いて機械の重複だと判断した人が
        消すことは別のこと**。実機で「同じ音声が 3 区間に書かれた」場面に
        当たり、結合を 2 回してから分割する遠回りを強いられた
        （2026-08-22）。

        転写が出した区間を消したときは、**消した中身を編集履歴に残す**。
        第三者が「ここで何が消されたか」を追えなければ、検証済みの記録に
        ならない。
        """
        if not (0 <= index < len(self.segments)):
            raise ValueError("その区間はありません。")
        seg = self.segments[index]
        added = self.is_added_utterance(seg)
        del self.segments[index]
        self.renumber()
        self.edit_log.append({
            "op": "remove_added_utterance" if added else "remove_segment",
            "at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
            "actor": "user",
            "orig_start": float(seg.orig_start),
            # **転写が出した区間を消したときは、中身も残す。**
            # 足した区間を消すのは「自分の入力の取り消し」だが、こちらは
            # 機械の出力を人の判断で落とすこと。第三者が「ここで何が
            # 消されたか」を追えなければ、検証済みの記録にならない。
            **({} if added else {
                "start": float(seg.start),
                "end": float(seg.end),
                "text": seg.text,
                "cluster": str(seg.cluster),
            }),
        })

    def merge_segments(self, index: int) -> Segment:
        """index の区間と、その次の区間を 1 つにまとめる。

        分割のやり直し(逆操作)と、同じ発言が 2 行に割れているときの整理に使う。
        話者が食い違うときは、どちらが正しいか機械には決められないので未確定に
        落とす。呼び出し側はその旨をユーザーに確認してから呼ぶこと。
        """
        if index < 0 or index + 1 >= len(self.segments):
            raise ValueError("次の区間がないので結合できません。")
        a, b = self.segments[index], self.segments[index + 1]
        notes = [n for n in (a.note, b.note) if n]
        merged = Segment(
            index=index,
            start=min(a.start, b.start),
            end=max(a.end, b.end),
            text=a.text + b.text,           # 日本語なので空白を挟まない
            cluster=a.cluster,              # 前側の声のまとまりを採用する
            chunk=a.chunk,
            speaker_id=a.speaker_id if a.speaker_id == b.speaker_id else None,
            reviewed=False,                 # 範囲が変わったので聴き直し対象
            note=" / ".join(notes),
            text_edited=a.text_edited or b.text_edited,
            time_edited=True,
            # 両端とも耳で確かめてあったときだけ確認済みのまま(片方が未確認なら未確認)
            time_reviewed=a.time_reviewed and b.time_reviewed,
            orig_start=a.orig_start,        # 系譜の始まりは前側
            orig_end=b.orig_end,            # 終わりは後側(再実行時の吸収に要る)
        )
        self.segments[index:index + 2] = [merged]
        self.renumber()
        return merged

    # -------------------------------------------------- 統計
    @property
    def assigned_count(self) -> int:
        return sum(1 for s in self.segments if s.speaker_id)

    @property
    def reviewed_count(self) -> int:
        """実際に音声を聴いて確定した区間の数。

        一括適用で埋めた区間は「確定はしているが未確認」なので、ここには入らない。
        あとから未確認だけを拾い直せるようにするための区別。
        """
        return sum(1 for s in self.segments if s.speaker_id and s.reviewed)

    @property
    def unreviewed_count(self) -> int:
        return sum(1 for s in self.segments if s.speaker_id and not s.reviewed)

    @property
    def total_count(self) -> int:
        return len(self.segments)

    def clusters(self) -> list[str]:
        seen: list[str] = []
        for s in self.segments:
            if s.cluster not in seen:
                seen.append(s.cluster)
        return seen

    def cluster_segments(self, cluster: str) -> list[Segment]:
        return [s for s in self.segments if s.cluster == cluster]

    # -------------------------------------------------- 永続化
    def to_dict(self) -> dict[str, Any]:
        return {
            "schema": SCHEMA_VERSION,
            "audio_path": self.audio_path,
            "duration": self.duration,
            "chunk_seconds": self.chunk_seconds,
            "model": self.model,
            "verbatim": self.verbatim,
            "audio_fingerprint": self.audio_fingerprint,
            "time_offset": self.time_offset,
            "source_sha256": self.source_sha256,
            "engine": self.engine,
            "doc_revision": self.doc_revision,
            "edit_log": self.edit_log,
            "speakers": [sp.to_dict() for sp in self.speakers],
            "segments": [sg.to_dict() for sg in self.segments],
        }

    @classmethod
    def from_dict(cls, d: dict[str, Any]) -> "Project":
        return cls(
            audio_path=str(d.get("audio_path", "")),
            duration=float(d.get("duration", 0.0)),
            chunk_seconds=int(d.get("chunk_seconds", 600)),
            model=str(d.get("model", "")),
            verbatim=bool(d.get("verbatim", False)),
            audio_fingerprint=str(d.get("audio_fingerprint", "")),
            time_offset=float(d.get("time_offset", 0.0) or 0.0),
            # v4 以前のファイルには無い。既定値で読めば移行処理は不要
            source_sha256=str(d.get("source_sha256", "")),
            engine=dict(d.get("engine") or {}),
            doc_revision=int(d.get("doc_revision", 0) or 0),
            edit_log=list(d.get("edit_log") or []),
            speakers=[Speaker.from_dict(x) for x in d.get("speakers", [])],
            segments=[Segment.from_dict(x) for x in d.get("segments", [])],
        )

    def save(self, path: Optional[Path | str] = None) -> Path:
        target = Path(path or self.json_path or "")
        if not str(target):
            raise ValueError("保存先が指定されていません。")
        target.parent.mkdir(parents=True, exist_ok=True)
        tmp = target.with_suffix(target.suffix + ".tmp")
        tmp.write_text(
            json.dumps(self.to_dict(), ensure_ascii=False, indent=1),
            encoding="utf-8",
        )
        tmp.replace(target)
        self.json_path = str(target)
        return target

    @classmethod
    def load(cls, path: Path | str) -> "Project":
        p = Path(path)
        data = json.loads(p.read_text(encoding="utf-8"))
        proj = cls.from_dict(data)
        proj.json_path = str(p)
        return proj

    @staticmethod
    def default_json_path(output_dir: Path | str, audio_path: Path | str) -> Path:
        return Path(output_dir) / f"{Path(audio_path).stem}.speakers.json"


# ----------------------------------------------------------------------
# Word 出力
# ----------------------------------------------------------------------

def short_labels(speakers: Sequence[Speaker]) -> dict[str, str]:
    """話者 ID → 埋め込みで使う短い呼び名。

    本文の中に「(氏名：はい)」の形で入れるとき、名前が長いと相づちが肩書に
    埋もれる(実データで 6 字の相づちが 40 字の肩書に埋もれた・2026-08-18)。
    そこで**最初の空白まで**を使う。「山本学　文科省 高等教育局…」→「山本学」。

    **短くした結果が他の人とかぶるなら、その人は短くしない。**取り違えは
    誰が言ったかの記録そのものを壊すので、読みやすさより優先する。
    空白の無い名前も、切りようが無いのでそのまま。

    記録を書き換えるわけではない(【 】のラベルと出席者一覧は全名のまま)。
    根拠は設計書 §11.7。
    """
    cand: dict[str, str] = {}
    for sp in speakers:
        # \s は全角スペース(U+3000)も含む
        head = re.split(r"\s", sp.name.strip(), maxsplit=1)[0].strip()
        cand[sp.id] = head if head else sp.name
    # 短くした形が 2 人以上でぶつかるなら、その人たちは全名に戻す
    seen: dict[str, list[str]] = {}
    for sid, short in cand.items():
        seen.setdefault(short, []).append(sid)
    full = {sp.id: sp.name for sp in speakers}
    for short, ids in seen.items():
        if len(ids) > 1:
            for sid in ids:
                cand[sid] = full[sid]
    # 他人の全名とぶつかる場合も戻す(「山本」と「山本学」が並ぶ等)
    names = set(full.values())
    for sid, short in list(cand.items()):
        if short != full[sid] and short in names:
            cand[sid] = full[sid]
    return cand


def speaker_label(
    proj: Project, sid: Optional[str], with_role: bool = False
) -> str:
    """本文の【 】に出す呼び名(設計書 §11.8)。

    with_role=False … 名前だけ。「山本学」
    with_role=True  … 役職も付ける。「山本学(文科省 高等教育局…)」

    **出席者一覧は常に両方を載せる**(記録として全名を残す)ので、
    ここで名前だけにしても、誰なのかは文書の冒頭から辿れる。
    """
    sp = proj.speaker(sid)
    if not sp:
        return UNKNOWN_LABEL
    return sp.display if with_role else sp.name


def _insert_cuts(proj: Project) -> dict:
    """「どの区間の、本文の何文字目に、どの追加発話が割り込んだか」を集める。

    割り込み位置は add_utterance が edit_log に残している(区間そのものは
    割らない。割ると、直すときに元の本文を復元できない → 設計書 §5.0.5)。

    鍵は `(parent_orig, parent_start)`。**`parent_start` が無い古い記録は
    `None`** にしておき、突き合わせのときに `parent_orig` だけで当てる。

    値は `(何文字目, 足した発話の身元)`。**身元は開始時刻ではなく
    `added_signature` と同じ 3 つ組。**開始時刻だけだと、同じ時刻の 2 件が
    区別できず、Word で片方が消える(実データで発生・2026-08-20)。

    **消した記録は見ない。**存在しない区間は差し込みようがないので、
    突き合わせのときに区間の側で落ちる(`added_signature` を見よ)。
    """
    cuts: dict[tuple[float, Optional[float]],
               list[tuple[int, tuple[float, float, str]]]] = {}
    for rec in proj.edit_log:
        if (rec.get("op") == "add_utterance" and rec.get("cut") is not None
                and rec.get("parent_orig") is not None
                and rec.get("end") is not None):
            ps = rec.get("parent_start")
            key = (round(float(rec["parent_orig"]), 3),
                   None if ps is None else round(float(ps), 3))
            sig = (round(float(rec["orig_start"]), 3),
                   round(float(rec["end"]), 3),
                   str(rec.get("cluster") or ""))
            cuts.setdefault(key, []).append((int(rec["cut"]), sig))
    return cuts


def _cuts_for(cuts: dict, seg: "Segment") -> list:
    """その区間に差し込むもの。組で当て、古い形(orig だけ)も拾う。"""
    k = segment_key(seg)
    return list(cuts.get(k, [])) + list(cuts.get((k[0], None), []))


def has_inserted_utterances(proj: Project) -> bool:
    """出力の本文に差し込まれる追加発話があるか(凡例を出すかの判断に使う)"""
    added = {added_signature(s)
             for s in proj.segments if proj.is_added_utterance(s)}
    return any(sig in added
               for lst in _insert_cuts(proj).values() for _cut, sig in lst)


def inserted_marks(proj: Project) -> tuple[dict, set]:
    """「どの区間の何文字目に、どの追加発話が入るか」を割り当てる。

    返すのは (親の鍵 → [(何文字目, 足した区間)], 割り当て済みの id 集合)。

    **Word の出力と、割当画面の一覧で同じものを使う。**別々に実装すると、
    片方だけ直したときに「画面ではここ、出力では別のところ」という食い違いが
    起きる。この製品でいちばん避けたい種類の欠陥
    （書いてあることと実物が違う）。
    """
    cuts = _insert_cuts(proj)

    # **身元ごとに「並び」で持つ。**辞書に 1 つずつ入れていたため、同じ
    # 身元の 2 件目が 1 件目を上書きし、上書きされたほうが差し込み先を
    # 失って**出力から丸ごと落ちていた**(実データで発生・2026-08-20)。
    pool: dict[tuple[float, float, str], list[Segment]] = {}
    for s in proj.segments:
        if proj.is_added_utterance(s):
            pool.setdefault(added_signature(s), []).append(s)

    # 差し込み記録に区間を 1 つずつ割り当てる。**記録より区間が多いときは
    # 余りを単独で出す**(消し損ねて残った区間を黙って落とさないため)。
    marks_by_parent: dict[tuple, list[tuple[int, Segment]]] = {}
    placed: set[int] = set()
    for pkey, lst in cuts.items():
        for cut, sig in lst:
            queue = pool.get(sig)
            if not queue:
                continue
            seg_add = queue.pop(0)
            placed.add(id(seg_add))
            marks_by_parent.setdefault(pkey, []).append((cut, seg_add))
    return marks_by_parent, placed


def marks_for_segment(marks_by_parent: dict, seg: "Segment") -> list:
    """その区間に入る割り込みを、位置の順に返す。"""
    k = segment_key(seg)
    return sorted(marks_by_parent.get(k, [])
                  + marks_by_parent.get((k[0], None), []),
                  key=lambda c: c[0])


def _merge_runs(
    proj: Project,
    merge_consecutive: bool = True,
    drop_noise: bool = True,
    insert_style: str = INSERT_STYLE_LINE,
) -> list[tuple[float, Optional[str], str, bool]]:
    """(開始秒, 話者ID, 本文, 続きか) の並びを作る。

    merge_consecutive=True なら、同一話者の連続区間を 1 段落にまとめる。
    drop_noise=True なら「発言なし・雑音」と印を付けた区間は出力しない。
    insert_style は人が足した相づちの書き方(設計書 §11)。
      INSERT_STYLE_LINE   … 行を分け、割られた前半の末尾に ,, を付ける
      INSERT_STYLE_INLINE … 元の本文の中に (氏名：本文) の形で埋め込む

    「続きか」が True の要素は、**割られた発言の後半**である。時刻を書いては
    いけない。前半と同じ開始時刻しか持っておらず、そのまま出すと時刻が
    戻って見える。**後半の本当の開始時刻は測っていないので書かない**
    (測っていないものを書けば記録として嘘になる → 設計書 §11.3)。
    """
    # **人が足した発話は、割り込んだ位置で元の本文に差し込む。**
    # これをやらないと Word が「長い発言 → 相づち」の順になり、しかも
    # 同じ話者の相づちが 1 段落にまとまる(実機で判明・2026-08-18)。
    cuts = _insert_cuts(proj)
    marks_by_parent, placed = inserted_marks(proj)

    inline = insert_style == INSERT_STYLE_INLINE

    def marks_for(seg: Segment) -> list[tuple[int, Segment]]:
        return marks_for_segment(marks_by_parent, seg)

    shorts = short_labels(proj.speakers) if inline else {}

    def label_of(seg: Segment) -> str:
        """埋め込みで使う呼び名。長い肩書に相づちが埋もれるのを避ける。"""
        sp = proj.speaker(seg.speaker_id)
        if not sp:
            return UNKNOWN_LABEL
        return shorts.get(sp.id, sp.name)

    def pieces(seg: Segment) -> list[tuple[Segment, bool, bool]]:
        """区間を割り込みの位置で切った断片に分ける(出力のためだけ)。

        返すのは (断片, 続きか, 閉じたか)。「閉じた」は末尾に ,, を付けた
        断片で、後ろに何も足してはいけない(足すと ,, が文中に埋もれ、
        そこで割り込まれたことが読み取れなくなる)。
        inline のときは切らず、本文に埋め込んだ 1 つの断片にする。
        """
        marks = marks_for(seg)
        if not marks:
            return [(seg, False, False)]
        if inline:
            out, prev = [], 0
            for cut, add in marks:
                cut = max(0, min(len(seg.text), cut))
                if drop_noise and add.speaker_id == SPECIAL_NOISE:
                    continue
                body = add.text.strip()
                if not body:
                    continue
                out.append(seg.text[prev:cut])
                out.append(f"({label_of(add)}：{body})")
                prev = cut
            out.append(seg.text[prev:])
            return [(replace(seg, text="".join(out)), False, False)]
        out2: list[tuple[Segment, bool, bool]] = []
        prev, cont = 0, False
        for cut, add in marks:
            cut = max(0, min(len(seg.text), cut))
            head = seg.text[prev:cut]
            if head.strip():
                # 下の行に続くことを示す(BTSJ 2.3.2)。ここで閉じる
                out2.append(
                    (replace(seg, text=head.rstrip() + CONTINUE_MARK), cont, True))
                cont = True
            out2.append((add, False, False))
            prev = cut
        tail = seg.text[prev:]
        if tail.strip():
            out2.append((replace(seg, text=tail), cont, False))
        return out2

    ordered: list[tuple[Segment, bool, bool]] = []
    for seg in proj.segments:
        # 差し込み先が決まったものは、そこで出すので単独では出さない。
        # **決まらなかったものは出す。**先が無いのに黙って落とすと、人が
        # 打った発言が成果物から消える。
        if id(seg) in placed:
            continue
        ordered.extend(pieces(seg))

    runs: list[tuple[float, Optional[str], list[str], bool]] = []
    closed = False          # 直前の段落が ,, で閉じられたか
    for seg, cont, shut in ordered:
        text = seg.text.strip()
        if not text:
            continue
        if drop_noise and seg.speaker_id == SPECIAL_NOISE:
            continue
        if (
            merge_consecutive
            and runs
            and runs[-1][1] == seg.speaker_id
            and seg.speaker_id is not None
            and not cont          # 割られた後半は、前半と同じ段落に戻さない
            and not closed        # ,, で閉じた段落の後ろにも足さない
        ):
            runs[-1][2].append(text)
        else:
            runs.append((seg.start, seg.speaker_id, [text], cont))
        closed = shut
    # 日本語なので連結時に空白を挟まない
    return [(start, sid, "".join(parts), cont) for start, sid, parts, cont in runs]


# 編集履歴の op → 検証要約に出す日本語（編集履歴設計書 §3）。
# **知らない op は「その他」にまとめず、op 名のまま出す。**黙って畳むと
# 新しい操作を足したときに履歴から消えたように見える。
LOG_LABELS = {
    "assign": "話者",
    "assign_bulk": "話者(まとめて)",
    "undo_assign": "話者の取り消し",
    "clear_speakers": "名簿から外れた割当",
    "edit_text": "本文",
    "replace_text_bulk": "語句(まとめて)",
    "replace_speaker_bulk": "話者の置き換え",
    "carry_speakers": "前の転写から写した割当",
    "edit_time": "時刻",
    "apply_times_bulk": "時刻(まとめて)",
    "revert_time": "時刻を戻した",
    "undo_times": "時刻の取り消し",
    "add_utterance": "相づちを足した",
    "remove_added_utterance": "相づちを消した",
    "remove_segment": "区間を消した(転写の重複)",
    "restore_lost_segments": "消えた区間を戻した",
    "split": "区間の分割",
    "merge": "区間の結合",
}


def fmt_log_at(iso: str) -> str:
    """記録の時刻（UTC の ISO）を、その場の時刻の見やすい形にする。"""
    if not iso:
        return ""
    try:
        dt = datetime.fromisoformat(iso)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone().strftime("%Y-%m-%d %H:%M")
    except Exception:
        return iso


def build_log_summary(proj: Project) -> str:
    """検証要約に出す編集履歴の 1 行（明細は Day 75）。"""
    counts = proj.log_counts()
    if not counts:
        return "記録なし"
    total = sum(counts.values())
    parts = [f"{LOG_LABELS.get(op, op)} {n}"
             for op, n in sorted(counts.items(), key=lambda x: (-x[1], x[0]))]
    last = fmt_log_at(proj.log_last_at())
    out = f"全 {total} 件（" + " / ".join(parts) + "）"
    if last:
        out += f"  最終 {last}"
    return out


def build_verification(proj: Project, revision: int) -> list[tuple[str, str]]:
    """docx 末尾の「検証要約」の (項目, 値) を作る。

    書くのは確認の履歴であって、正しさの保証ではない。何をどの経路で処理し、
    人がどこまで確認したかを、第三者が検算できる形で残す(スキーマ v5)。
    """
    rows: list[tuple[str, str]] = [
        ("元音声", f"{Path(proj.audio_path).name} ({fmt_hms(proj.duration)})"),
    ]
    if proj.source_sha256:
        rows.append(("SHA-256", proj.source_sha256))
    if proj.engine:
        mode = ENGINE_LABELS.get(
            str(proj.engine.get("mode", "")), str(proj.engine.get("mode", "")))
        model = str(proj.engine.get("model", ""))
        at = str(proj.engine.get("at", ""))
        rows.append(("処理経路", " / ".join(x for x in (mode, model, at) if x)))
    rows.append(("版", f"revision {revision} (schema {SCHEMA_VERSION})"))

    # 数え方は「区間の数」で統一する。分母は必ず全区間数を書き、内訳は
    # 読点で区切る。区切りに「/」を使うと分数に見え、「聴いて確認 41 /
    # 適用のみ 40」が「41 分の 40」と読まれる(実出力で発生した)。
    # 検証要約は確認の履歴そのものなので、読み違えられる表示は信用を損なう。
    total = proj.total_count
    heard = proj.reviewed_count
    bulk = proj.unreviewed_count
    unassigned = total - proj.assigned_count
    rows.append((
        "話者の確認",
        f"全 {total} 区間 — 聴いて確定 {heard} 区間、"
        f"まとめて適用 {bulk} 区間、未確定 {unassigned} 区間",
    ))
    t_heard = sum(1 for s in proj.segments if s.time_edited and s.time_reviewed)
    t_bulk = sum(1 for s in proj.segments if s.time_edited and not s.time_reviewed)
    rows.append((
        "時刻の修正",
        f"全 {total} 区間中 {t_heard + t_bulk} 区間 — "
        f"聴いて確認 {t_heard} 区間、適用のみ {t_bulk} 区間",
    ))
    # **人が何回、何を触ったか。**「どこまで人が原音で確認したかを成果物に
    # 残せるか」が本製品の差別化そのもの(事業計画 v29)。件数の集計だけでは
    # 「いつまで手を入れたか」が分からない。明細は Day 75。
    rows.append(("編集の履歴", build_log_summary(proj)))
    rows.append((
        "凡例",
        "「聴いて確定」「聴いて確認」＝その区間の音声を人が聴いて決めたもの。"
        "「まとめて適用」「適用のみ」＝機械の結果をまとめて当てただけで、"
        "その区間を個別には聴いていないもの。数はいずれも区間の数です。",
    ))
    rows.append(("注意", "本書の記載は確認の履歴であり、内容の正しさや"
                        "法的効力を保証するものではありません。"))
    return rows


def build_note(proj: Project) -> str:
    """docx 冒頭に入れる但し書き。何がどこまで人手で確認されたかを明示する。"""
    parts = [
        "※ 話者ラベルはユーザーが音声を聴いて割り当てたものです",
        f"(全 {proj.total_count} 区間中、聴いて確定 {proj.reviewed_count} 区間",
    ]
    if proj.unreviewed_count:
        parts.append(f"、まとめて適用 {proj.unreviewed_count} 区間")
    unassigned = proj.total_count - proj.assigned_count
    if unassigned:
        parts.append(f"、未確定 {unassigned} 区間")
    parts.append(")。")
    return "".join(parts)


def write_docx(
    proj: Project,
    output_path: Path | str,
    title: Optional[str] = None,
    with_timestamps: bool = True,
    merge_consecutive: bool = True,
    include_note: bool = True,
    include_attendees: bool = True,
    drop_noise: bool = True,
    include_verification: bool = True,
    revision: Optional[int] = None,
    insert_style: str = INSERT_STYLE_LINE,
    with_role: bool = False,
) -> Path:
    """割当結果を Word ファイルに書き出す。

    include_verification: 末尾に検証要約(元音声・SHA-256・処理経路・版・
    確認状態)を付ける。revision はこの出力の版番号(省略時は記録済みの値)。
    insert_style: 人が足した相づちの書き方(設計書 §11)。
    with_role: 本文の【 】に企業・役職も入れる(設計書 §11.8)。
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    output_path = Path(output_path)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "游明朝"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "游明朝")

    doc.add_heading(title or Path(proj.audio_path).stem, level=1)

    if include_attendees and proj.speakers:
        p = doc.add_paragraph()
        head = p.add_run("出席者: ")
        head.bold = True
        p.add_run("、".join(sp.display for sp in proj.speakers))

    # 記号だけ出しても受け取った人が読めない。差し込みがあるときだけ添える
    if has_inserted_utterances(proj):
        p = doc.add_paragraph()
        run = p.add_run(INSERT_LEGEND[insert_style])
        run.italic = True
        run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    if include_note:
        p = doc.add_paragraph()
        run = p.add_run(build_note(proj))
        run.italic = True
        run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
        doc.add_paragraph()

    for start, sid, text, cont in _merge_runs(
            proj, merge_consecutive, drop_noise, insert_style):
        p = doc.add_paragraph()
        # 割られた発言の後半には時刻を書かない(測っていない → 設計書 §11.3)
        if with_timestamps and not cont:
            ts = p.add_run(f"[{fmt_hms(start)}] ")
            ts.bold = True
        label = speaker_label(proj, sid, with_role)
        name_run = p.add_run(f"【{label}】 ")
        name_run.bold = True
        if sid is None:
            # 一度も触れられていない区間。要確認なので赤で目立たせる
            name_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif sid in SPECIAL_SPEAKERS:
            # ユーザーが意図的に「不明」等と判断した区間はグレー
            name_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        p.add_run(text)

    if include_verification:
        doc.add_paragraph()
        head = doc.add_paragraph()
        run = head.add_run("―― 検証要約 ――")
        run.bold = True
        for label, value in build_verification(
                proj, revision if revision is not None else proj.doc_revision):
            p = doc.add_paragraph()
            r = p.add_run(f"{label}: ")
            r.bold = True
            r.font.size = Pt(9)
            v = p.add_run(value)
            v.font.size = Pt(9)
            v.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.save(str(output_path))
    return output_path


def write_text(
    proj: Project,
    output_path: Path | str,
    merge_consecutive: bool = True,
    drop_noise: bool = True,
    insert_style: str = INSERT_STYLE_LINE,
    with_role: bool = False,
) -> Path:
    """プレーンテキスト出力(自分のテンプレートに貼り込む場合や、差分取り用)"""
    output_path = Path(output_path)
    lines = []
    if has_inserted_utterances(proj):
        lines.append(INSERT_LEGEND[insert_style])
        lines.append("")
    for start, sid, text, cont in _merge_runs(
            proj, merge_consecutive, drop_noise, insert_style):
        # 後半には時刻を書かない。桁だけ空けて縦を揃える(設計書 §11.3)
        head = " " * (len(fmt_hms(start)) + 3) if cont else f"[{fmt_hms(start)}] "
        lines.append(
            f"{head}【{speaker_label(proj, sid, with_role)}】 {text}")
    output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return output_path
