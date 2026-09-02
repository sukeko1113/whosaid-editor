"""コア部分(GUI 以外)のユニットテスト。

    python -m pytest tests/ -q
    python tests/test_core.py     # pytest が無くても動く

GUI・Gemini API・ffmpeg には依存しない。
"""
from __future__ import annotations

import json
import os
import re
import sys
import tempfile
from collections import Counter
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

# Windows の既定コンソールは cp932 / cp1252 なので、日本語のラベルを print した
# 時点で UnicodeEncodeError になる(テストの中身とは無関係に落ちる)。
# 出力先のエンコーディングをここで固定しておく。
for _s in (sys.stdout, sys.stderr):
    if hasattr(_s, "reconfigure"):
        _s.reconfigure(encoding="utf-8", errors="replace")


from src.segments import (  # noqa: E402
    SCHEMA_VERSION,
    Project,
    SPECIAL_UNKNOWN,
    Segment,
    fmt_hms,
    fmt_hms_frac,
    parse_hms,
    parse_roster,
    roster_to_text,
)
from src.suggest import (  # noqa: E402
    SpeakerSuggester,
    next_unassigned,
    next_unreviewed,
)
from src.transcribe import (  # noqa: E402
    build_prompt,
    normalize_cluster_label,
    parse_segments,
    parse_utterances,
)
from src.align import AlignUnavailable  # noqa: E402
from src import diarize, evaluate  # noqa: E402
from src.local_asr import LocalTranscriber  # noqa: E402
from src.segments import (  # noqa: E402
    _merge_runs, has_inserted_utterances, write_text, UNKNOWN_LABEL,
    short_labels, Speaker, suggest_split, suggest_roster_rows,
    speaker_label, build_log_summary, fmt_log_at, LOG_LABELS,
    build_verification,
    INSERT_STYLE_LINE, INSERT_STYLE_INLINE, INSERT_LEGEND, CONTINUE_MARK)
from src.segments import TextHit, CONTEXT_CHARS  # noqa: E402,F401
from src.segments import segment_key  # noqa: E402
from src.segments import (  # noqa: E402
    PSEUDO_UNKNOWN,
    Utterance,
    assign_speaker_clusters,
    merge_same_speaker,
    utterances_to_segments,
)


# ======================================================================
# 音声のハッシュ(指紋と SHA-256)
# ======================================================================

def test_audio_hashes_computes_both_in_one_pass():
    """指紋と SHA-256 を同時に返し、SHA-256 は外部ツールの計算と一致する。"""
    import hashlib

    from src.audio import audio_fingerprint, audio_hashes

    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "a.wav"
        payload = b"RIFF" + bytes(range(256)) * 512
        p.write_bytes(payload)

        fp, sha = audio_hashes(p)
        # SHA-256 は素のファイル内容のハッシュ(certutil 等で検算できる値)
        assert sha == hashlib.sha256(payload).hexdigest()
        assert len(sha) == 64
        # 指紋は既存アルゴリズム(サイズをシードに含む BLAKE2b 64bit)と同一
        assert fp == audio_fingerprint(p)
        h = hashlib.blake2b(digest_size=8)
        h.update(str(len(payload)).encode("ascii"))
        h.update(payload)
        assert fp == h.hexdigest()


def test_audio_hashes_missing_file():
    from src.audio import audio_fingerprint, audio_hashes

    missing = Path("C:/no/such/dir/none.wav")
    assert audio_hashes(missing) == ("", "")
    assert audio_fingerprint(missing) == ""


def test_audio_hashes_change_with_content():
    """1 バイト変われば両方とも別の値になる(同一性判定の前提)。"""
    from src.audio import audio_hashes

    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "a.wav"
        p.write_bytes(b"abcdefg")
        before = audio_hashes(p)
        p.write_bytes(b"abcdefh")
        after = audio_hashes(p)
    assert before[0] != after[0] and before[1] != after[1]


# ======================================================================
# 名簿パース
# ======================================================================

def test_parse_roster_formats():
    text = """
    佐藤(理事長)
    田中（事務局長）: 議事進行を担当
    - 鈴木
    高橋: 会計
    """
    sp = parse_roster(text)
    assert [s.name for s in sp] == ["佐藤", "田中", "鈴木", "高橋"]
    assert sp[0].note == "理事長"
    assert sp[1].note == "事務局長 / 議事進行を担当"
    assert sp[2].note == ""
    assert sp[3].note == "会計"
    assert [s.order for s in sp] == [0, 1, 2, 3]
    # ID は一意
    assert len({s.id for s in sp}) == 4


def test_roster_round_trip():
    sp = parse_roster("佐藤(理事長)\n鈴木")
    assert roster_to_text(sp) == "佐藤(理事長)\n鈴木"


# ======================================================================
# クラスタラベルの正規化
# ======================================================================

def test_normalize_cluster_label():
    assert normalize_cluster_label("A") == "A"
    assert normalize_cluster_label("発言者B") == "B"
    assert normalize_cluster_label("話者 c") == "C"
    assert normalize_cluster_label("Speaker D") == "D"
    assert normalize_cluster_label("A(男性)") == "A"
    assert normalize_cluster_label("発言者不明") == "?"
    assert normalize_cluster_label("発言者複数・重複") == "*"
    assert normalize_cluster_label("?") == "?"
    assert normalize_cluster_label("*") == "*"
    assert normalize_cluster_label(None) == "?"
    assert normalize_cluster_label("") == "?"


# ======================================================================
# セグメント分解
# ======================================================================

SAMPLE = """[00:00] 【A】 本日はお忙しい中ありがとうございます。
[00:25] 【B】 前回の議事録は配布済みですか?
続きの行です。
[00:32] 【A】 はい、配布済みです。
[01:10] 【?】 (聴取不能)
"""


def test_parse_segments_basic():
    segs = parse_segments(SAMPLE, chunk_index=0, offset_seconds=0, chunk_seconds=600)
    assert len(segs) == 4
    assert segs[0]["start"] == 0.0
    assert segs[0]["cluster"] == "0:A"
    assert segs[1]["start"] == 25.0
    # 時刻の無い行は直前に連結される
    assert segs[1]["text"].endswith("続きの行です。")
    # 終了時刻は次の開始
    assert segs[1]["end"] == 32.0
    # 最終区間はチャンク末尾まで
    assert segs[3]["end"] == 600.0
    assert segs[3]["cluster"] == "0:?"
    assert [s["index"] for s in segs] == [0, 1, 2, 3]


def test_parse_segments_offset_and_chunk_namespacing():
    segs = parse_segments(SAMPLE, chunk_index=2, offset_seconds=1200, chunk_seconds=600,
                          start_index=10)
    assert segs[0]["start"] == 1200.0
    assert segs[1]["start"] == 1225.0
    assert segs[0]["cluster"] == "2:A"        # チャンクごとに別クラスタ扱い
    assert segs[0]["index"] == 10
    assert segs[-1]["end"] == 1800.0


def test_parse_segments_handles_messy_output():
    messy = """
    ここに前置きが入ってしまった
    [0:05]【発言者A】ラベルが全角括弧でない場合
    ［01:00］ 【B】 全角の括弧
    [00:30] 【C】 時刻が巻き戻っているケース
    [02:00]
    [02:10] 【B】 空行のあとの発言
    """
    segs = parse_segments(messy, chunk_index=1, offset_seconds=600, chunk_seconds=600)
    starts = [s["start"] for s in segs]
    # 単調増加が保たれる(巻き戻りは直前に合わせる)
    assert starts == sorted(starts)
    clusters = [s["cluster"] for s in segs]
    assert "1:A" in clusters and "1:B" in clusters and "1:C" in clusters
    # 本文が空の行は落ちる
    assert all(s["text"] for s in segs)


def test_parse_segments_empty():
    assert parse_segments("", 0, 0, 600) == []


# ======================================================================
# 同じ話者の連続をまとめる(v2.0.3)
#
# Gemini は指示しても息継ぎごとに行を分けてくることがある。
# 実機では 52 分の音声が 615 区間に割れ、割当作業が現実的でなくなった。
# ======================================================================

FRAGMENTED = """[00:00] 【A】 吉沢と申しますけども
[00:01] 【A】 えー、まずその6つのですね
[00:24] 【A】 工程表
[00:26] 【A】 えー、財源計画
[00:29] 【A】 年度別資金繰り
[00:47] 【A】 あの、第1の、あの、お願いでございます。
[01:05] 【B】 はい。
[01:06] 【A】 はい、以上です。
[01:06] 【B】 で、は、あの、代わりまして西村と申します。
[01:11] 【B】 えー、校長をしておりました。はい。
"""


def test_merge_collapses_same_speaker_run():
    merged = parse_segments(FRAGMENTED, 0, 0, 600)
    clusters = [s["cluster"] for s in merged]
    assert clusters == ["0:A", "0:B", "0:A", "0:B"], clusters
    # 1件目は 6 行分がまとまり、開始と終了が run 全体を覆う
    assert merged[0]["start"] == 0.0
    assert merged[0]["end"] == 65.0
    assert "吉沢と申しますけども" in merged[0]["text"]
    assert "お願いでございます。" in merged[0]["text"]


def test_merge_can_be_disabled():
    raw = parse_segments(FRAGMENTED, 0, 0, 600, merge_same_speaker=False)
    assert len(raw) == 10


def test_merge_inserts_reading_comma():
    """断片をそのまま繋ぐと「けどもえー」になる。間には読点を補う。"""
    merged = parse_segments(FRAGMENTED, 0, 0, 600)
    text = merged[0]["text"]
    assert "けども、えー" in text
    assert "けどもえー" not in text
    # 既に句読点で終わっていれば二重にしない
    assert "。、" not in text


def test_merge_skips_pseudo_clusters():
    """【?】【*】は中身がばらばらなので連結しない"""
    text = """[00:00] 【?】 うん
[00:02] 【?】 ええ
[00:04] 【*】 (重なり)
[00:06] 【*】 (重なり)
"""
    merged = parse_segments(text, 0, 0, 600)
    assert len(merged) == 4


def test_merge_respects_gap_and_max_length():
    from src.transcribe import merge_consecutive

    raw = [
        {"rel": 0.0, "rel_end": 10.0, "cluster": "A", "text": "あ"},
        {"rel": 60.0, "rel_end": 70.0, "cluster": "A", "text": "い"},   # 50秒の間
    ]
    assert len(merge_consecutive(raw, max_gap=20.0)) == 2
    assert len(merge_consecutive(raw, max_gap=90.0)) == 1

    long_run = [
        {"rel": float(i * 30), "rel_end": float(i * 30 + 30), "cluster": "A", "text": f"t{i}"}
        for i in range(10)      # 合計 300 秒
    ]
    out = merge_consecutive(long_run, max_seconds=120.0)
    assert len(out) > 1
    assert all(s["rel_end"] - s["rel"] <= 130.0 for s in out)


def test_merge_keeps_absolute_times_with_offset():
    merged = parse_segments(FRAGMENTED, chunk_index=2, offset_seconds=1200,
                            chunk_seconds=600)
    assert merged[0]["start"] == 1200.0
    assert merged[0]["cluster"] == "2:A"
    starts = [s["start"] for s in merged]
    assert starts == sorted(starts)


# ======================================================================
# 時刻が詰まった箇所の割り振り直し(v2.0.4 / v2.0.6)
#
# Gemini は発言が立て込むと 1 秒刻みで時刻を振る。そのまま使うと
# 長い発言が 1 秒で切られ、単に延ばすと隣同士の再生窓が重なって
# 同じ音声が聞こえる。本文の長さで按分して、重ならないようにする。
# ======================================================================

def test_estimate_speech_seconds():
    from src.transcribe import MIN_SEGMENT_SECONDS, estimate_speech_seconds

    assert estimate_speech_seconds("") == MIN_SEGMENT_SECONDS
    assert estimate_speech_seconds("うん。") == MIN_SEGMENT_SECONDS   # 短い相づちは下限
    long_line = "彼が要するに院外理事で、あの、理事で、あとはみんな学園の教授の方とか。"
    assert 6.0 < estimate_speech_seconds(long_line) < 12.0
    # 長さに比例する
    assert estimate_speech_seconds("あ" * 100) > estimate_speech_seconds("あ" * 50)


def test_overlapping_backchannel_does_not_truncate():
    """相づちが発言に重なると、終了時刻が次の開始で潰されていた。

    実機で確認された事例: 37文字の発言の再生窓が 1 秒になり、
    再生すると「彼」だけ聞こえて切れた。
    """
    text = """[07:02] 【C】 卒業生。
[07:03] 【A】 彼が要するに院外理事で、あの、理事で、あとはみんな学園の教授の方とか。
[07:04] 【C】 うん。
[07:30] 【A】 次の発言。
"""
    segs = parse_segments(text, chunk_index=5, offset_seconds=2700, chunk_seconds=540)
    long_seg = next(s for s in segs if "院外理事" in s["text"])
    # 次の開始(07:04)で切られず、本文に見合う長さが確保されている
    assert long_seg["end"] - long_seg["start"] > 6.0, long_seg
    # 後ろに余裕があるので開始は動かない
    assert long_seg["start"] == 2700 + 7 * 60 + 3
    # 余裕があっても必要以上には伸ばさない
    assert long_seg["end"] - long_seg["start"] < 12.0


def test_no_overlapping_playback_windows():
    """隣り合う区間が重なると、再生したときに同じ音声が聞こえてしまう。

    実機報告: 51:47 と 51:48、51:50 と 51:51、52:03 と 52:04 で
    同じ音声が流れた。
    """
    text = """[06:47] 【B】 それができればいいんだけど。
[06:48] 【A】 理事がもう。
[06:49] 【B】 はい。
[06:50] 【A】 ま、梅田さんとあともう1人の県議以外は全部内部理事。
[06:51] 【C】 うん。
[06:56] 【B】 県議で。
[06:57] 【A】 はい。
"""
    segs = parse_segments(text, chunk_index=5, offset_seconds=2700, chunk_seconds=540)
    for a, b in zip(segs, segs[1:]):
        assert b["start"] >= a["end"] - 0.01, (a["text"], b["text"], a["end"], b["start"])
    # 詰まった範囲でも、長い発言には短い相づちより多くの時間が割り当てられる
    # (最後の区間はチャンク末尾まで伸びるので比較から外す)
    long_seg = next(s for s in segs if "梅田" in s["text"])
    short = [s for s in segs[:-1] if s["text"].strip() in ("はい。", "うん。")]
    assert short, "比較対象の相づちが無い"
    assert all(long_seg["end"] - long_seg["start"] > s["end"] - s["start"] for s in short)
    assert long_seg["end"] - long_seg["start"] > 3.0


def test_roomy_timestamps_are_left_alone():
    """余裕がある箇所は Gemini の時刻をそのまま使う"""
    text = """[00:00] 【A】 短い。
[00:30] 【B】 これも短い。
[01:00] 【A】 やはり短い。
"""
    segs = parse_segments(text, 0, 0, 600)
    assert [s["start"] for s in segs] == [0.0, 30.0, 60.0]
    assert segs[0]["end"] == 30.0 and segs[1]["end"] == 60.0


def test_extension_is_clamped_to_chunk_end():
    """チャンク(音声)の末尾を越えて伸ばさない"""
    text = "[08:58] 【A】 " + "あ" * 300 + "\n"
    segs = parse_segments(text, chunk_index=0, offset_seconds=0, chunk_seconds=540)
    assert segs[-1]["end"] == 540.0


def test_time_is_shared_in_proportion_to_text():
    """詰まった範囲では、時間を本文の長さの比で分け合う。

    実際より時間が足りない以上、全員に必要な長さを与えることはできない。
    せめて長い発言に多く、短い相づちに少なく配ることで、
    再生したときに隣と同じ音にならないようにする。
    """
    text = """[00:00] 【A】 これはかなり長い発言でして、相づちが重なっても最後まで再生できる必要があります。
[00:01] 【B】 うん。
[00:02] 【A】 続けてもう一つ、これも長めの発言をしておきます。切られては困ります。
[00:03] 【C】 はい。
[00:20] 【A】 次の発言です。
"""
    from src.transcribe import estimate_speech_seconds

    segs = parse_segments(text, 0, 0, 600)
    dur = {s["text"][:8]: s["end"] - s["start"] for s in segs}

    # 0〜2 秒に 2 発言。時間が足りないので比で分ける
    assert dur["これはかなり長い"] > dur["うん。"] * 3

    # 後ろに余裕がある発言は、必要な長さを丸ごと確保できる
    need2 = estimate_speech_seconds(
        "続けてもう一つ、これも長めの発言をしておきます。切られては困ります。")
    assert dur["続けてもう一つ、"] >= need2 - 0.05

    # 隙間なく敷き詰められ、重ならない
    for a, b in zip(segs, segs[1:]):
        assert abs(b["start"] - a["end"]) < 0.02


def test_cluster_prompt_forbids_fragmenting():
    p = build_prompt(True, True, cluster_only=True)
    assert "同じ人が話し続けている間は、絶対に行を分けない" in p
    assert "息継ぎ" in p


def test_classify_api_error():
    """再試行しても直らないエラーを見分けられること"""
    from src.transcribe import classify_api_error

    depleted = (
        "429 RESOURCE_EXHAUSTED. {'error': {'code': 429, 'message': "
        "'Your prepayment credits are depleted. Please go to AI Studio at "
        "https://ai.studio/projects to manage your project and billing.', "
        "'status': 'RESOURCE_EXHAUSTED'}}"
    )
    msg = classify_api_error(Exception(depleted))
    assert msg is not None
    assert "残高が尽きています" in msg
    assert "aistudio.google.com" in msg

    quota = classify_api_error(Exception("You exceeded your current quota, please check your plan"))
    assert quota is not None and "利用上限" in quota

    for key_err in ("API key not valid. Please pass a valid API key.",
                    "403 PERMISSION_DENIED",
                    "UNAUTHENTICATED: request is missing credentials"):
        m = classify_api_error(Exception(key_err))
        assert m is not None, key_err
        assert "API キー" in m

    # 一時的なものは再試行させる(None)
    assert classify_api_error(Exception("Connection reset by peer")) is None
    assert classify_api_error(Exception("503 Service Unavailable")) is None
    assert classify_api_error(Exception("500 Internal error")) is None
    # 素の 429(分あたりのレート上限)は待てば通るので再試行対象
    assert classify_api_error(Exception("429 Too Many Requests")) is None


def test_rate_limit_detection():
    from src.transcribe import _is_rate_limited

    assert _is_rate_limited(Exception("429 Too Many Requests")) is True
    assert _is_rate_limited(Exception("RESOURCE_EXHAUSTED")) is True
    assert _is_rate_limited(Exception("Connection reset")) is False


def test_build_prompt_cluster_only_has_no_names():
    p = build_prompt(True, True, roster="佐藤(理事長)", verbatim=False, cluster_only=True)
    assert "佐藤" not in p              # 名簿は渡さない
    assert "【A】" in p
    assert "名前や役職は絶対に書かない" in p


def test_verbatim_prompt_does_not_absorb_backchannels():
    """逐語では相づちを前の行に含めさせない。

    **「一切要約・整文しない」と正面から矛盾する指示だった。**しかも相づちは
    消えるのではなく前の話者の行に混ざるので、「誰が言ったか」の記録としては
    消えるより悪い（別人の発言として残る）。
    CLAUDE.md の原則:「短い相づちの自動削除・自動重複除去はしない」。
    """
    p = build_prompt(True, True, verbatim=True)
    assert "前の発言と同じ行に含めてよい" not in p
    assert "独立した行" in p


def test_verbatim_prompt_limits_line_length():
    """同じ話者が続いても行を分けさせる。

    区切りが話者交代だけだと 1 区間が 112 秒に達し（実測）、**その区間には
    話者を割り当てられない**——間に何人も話しているため。
    """
    p = build_prompt(True, True, verbatim=True)
    assert "20 秒" in p


def test_verbatim_example_shows_short_lines():
    """例は指示より強く効く。相づちが独立した行になっている例を見せる。"""
    p = build_prompt(True, True, verbatim=True)
    assert "[00:08] 【発言者B】 はい。" in p


def test_cleanup_prompt_keeps_the_old_rule():
    """整文モード（既定）は従来どおり。手数を減らすための指示なので残す。"""
    p = build_prompt(True, True, verbatim=False)
    assert "前の発言と同じ行に含めてよい" in p
    assert "20 秒" not in p


def _long_turn(lines: int = 19) -> str:
    """同じ話者が 6 秒おきに話し続ける出力。実データと同じ形。"""
    out = ["[00:00] 【A】 よろしくお願いいたします。",
           "[00:05] 【B】 はい。",
           "[00:06] 【A】 はい、以上。"]
    for i in range(lines - 3):
        s = 7 + i * 6
        out.append(f"[{s//60:02d}:{s%60:02d}] 【C】 えー、あの、それでですね、続きの話をします。")
    return "\n".join(out)


def test_verbatim_does_not_merge_a_turn_into_one_block():
    """逐語では、同じ人が話し続けても 1 区間に潰さない。

    **Gemini は 5〜8 秒ごとに出しているのに、後処理が 112 秒の塊にしていた**
    （実測）。その長さの区間には話者を割り当てられない——間に何人も話すため。
    """
    u = parse_utterances(_long_turn(), chunk_seconds=103.0, verbatim=True)
    assert len(u) > 6, f"潰れている: {len(u)} 区間"
    # **最後の区間は除く。**「区間の終わりは次の区間の開始、最後だけチャンク
    # 末尾」という別の規則で伸びるので、連結の上限とは無関係に長くなる。
    assert max(x.rel_end - x.rel_start for x in u[:-1]) <= 21.0


def test_cleanup_still_merges_long_turns():
    """整文モード（既定）は従来どおり連結する。

    52 分の音声で 600 区間を超えると割当が現実的でなくなる、という判断は
    そのまま残す。
    """
    u = parse_utterances(_long_turn(), chunk_seconds=103.0, verbatim=False)
    assert max(x.rel_end - x.rel_start for x in u) > 60.0


def test_verbatim_still_joins_fragments():
    """断片の連結は逐語でも効く。1 文が数行に割れるのは読みづらい。"""
    text = ("[00:00] 【A】 求めるものは、\n"
            "[00:02] 【A】 工程表、\n"
            "[00:04] 【A】 えー、財源計画、\n"
            "[00:06] 【A】 年度別資金繰り計画です。")
    u = parse_utterances(text, chunk_seconds=30.0, verbatim=True)
    assert len(u) == 1, f"断片が連結されていない: {len(u)} 区間"


def test_verbatim_keeps_another_speaker_separate():
    """別の人の相づちは、逐語でも整文でも連結しない（元からの挙動）。"""
    text = "[00:00] 【A】 そうですね。\n[00:02] 【B】 はい。\n[00:03] 【A】 それで、"
    for vb in (True, False):
        u = parse_utterances(text, chunk_seconds=30.0, verbatim=vb)
        assert len(u) == 3, f"verbatim={vb} で {len(u)} 区間"


def test_verbatim_cache_key_carries_the_prompt_version():
    """逐語のキャッシュキーにプロンプトの版が入る。

    **無いと、指示を変えても古い転写を使い回す。**CLAUDE.md:
    「どれかを欠くと古い転写の使い回し事故になる」。
    """
    from src.pipeline import VERBATIM_PROMPT_VER, _cache_suffix
    now = _cache_suffix(True, True, True, "", 10, "abc")
    assert f"vb{VERBATIM_PROMPT_VER}" in now
    assert ".vb." not in now            # 版なしの古い形と衝突しない
    # 逐語でなければ版は付かない
    assert "vb" not in _cache_suffix(True, True, False, "", 10, "abc")


# ======================================================================
# 候補の学習・並べ替え
# ======================================================================

def _make_project() -> Project:
    proj = Project(audio_path="dummy.m4a", duration=600.0, chunk_seconds=600)
    proj.speakers = parse_roster("佐藤\n田中\n鈴木")
    texts = ["あ", "い", "う", "え", "お", "か", "が", "き", "く", "け"]
    clusters = ["0:A", "0:B", "0:A", "0:B", "0:A", "0:C", "0:A", "0:B", "0:A", "0:C"]
    proj.segments = [
        Segment(index=i, start=i * 10.0, end=i * 10.0 + 9, text=t, cluster=c, chunk=0)
        for i, (t, c) in enumerate(zip(texts, clusters))
    ]
    return proj


def test_cluster_vote_dominates():
    proj = _make_project()
    sato, tanaka, suzuki = proj.speakers
    # クラスタ 0:A を佐藤で 2 回確定
    proj.segments[0].speaker_id = sato.id
    proj.segments[2].speaker_id = sato.id
    s = SpeakerSuggester(proj)

    # 同じクラスタの未確定区間 → 佐藤が 1 位
    top = s.rank(4)[0]
    assert top.speaker.id == sato.id
    assert "同じ声のまとまり" in top.reason_text

    # 別クラスタでは佐藤が 1 位にはならない(頻度の加点だけ)
    ranked_b = s.rank(1)
    assert ranked_b[0].speaker.id != sato.id or ranked_b[0].score < top.score


def test_learning_improves_with_more_assignments():
    proj = _make_project()
    sato, tanaka, suzuki = proj.speakers
    s = SpeakerSuggester(proj)
    # 何も確定していない段階では名簿順
    assert [c.speaker.id for c in s.rank(1)] == [sato.id, tanaka.id, suzuki.id]

    # 0:B を田中で確定していくと、0:B の未確定は田中が 1 位になる
    proj.segments[1].speaker_id = tanaka.id
    proj.segments[3].speaker_id = tanaka.id
    s.refresh()
    assert s.rank(7)[0].speaker.id == tanaka.id
    assert s.dominant_speaker("0:B") == tanaka.id


def test_repeat_penalty_and_transition():
    proj = _make_project()
    sato, tanaka, suzuki = proj.speakers
    # 直前が佐藤。遷移統計が薄いので佐藤は減点されるはず
    proj.segments[0].speaker_id = sato.id
    s = SpeakerSuggester(proj)
    ranked = {c.speaker.id: c.score for c in s.rank(1)}
    assert ranked[sato.id] < ranked[tanaka.id]


def test_special_labels_excluded_from_stats():
    proj = _make_project()
    proj.segments[0].speaker_id = SPECIAL_UNKNOWN
    proj.segments[2].speaker_id = SPECIAL_UNKNOWN
    s = SpeakerSuggester(proj)
    # 不明だけのクラスタは、実在の話者を押し上げない
    ranked = s.rank(4)
    assert all(not c.reasons for c in ranked)
    assert s.dominant_speaker("0:A") is None


def test_rank_ignores_own_current_assignment():
    """付け替えたいときに、自分自身の割当が候補を固着させないこと"""
    proj = _make_project()
    sato, tanaka, _ = proj.speakers
    proj.segments[0].speaker_id = sato.id   # このクラスタで唯一の確定
    s = SpeakerSuggester(proj)
    ranked = s.rank(0)                      # 自分自身を評価
    assert not ranked[0].reasons            # クラスタ投票の根拠は消えている


def test_pseudo_clusters_are_not_learned():
    """【?】【*】は「同じ声のまとまり」ではないので、投票にも一括適用にも使わない"""
    proj = _make_project()
    sato = proj.speakers[0]
    for i, seg in enumerate(proj.segments):
        seg.cluster = "0:?"
    proj.segments[0].speaker_id = sato.id
    s = SpeakerSuggester(proj)

    top = s.rank(4)[0]
    assert not top.reasons, "判別不能の区間から他の区間を推測してはいけない"
    assert s.dominant_speaker("0:?") is None
    assert proj.segments[4].is_pseudo_cluster is True
    assert "一括適用は不可" in s.cluster_summary("0:?")


def test_boundary_continuity_bonus():
    """チャンク境界をまたぐ隣接区間は同一話者の可能性が高い"""
    proj = Project(audio_path="x.m4a", duration=1200.0, chunk_seconds=600)
    proj.speakers = parse_roster("佐藤\n田中\n鈴木")
    sato, tanaka, _ = proj.speakers
    proj.segments = [
        Segment(index=0, start=560.0, end=595.0, text="…そこで", cluster="0:B", chunk=0),
        Segment(index=1, start=600.0, end=630.0, text="申し上げた", cluster="1:A", chunk=1),
        Segment(index=2, start=630.0, end=660.0, text="別の発言", cluster="1:B", chunk=1),
    ]
    proj.segments[0].speaker_id = tanaka.id
    s = SpeakerSuggester(proj)

    ranked = s.rank(1)
    assert ranked[0].speaker.id == tanaka.id
    assert "チャンク境界" in ranked[0].reason_text
    # 境界でない区間には効かない
    assert "チャンク境界" not in " ".join(c.reason_text for c in s.rank(2))


def test_letter_prior_across_chunks():
    """前のチャンクの A が誰だったかは、次のチャンクの A の弱い手がかりになる"""
    proj = Project(audio_path="x.m4a", duration=1800.0, chunk_seconds=600)
    proj.speakers = parse_roster("佐藤\n田中\n鈴木")
    sato = proj.speakers[0]
    segs = []
    for i in range(6):
        chunk = i // 2
        segs.append(Segment(index=i, start=i * 100.0, end=i * 100.0 + 50,
                            text=f"t{i}", cluster=f"{chunk}:{'AB'[i % 2]}", chunk=chunk))
    proj.segments = segs
    # チャンク 0 と 1 の A を佐藤で確定
    proj.segments[0].speaker_id = sato.id
    proj.segments[2].speaker_id = sato.id
    s = SpeakerSuggester(proj)

    # チャンク 2 の A(未確定・クラスタ票なし)でも佐藤が 1 位
    ranked = s.rank(4)
    assert ranked[0].speaker.id == sato.id
    assert "別チャンクの 【A】" in ranked[0].reason_text


def test_next_unreviewed():
    proj = _make_project()
    sid = proj.speakers[0].id
    proj.segments[0].speaker_id = sid
    proj.segments[0].reviewed = True
    proj.segments[1].speaker_id = sid     # 一括適用ぶん(未確認)
    assert proj.reviewed_count == 1
    assert proj.unreviewed_count == 1
    assert next_unreviewed(proj, 0) == 1
    assert next_unassigned(proj, 0) == 2


def test_next_unassigned():
    proj = _make_project()
    proj.segments[0].speaker_id = "sp01"
    proj.segments[1].speaker_id = "sp01"
    assert next_unassigned(proj, 0) == 2
    assert next_unassigned(proj, 1) == 2
    assert next_unassigned(proj, 9) is None
    assert next_unassigned(proj, 3, forward=False) == 2


# ======================================================================
# 保存・読み込み・出力
# ======================================================================

def test_save_load_round_trip():
    proj = _make_project()
    proj.segments[0].speaker_id = proj.speakers[0].id
    proj.segments[0].reviewed = True
    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "x.speakers.json"
        proj.save(p)
        loaded = Project.load(p)
    assert loaded.total_count == proj.total_count
    assert loaded.assigned_count == 1
    assert loaded.segments[0].reviewed is True
    assert [s.name for s in loaded.speakers] == ["佐藤", "田中", "鈴木"]
    assert loaded.clusters() == ["0:A", "0:B", "0:C"]


def test_remove_speaker_clears_assignments():
    proj = _make_project()
    sid = proj.speakers[0].id
    proj.segments[0].speaker_id = sid
    proj.remove_speaker(sid)
    assert proj.segments[0].speaker_id is None
    assert proj.assigned_count == 0
    assert [s.order for s in proj.speakers] == [0, 1]


def test_write_docx_merges_consecutive():
    from docx import Document
    from src.segments import SPECIAL_NOISE, write_docx

    proj = _make_project()
    sid = proj.speakers[0].id
    for i in (0, 1, 2):
        proj.segments[i].speaker_id = sid
    proj.segments[3].speaker_id = SPECIAL_NOISE    # 出力から省かれるはず
    with tempfile.TemporaryDirectory() as d:
        out = Path(d) / "out.docx"
        write_docx(proj, out, title="第1回理事会")
        doc = Document(str(out))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
    body = [p for p in paras if p.startswith("[")]
    # 連続する 3 区間が 1 段落にまとまる(日本語なので空白を挟まない)
    assert body[0].startswith("[00:00:00] 【佐藤】")
    assert "あいう" in body[0]
    # 表題・出席者一覧が入る
    assert paras[0] == "第1回理事会"
    assert any(p.startswith("出席者: 佐藤、田中、鈴木") for p in paras)
    # 雑音と印を付けた区間は出ない
    assert not any("え" == p[-1] for p in body if "【発言なし" in p)
    assert all("発言なし・雑音" not in p for p in paras)
    # 未確定は【発言者不明】
    assert "【発言者不明】" in body[1]


def test_write_docx_options():
    from docx import Document
    from src.segments import write_docx

    proj = _make_project()
    sid = proj.speakers[0].id
    for i in (0, 1):
        proj.segments[i].speaker_id = sid
    with tempfile.TemporaryDirectory() as d:
        out = Path(d) / "out.docx"
        write_docx(proj, out, with_timestamps=False, merge_consecutive=False,
                   include_attendees=False, include_note=False)
        paras = [p.text for p in Document(str(out)).paragraphs if p.text.strip()]
    assert not any(p.startswith("[") for p in paras)
    assert not any(p.startswith("出席者:") for p in paras)
    assert paras[1] == "【佐藤】 あ"      # まとめない
    assert paras[2] == "【佐藤】 い"


def test_write_docx_uses_edited_times():
    """時刻を直した区間は、直した時刻で出力される。"""
    from docx import Document
    from src.segments import write_docx

    proj = _make_project()
    proj.time_offset = 4.0          # ずれ補正は出力に効かない(再生だけのもの)
    seg = proj.segments[3]          # 既定では 30.0〜39.0
    seg.speaker_id = proj.speakers[0].id
    seg.start, seg.end = 37.0, 45.0
    seg.time_edited = True
    with tempfile.TemporaryDirectory() as d:
        out = Path(d) / "out.docx"
        write_docx(proj, out, merge_consecutive=False)
        paras = [p.text for p in Document(str(out)).paragraphs if p.text.strip()]
    assert any(p.startswith("[00:00:37]") for p in paras)
    assert not any(p.startswith("[00:00:30]") for p in paras)


def test_build_note_distinguishes_reviewed():
    from src.segments import build_note

    proj = _make_project()
    sid = proj.speakers[0].id
    proj.segments[0].speaker_id = sid
    proj.segments[0].reviewed = True
    proj.segments[1].speaker_id = sid          # 一括適用ぶん
    note = build_note(proj)
    assert "聴いて確定 1 区間" in note
    assert "まとめて適用 1 区間" in note
    assert "未確定 8 区間" in note


def test_fmt_hms():
    assert fmt_hms(0) == "00:00:00"
    assert fmt_hms(59.6) == "00:01:00"
    assert fmt_hms(3725) == "01:02:05"


# ======================================================================
# 時刻の表示・入力(0.1 秒精度)
# ======================================================================

def test_fmt_hms_frac():
    assert fmt_hms_frac(0) == "00:00:00.0"
    assert fmt_hms_frac(2631.5) == "00:43:51.5"
    assert fmt_hms_frac(3725.04) == "01:02:05.0"
    assert fmt_hms_frac(59.96) == "00:01:00.0"      # 秒への繰り上がり
    assert fmt_hms_frac(-3.0) == "00:00:00.0"       # 負は 0 に丸める


def test_parse_hms_formats():
    assert parse_hms("01:02:05.4") == 3725.4
    assert parse_hms("43:51.5") == 2631.5
    assert parse_hms("7.5") == 7.5
    assert parse_hms("7") == 7.0
    assert parse_hms("00:00:00.0") == 0.0
    # 最上位の桁だけは 60 以上を許す(90 秒、75 分)
    assert parse_hms("90") == 90.0
    assert parse_hms("75:00") == 4500.0
    # 前後の空白と、全角のまま打たれた数字・区切りも受ける
    assert parse_hms("  01:02:05.4  ") == 3725.4
    assert parse_hms("００:４３:５１．５") == 2631.5


def test_parse_hms_rejects_bad_values():
    bad = [
        "", "   ",
        "abc",
        "01:02:03:04",       # 桁が多い
        "01::05",            # 空フィールド
        ":30",
        "-5",                # 負
        "1.5:30",            # 小数は最下位だけ
        "01:75:00",          # 分が 60 以上
        "43:75.0",           # 秒が 60 以上
        "12:34:5x",
    ]
    for s in bad:
        try:
            parse_hms(s)
        except ValueError:
            continue
        raise AssertionError(f"拒否されるべき入力が通った: {s!r}")


def test_hms_round_trip():
    for sec in (0.0, 0.1, 7.5, 59.9, 2631.5, 3725.4, 36000.0):
        assert parse_hms(fmt_hms_frac(sec)) == sec


# ======================================================================
# 時刻編集のスキーマ(schema 3)
# ======================================================================

def test_orig_times_default_to_start_end():
    seg = Segment(index=0, start=10.0, end=20.0, text="あ", cluster="0:A")
    assert seg.orig_start == 10.0
    assert seg.orig_end == 20.0
    assert seg.time_edited is False
    # 明示指定したときはそちらが勝つ(分割で親の値を共有させるため)
    child = Segment(index=1, start=15.0, end=20.0, text="い", cluster="0:A",
                    orig_start=10.0, orig_end=20.0)
    assert (child.orig_start, child.orig_end) == (10.0, 20.0)


def test_load_schema2_file_fills_orig_times():
    """orig_* を持たない旧ファイルは既定値だけで読め、保存は今の版になる。"""
    old = {
        "schema": 2,
        "audio_path": "a.m4a",
        "duration": 100.0,
        "speakers": [{"id": "sp01", "name": "佐藤", "note": "", "order": 0}],
        "segments": [
            {"index": 0, "start": 10.0, "end": 20.0, "text": "あ", "cluster": "0:A",
             "chunk": 0, "speaker_id": "sp01", "reviewed": True, "note": "",
             "text_edited": False},
        ],
    }
    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "old.speakers.json"
        p.write_text(json.dumps(old, ensure_ascii=False), encoding="utf-8")

        proj = Project.load(p)
        seg = proj.segments[0]
        assert (seg.orig_start, seg.orig_end) == (10.0, 20.0)
        assert seg.time_edited is False

        # 時刻を直して保存 → 今の schema で書かれ、orig は元の時刻のまま残る
        seg.start, seg.end = 16.0, 26.0
        seg.time_edited = True
        proj.save(p)
        data = json.loads(p.read_text(encoding="utf-8"))
        assert data["schema"] == SCHEMA_VERSION

        again = Project.load(p)
        s2 = again.segments[0]
        assert (s2.start, s2.end) == (16.0, 26.0)
        assert (s2.orig_start, s2.orig_end) == (10.0, 20.0)
        assert s2.time_edited is True
        assert s2.reviewed is True and s2.speaker_id == "sp01"


# ======================================================================
# 区間の分割・結合
# ======================================================================

def _splittable() -> Project:
    proj = Project(audio_path="a.m4a", duration=300.0)
    proj.speakers = parse_roster("佐藤\n田中")
    proj.segments = [
        Segment(index=0, start=90.0, end=100.0, text="前の発言", cluster="0:A"),
        Segment(index=1, start=100.0, end=110.0, text="そうですねいいと思います",
                cluster="0:B", chunk=0, speaker_id="sp01", reviewed=True, note="めも"),
        Segment(index=2, start=110.0, end=120.0, text="次の発言", cluster="0:C"),
    ]
    return proj


# ======================================================================
# 相づちを足す(設計書 claude_相づちを足す_設計書.md)
# ======================================================================

def test_add_utterance_inserts_in_time_order():
    """時間順の正しい位置に入り、番号が振り直される。"""
    proj = _splittable()
    seg = proj.add_utterance(105.0, 105.8, "はいはい", cluster="g:B")
    assert len(proj.segments) == 4
    assert [s.index for s in proj.segments] == [0, 1, 2, 3]
    assert proj.segments[2] is seg                    # 100〜110 の次
    assert (seg.start, seg.end, seg.text) == (105.0, 105.8, "はいはい")
    assert seg.cluster == "g:B"


def test_add_utterance_does_not_assign_a_speaker():
    """話者は付けずに返す。**機械が ✓ を立てる経路は作らない。**

    話者を付けるのは呼び出し側の通常の割当操作で、そのときの ✓/△ は
    既存の意味論に従う。
    """
    proj = _splittable()
    seg = proj.add_utterance(105.0, 105.8, "うん")
    assert seg.speaker_id is None
    assert seg.reviewed is False


def test_add_utterance_marks_text_edited_not_time_edited():
    """本文は人が打った(text_edited)。時刻は turn 由来の機械値。

    time_edited を立てると、再実行の引き継ぎで「旧側が正しい」枝に入り、
    **他の区間を置き換えて消す**(設計書 §4)。
    """
    proj = _splittable()
    seg = proj.add_utterance(105.0, 105.8, "ええ")
    assert seg.text_edited is True
    assert seg.time_edited is False


def test_add_utterance_allows_overlap():
    """既存区間と時間的に重なってよい。**相づちは重なるのが本性。**"""
    proj = _splittable()
    seg = proj.add_utterance(102.0, 103.0, "はい")     # 100〜110 の内側
    assert len(proj.segments) == 4
    others = [s for s in proj.segments if s is not seg]
    assert any(s.start < seg.start and s.end > seg.end for s in others)


def test_add_utterance_rejects_empty_text():
    proj = _splittable()
    try:
        proj.add_utterance(105.0, 105.8, "   ")
    except ValueError:
        pass
    else:
        raise AssertionError("空の本文を受け付けてしまった")


def test_add_utterance_records_the_key_in_edit_log():
    """識別は edit_log。**スキーマ(区間のフラグ)を増やさない。**"""
    proj = _splittable()
    seg = proj.add_utterance(105.0, 105.8, "はい")
    rec = [r for r in proj.edit_log if r.get("op") == "add_utterance"]
    assert len(rec) == 1
    assert rec[0]["orig_start"] == seg.orig_start
    assert rec[0]["actor"] == "user"
    assert proj.added_utterance_keys() == {round(seg.orig_start, 3)}
    assert proj.is_added_utterance(seg) is True
    assert proj.is_added_utterance(proj.segments[0]) is False


def _base_for_insert() -> Project:
    proj = Project(audio_path="a.m4a", duration=300.0)
    proj.speakers = parse_roster("西村" + chr(10) + "山本")
    proj.segments = [
        Segment(index=0, start=90.0, end=100.0, text="前の発言", cluster="g:A"),
        Segment(index=1, start=100.0, end=110.0,
                text="ながらけれどもそれで", cluster="g:B",
                speaker_id="sp01", reviewed=True),
        Segment(index=2, start=110.0, end=120.0, text="次の発言", cluster="g:C"),
    ]
    return proj


def _add_at(proj, base, cut, at, text, sid=None):
    seg = proj.add_utterance(at, at + 0.6, text, "g:D",
                             cut=cut, parent_orig=base.orig_start)
    seg.speaker_id = sid
    seg.reviewed = bool(sid)
    return seg


def test_added_utterance_records_where_it_cut_in():
    """**どの区間の本文の、どこに割り込んだかを残す。**

    区間は割らない(割ると、直すときに元の本文を復元できない)。
    代わりにこれを覚えておき、Word に出すときだけ差し込む。
    """
    proj = _base_for_insert()
    base = proj.segments[1]
    _add_at(proj, base, 3, 103.0, "はい", "sp02")
    rec = [r for r in proj.edit_log if r.get("op") == "add_utterance"][0]
    assert rec["cut"] == 3
    assert rec["parent_orig"] == base.orig_start
    # 元の区間はそのまま（割れていない）
    assert proj.segments[1].text == "ながらけれどもそれで"


def test_word_output_puts_the_added_one_inside_the_text():
    """**Word 出力が読んだとおりの順になる。**これが cut を残す理由そのもの。

    割り込み位置で本文を切り、その間に足した発話を挟む。やらないと
    「長い発言 → 相づち」の順になり、しかも同じ話者の相づちが 1 段落に
    まとまる（実機で「あ、はいはいはいええはいはい」になった）。
    """
    proj = _base_for_insert()
    base = proj.segments[1]
    _add_at(proj, base, 3, 103.0, "はい", "sp02")
    _add_at(proj, base, 7, 107.0, "ええ", "sp02")
    bodies = [text for _s, _sid, text, _c in _merge_runs(proj, True)]
    assert "ながら" + CONTINUE_MARK in bodies, bodies
    assert "はい" in bodies and "ええ" in bodies, bodies
    assert "はいええ" not in "".join(bodies), "相づちがひとまとまりになっている"
    # 元の話者の断片が、相づちを挟んで前後に出る
    assert bodies.index("はい") == bodies.index("ながら" + CONTINUE_MARK) + 1
    assert bodies.index("ええ") > bodies.index("はい")


def test_word_output_does_not_duplicate_the_added_one():
    """差し込んだものを、単独でももう一度出さない。"""
    proj = _base_for_insert()
    _add_at(proj, proj.segments[1], 3, 103.0, "はい", "sp02")
    bodies = [text for _s, _sid, text, _c in _merge_runs(proj, True)]
    assert sum(1 for b in bodies if b == "はい") == 1, bodies


def test_word_output_keeps_added_ones_without_a_cut():
    """位置を持たない足し方（データ層の直呼び）でも消えない。"""
    proj = _base_for_insert()
    proj.add_utterance(103.0, 103.6, "うん", "g:D")
    bodies = [text for _s, _sid, text, _c in _merge_runs(proj, True)]
    assert "うん" in "".join(bodies)


def test_word_output_forgets_removed_ones():
    """消した発話は差し込まれない。"""
    proj = _base_for_insert()
    added = _add_at(proj, proj.segments[1], 3, 103.0, "はい", "sp02")
    proj.remove_added_utterance(added.index)
    bodies = [text for _s, _sid, text, _c in _merge_runs(proj, True)]
    assert "はい" not in "".join(bodies)
    assert "ながらけれどもそれで" in bodies, bodies


# --- 出力の表記法(設計書 §11)------------------------------------------
# 標準を調べたら BTSJ(基本的な文字化の原則)に規定があった。データは 1 つの
# まま、出し方だけを 2 通りから選ぶ。壊れたら落とす。

def test_line_style_marks_that_the_utterance_continues():
    """**割られた前半の末尾に ,, を付ける(BTSJ 2.3.2)。**

    これが無いと、読む人は「1 つの発言が割り込まれた」のか「同じ人が 2 回
    別々に発言した」のか区別できない。記録としての欠陥になる。
    """
    proj = _base_for_insert()
    _add_at(proj, proj.segments[1], 3, 103.0, "はい", "sp02")
    runs = _merge_runs(proj, True, True, INSERT_STYLE_LINE)
    bodies = [r[2] for r in runs]
    assert "ながら" + CONTINUE_MARK in bodies, bodies
    assert "けれどもそれで" in bodies, bodies        # 後半には付けない
    assert not any(b.endswith(CONTINUE_MARK) for b in bodies if b == "はい")


def test_line_style_does_not_put_a_time_on_the_second_half():
    """**割られた後半に時刻を書かない(設計書 §11.3)。**

    後半は前半と同じ開始時刻しか持っておらず、そのまま出すと時刻が戻って
    見える。**測っていないものを書けば記録として嘘になる。**
    """
    proj = _base_for_insert()
    _add_at(proj, proj.segments[1], 3, 103.0, "はい", "sp02")
    runs = _merge_runs(proj, True, True, INSERT_STYLE_LINE)
    by_text = {r[2]: r for r in runs}
    assert by_text["ながら" + CONTINUE_MARK][3] is False   # 前半は時刻を出す
    assert by_text["けれどもそれで"][3] is True            # 後半は出さない
    assert by_text["はい"][3] is False                     # 相づちは出す


def test_line_style_keeps_the_two_halves_in_separate_paragraphs():
    """後半を、前半と同じ段落にまとめ直さない(まとめると ,, が文中に埋もれる)"""
    proj = _base_for_insert()
    # 相づちを、割られた前後と同じ話者にする（まとめの条件を踏ませる）
    _add_at(proj, proj.segments[1], 3, 103.0, "はい", proj.segments[1].speaker_id)
    bodies = [r[2] for r in _merge_runs(proj, True, True, INSERT_STYLE_LINE)]
    assert "ながら" + CONTINUE_MARK in bodies, bodies
    assert "けれどもそれで" in bodies, bodies


def test_inline_style_puts_the_backchannel_inside_the_text():
    """**行に埋め込む形(BTSJ 3.2.3 例49)。**

    多人数の会議なので、BTSJ と違って氏名を入れる(誰が言ったかの記録が
    製品価値そのものなので、ここは譲れない)。
    """
    proj = _base_for_insert()
    _add_at(proj, proj.segments[1], 3, 103.0, "はい", "sp02")
    _add_at(proj, proj.segments[1], 7, 107.0, "ええ", "sp02")
    runs = _merge_runs(proj, True, True, INSERT_STYLE_INLINE)
    bodies = [r[2] for r in runs]
    joined = "".join(bodies)
    assert "ながら(山本：はい)けれども(山本：ええ)それで" in joined, bodies
    assert not any(b == "はい" for b in bodies), "単独でも出してしまっている"
    assert all(r[3] is False for r in runs), "埋め込みでは後半という概念が無い"


def test_inline_style_uses_the_unknown_label_when_not_assigned():
    """話者が決まっていなくても落ちない(名前の代わりに既定の呼び名)"""
    proj = _base_for_insert()
    _add_at(proj, proj.segments[1], 3, 103.0, "はい", None)
    joined = "".join(r[2] for r in _merge_runs(proj, True, True, INSERT_STYLE_INLINE))
    assert f"({UNKNOWN_LABEL}：はい)" in joined, joined


def test_both_styles_agree_on_what_was_said():
    """**どちらで出しても、言葉そのものは同じ。**出し方だけが違う。"""
    proj = _base_for_insert()
    _add_at(proj, proj.segments[1], 3, 103.0, "はい", "sp02")

    def words(style):
        joined = "".join(r[2] for r in _merge_runs(proj, True, True, style))
        for ch in "()（）：,山本":
            joined = joined.replace(ch, "")
        return joined
    assert words(INSERT_STYLE_LINE) == words(INSERT_STYLE_INLINE)


def test_short_label_cuts_at_the_first_space():
    """**長い肩書に相づちが埋もれるのを避ける(設計書 §11.7)。**

    実データで、6 字の相づちが 40 字の肩書に埋もれた。
    """
    sps = parse_roster(chr(10).join(
        ["山本学　文科省 高等教育局私学部参事官付 企画官", "西村香介"]))
    got = short_labels(sps)
    assert got[sps[0].id] == "山本学", got
    assert got[sps[1].id] == "西村香介", "空白が無ければそのまま"


def test_short_label_keeps_the_full_name_when_it_would_collide():
    """**短くしてかぶるなら短くしない。**取り違えは読みやすさより重い。"""
    sps = [Speaker(id="a", name="山本 学"), Speaker(id="b", name="山本 花子")]
    got = short_labels(sps)
    assert got["a"] == "山本 学" and got["b"] == "山本 花子", got


def test_short_label_avoids_colliding_with_someone_elses_full_name():
    """「山本」と「山本学　…」が並ぶとき、短くすると別人と読める。"""
    sps = [Speaker(id="a", name="山本学　文科省"), Speaker(id="b", name="山本学")]
    got = short_labels(sps)
    assert got["a"] == "山本学　文科省", got


def test_short_label_only_affects_the_inline_form():
    """**記録は書き換えない。**【 】と出席者一覧は全名のまま。"""
    proj = _base_for_insert()
    proj.speakers[0].name = "佐藤　△△株式会社　代表取締役"
    proj.speakers[1].name = "山本　□□省　企画官"
    _add_at(proj, proj.segments[1], 3, 103.0, "はい", "sp02")

    inline = _merge_runs(proj, True, True, INSERT_STYLE_INLINE)
    assert "(山本：はい)" in "".join(r[2] for r in inline)

    # 行を分ける形では短くしない（そこは話者そのものの表示なので）
    line = _merge_runs(proj, True, True, INSERT_STYLE_LINE)
    sid = proj.speakers[1].id
    assert any(r[1] == sid and r[2] == "はい" for r in line), line
    assert proj.speakers[1].name == "山本　□□省　企画官", "名簿を書き換えている"


def test_inline_legend_says_the_name_is_abbreviated():
    """略記だと書いておかないと、受け取った人が別人と思う。"""
    assert "略記" in INSERT_LEGEND[INSERT_STYLE_INLINE]


# --- 名前と企業・役職に分ける(設計書 §11.8)--------------------------
# ユーザーの提案。表示名という別の項目を持つと同じ情報が 2 か所になり、
# 片方だけ直したとき食い違う。名前と役職に分けておけば表示は組み立てるだけ。

def test_suggest_split_cuts_at_a_role_word():
    """「三ツ林衆議院議員」→ 名前「三ツ林」／役職「衆議院議員」"""
    assert suggest_split("三ツ林衆議院議員") == ("三ツ林", "衆議院議員")
    assert suggest_split("吉沢忠一加茂暁星高校同窓会長")[0] == "吉沢忠一加茂暁星"


def test_suggest_split_cuts_at_a_space():
    """空白があればそこで切る（あなたの名簿の多くはこの形）"""
    n, note = suggest_split("山本学　文科省 高等教育局私学部参事官付 企画官")
    assert n == "山本学" and note.startswith("文科省")


def test_suggest_split_uses_other_rows_to_find_the_organisation():
    """**他の人の行にも出てくる並びは、個人名ではなく所属。**

    「加茂暁星」が梅田さんの行にも出るので、吉沢さんの名前ではない。
    """
    others = ["梅田茂　加茂暁星学園理事"]
    assert suggest_split("吉沢忠一加茂暁星高校同窓会長", others)[0] == "吉沢忠一"


def test_suggest_split_never_eats_a_real_surname():
    """**1 文字の「市」「村」「部」を手がかりにしない。**

    使うと「田村」さんが「田」になる。役職語は 2 文字以上のものだけ。
    """
    assert suggest_split("田村") == ("田村", "")
    assert suggest_split("志村賢一衆議院議員政策担当秘書")[0] == "志村賢一"
    assert suggest_split("西村香介") == ("西村香介", "")


def test_suggest_split_leaves_a_name_of_at_least_two_characters():
    """削りすぎるくらいなら分けない。"""
    assert suggest_split("市長") == ("市長", ""), "全部消してはいけない"


def test_suggest_roster_rows_respects_lines_already_split():
    """すでに「名前(役職)」で書いてある行は、推測で壊さない。"""
    rows = suggest_roster_rows(chr(10).join(
        ["佐藤太郎(株式会社ABC 部長)", "三ツ林衆議院議員"]))
    assert rows[0] == ("佐藤太郎", "株式会社ABC 部長"), rows
    assert rows[1] == ("三ツ林", "衆議院議員"), rows


def test_roster_line_accepts_nested_parentheses():
    """**役職に括弧が入れ子で出ると、行全体が名前になっていた。**

    「企画官（命）学校法人経営指導室長」で実際に起きた(2026-08-18)。
    """
    sp = parse_roster(
        "山本学(文科省 高等教育局私学部参事官付 企画官（命）学校法人経営指導室長)")[0]
    assert sp.name == "山本学", sp.name
    assert sp.note.endswith("学校法人経営指導室長"), sp.note


def test_speaker_label_can_add_the_role():
    """本文の【 】に役職も入れるかを選べる(設計書 §11.8)。"""
    proj = _base_for_insert()
    proj.speakers[0].name = "山本学"
    proj.speakers[0].note = "文科省 高等教育局私学部参事官付 企画官"
    sid = proj.speakers[0].id
    assert speaker_label(proj, sid, False) == "山本学"
    assert speaker_label(proj, sid, True).startswith("山本学(文科省")
    assert speaker_label(proj, None, True) == UNKNOWN_LABEL


def test_inline_form_never_adds_the_role(tmp_path=None):
    """**埋め込みの ( ) は常に名前だけ。**本文の中なので長いと意味がない。"""
    import tempfile
    proj = _base_for_insert()
    proj.speakers[1].name = "山本学"
    proj.speakers[1].note = "文科省 高等教育局私学部参事官付 企画官"
    _add_at(proj, proj.segments[1], 3, 103.0, "はい", proj.speakers[1].id)
    with tempfile.TemporaryDirectory() as d:
        out = write_text(proj, Path(d) / "a.txt",
                         insert_style=INSERT_STYLE_INLINE, with_role=True)
        body = out.read_text(encoding="utf-8")
    assert "(山本学：はい)" in body, body
    assert "(山本学(文科省" not in body, "埋め込みに役職が入っている"
    assert "【山本学(文科省" in body or "山本学" in body


def test_short_time_drops_the_leading_hours():
    """候補の選択肢に出す短い時刻(設計書 §10.3.2)。

    HH:MM:SS だと幅に収まらず「00:00:29 声!」と切れた(実機・2026-08-19)。
    **1 時間を超えたら時も出す**——落とすと別の場所と区別できない。
    """
    from src.assign_gui import fmt_short_time
    assert fmt_short_time(29) == "0:29"
    assert fmt_short_time(1504) == "25:04"
    assert fmt_short_time(3600) == "1:00:00"
    assert fmt_short_time(3930) == "1:05:30"
    assert fmt_short_time(-1) == "0:00", "負でも落ちない"
    # 選択肢の幅(11 文字)に収まること。「1:05:30 声B」で 11 文字
    assert len(fmt_short_time(3930) + " 声B") <= 11


# --- 編集履歴（設計書 §12）--------------------------------------------
# 事業計画 v29 で差別化はこの一点に絞られた
#   「どこまで人が原音で確認したかを成果物に残せるか」

def _logged():
    """履歴の検査用。**割当を空に戻してから始める。**

    _base_for_insert() は 1 区間を割当済みにしてあり、そのままだと
    「変わらないものは記録しない」規則に当たって記録が出ない。
    """
    proj = _base_for_insert()
    for s in proj.segments:
        s.speaker_id, s.reviewed = None, False
    proj.edit_log.clear()
    return proj


def test_nothing_is_recorded_when_nothing_changed():
    """**変わらないものは記録しない。**履歴が水増しされると読めなくなる。

    _base_for_insert() が 1 区間を割当済みにしているので、そのまま同じ話者を
    当てても記録は出ない（この規則を実際に踏んで気付いた・2026-08-19）。
    """
    proj = _base_for_insert()
    proj.edit_log.clear()
    assert proj.segments[1].speaker_id == "sp01"
    proj.assign_speaker(1, "sp01")
    proj.apply_speaker_to([1], "sp01", heard_index=1)
    proj.edit_text(1, proj.segments[1].text)
    assert proj.edit_log == [], proj.edit_log


def test_assign_is_recorded_with_the_previous_value():
    """**before を残す。**「機械が何と言い、人が何に直したか」が要る。"""
    proj = _logged()
    sid = proj.speakers[0].id
    proj.assign_speaker(1, sid)
    rec = proj.edit_log[-1]
    assert rec["op"] == "assign"
    assert rec["before"] is None and rec["after"] == sid
    assert rec["reviewed"] is True
    assert rec["actor"] == "user"
    assert proj.segments[1].speaker_id == sid


def test_assign_does_not_record_a_no_op():
    """変わっていないなら記録しない（履歴が水増しされる）。"""
    proj = _logged()
    sid = proj.speakers[0].id
    proj.assign_speaker(1, sid)
    n = len(proj.edit_log)
    proj.assign_speaker(1, sid)
    assert len(proj.edit_log) == n


def test_bulk_apply_is_one_record_not_one_per_segment():
    """**記録するのは人の判断。**一括適用は 50 区間変えても判断は 1 回。

    区間ごとに 1 件ずつ残すと実データで 95KB になり(実測 2026-08-19)、
    しかも「何回判断したか」が読み取れなくなる。
    """
    proj = _logged()
    sid = proj.speakers[0].id
    proj.apply_speaker_to([0, 1, 2], sid, heard_index=1)
    assert len(proj.edit_log) == 1
    rec = proj.edit_log[0]
    assert rec["op"] == "assign_bulk" and rec["count"] == 3
    assert len(rec["heard"]) == 1 and len(rec["bulk"]) == 2


def test_bulk_apply_keeps_the_reviewed_distinction():
    """**聴いた 1 区間だけ ✓。**残りは △（機械の結果を当てただけ）。

    この区別が製品価値そのもの（CLAUDE.md）。
    """
    proj = _logged()
    sid = proj.speakers[0].id
    proj.apply_speaker_to([0, 1, 2], sid, heard_index=1)
    assert proj.segments[1].reviewed is True
    assert proj.segments[0].reviewed is False
    assert proj.segments[2].reviewed is False
    # 記録の側でも区別できる
    rec = proj.edit_log[0]
    from src.segments import segment_key
    assert rec["heard"] == [list(segment_key(proj.segments[1]))]


def test_undo_is_recorded_too():
    """**取り消したことも残す。**消すと「当てたが戻した」経緯が読めない。"""
    proj = _logged()
    sid = proj.speakers[0].id
    snap = [(1, proj.segments[1].speaker_id, proj.segments[1].reviewed)]
    proj.assign_speaker(1, sid)
    proj.restore_assignments(snap)
    assert proj.segments[1].speaker_id is None
    assert [r["op"] for r in proj.edit_log] == ["assign", "undo_assign"]


def test_undo_ignores_indexes_that_no_longer_exist():
    """分割・結合のあとでも落ちない（戻せないぶんは触らない）。"""
    proj = _logged()
    proj.restore_assignments([(999, "sp01", True)])
    assert proj.edit_log == []


def test_text_edit_keeps_both_sides():
    proj = _logged()
    before = proj.segments[1].text
    assert proj.edit_text(1, "直しました") is True
    rec = proj.edit_log[-1]
    assert rec["op"] == "edit_text"
    assert rec["before"] == before and rec["after"] == "直しました"
    assert proj.segments[1].text_edited is True, "再実行で上書きされてしまう"
    assert proj.edit_text(1, "直しました") is False, "同じ内容で記録している"


def test_time_edit_records_before_and_after():
    proj = _logged()
    seg = proj.segments[1]
    b = [round(seg.start, 3), round(seg.end, 3)]
    assert proj.edit_time(1, seg.start + 0.5, seg.end, reviewed=True) is True
    rec = proj.edit_log[-1]
    assert rec["op"] == "edit_time"
    assert rec["before"] == b and rec["after"][0] == round(b[0] + 0.5, 3)
    assert rec["reviewed"] is True


def test_time_edit_records_a_confirmation_without_a_change():
    """値が同じでも ✎△ → ✎ は変化。**確認したことが履歴の中身。**"""
    proj = _logged()
    seg = proj.segments[1]
    assert proj.edit_time(1, seg.start, seg.end, reviewed=True) is True
    assert proj.edit_log[-1]["reviewed"] is True


def test_log_uses_the_segment_key_not_index():
    """**鍵は (orig_start, start) の組。**

    `index` は分割・結合・再実行で振り直る。**`orig_start` だけでも足りない**
    ——分割した 2 つが同じ値を持つため(設計書 §10.3.4)。
    """
    from src.segments import segment_key
    proj = _logged()
    key = list(segment_key(proj.segments[1]))
    proj.assign_speaker(1, proj.speakers[0].id)
    proj.segments[1].index = 99
    assert proj.edit_log[-1]["segment"] == key
    assert "orig_start" not in proj.edit_log[-1], "古い形が残っている"


def test_split_halves_get_different_log_keys():
    """分割した 2 つの記録が、どちらのものか区別できること。"""
    from src.segments import segment_key
    proj = _logged()
    head, tail = proj.split_segment(1, 105.0, 5)
    proj.edit_log.clear()
    proj.assign_speaker(head.index, proj.speakers[0].id)
    proj.assign_speaker(tail.index, proj.speakers[1].id)
    a, b = proj.edit_log[0]["segment"], proj.edit_log[1]["segment"]
    assert a[0] == b[0], "前提: 分割は orig_start を共有する"
    assert a != b, "**記録がどちらの区間か区別できていない**"


def test_verification_summary_shows_the_edit_history():
    """**検証要約に履歴の行を出す。**差別化はこの一点(事業計画 v29)。"""
    proj = _logged()
    proj.apply_speaker_to([0, 1], proj.speakers[0].id, heard_index=1)
    proj.edit_text(1, "直しました")
    rows = dict(build_verification(proj, 1))
    assert "編集の履歴" in rows
    body = rows["編集の履歴"]
    assert "全 2 件" in body
    assert "話者(まとめて) 1" in body and "本文 1" in body
    assert "最終" in body, "**いつまで手を入れたかが分からない**"


def test_summary_says_so_when_there_is_no_history():
    """古い作業ファイルには履歴が無い。**黙って空欄にしない。**"""
    assert build_log_summary(Project(audio_path="x.m4a")) == "記録なし"


def test_summary_does_not_hide_an_unknown_operation():
    """**知らない op も名前のまま出す。**畳むと履歴から消えたように見える。"""
    proj = _logged()
    proj._log("brand_new_op", orig_start=1.0)
    assert "brand_new_op 1" in build_log_summary(proj)


def test_log_labels_cover_every_op_we_write():
    """記録する op には日本語名を用意しておく（生の英語を出さない）。"""
    import re
    src = (Path(__file__).resolve().parent.parent
           / "src" / "segments.py").read_text(encoding="utf-8")
    ops = set(re.findall(r'_log\(\s*"([a-z_]+)"', src))
    ops |= set(re.findall(r'"op": "([a-z_]+)"', src))
    missing = sorted(ops - set(LOG_LABELS))
    assert missing == [], f"日本語名が無い op: {missing}"


def test_log_time_is_shown_in_local_time():
    """UTC で持ち、表示は手元の時刻に直す（読む人は現地時刻で読む）。"""
    assert fmt_log_at("") == ""
    got = fmt_log_at("2026-08-19T05:03:40+00:00")
    assert len(got) == 16 and got.startswith("2026-08-19"), got
    assert fmt_log_at("こわれている") == "こわれている", "落ちてはいけない"


def test_gui_never_writes_the_fields_directly():
    """**画面から直接の書き換えを禁じる。**

    seg.speaker_id = … と画面で書くと編集履歴を素通りする。経路が増える
    たびに記録し忘れるので、Project のメソッドを必ず通す(設計書 §1.3)。
    **この検査は、うっかり直接書いたときに落ちるためにある。**
    """
    import re
    src = (Path(__file__).resolve().parent.parent
           / "src" / "assign_gui.py").read_text(encoding="utf-8")
    banned = re.compile(
        r"^\s*\w+\.(speaker_id|reviewed|text_edited|time_edited|"
        r"time_reviewed)\s*=\s*(?!=)", re.M)
    hits = [m.group(0).strip() for m in banned.finditer(src)]
    assert hits == [], f"画面が直接書いている: {hits}"


def test_bulk_times_are_one_record():
    """時刻のまとめて適用も 1 件（百件の判断ではない）。"""
    proj = _logged()
    n = proj.apply_times_to(
        [(0, 90.5, 100.0), (2, 110.5, 120.0)], reviewed=False)
    assert n == 2
    assert len(proj.edit_log) == 1
    rec = proj.edit_log[0]
    assert rec["op"] == "apply_times_bulk" and rec["count"] == 2
    assert rec["reviewed"] is False, "**まとめて当てただけは ✎△**"
    assert proj.segments[0].time_reviewed is False


def test_revert_time_is_recorded():
    proj = _logged()
    proj.edit_time(1, 101.0, 111.0, reviewed=True)
    assert proj.revert_time(1) is True
    assert proj.edit_log[-1]["op"] == "revert_time"
    assert proj.segments[1].time_edited is False
    assert proj.revert_time(1) is False, "直っていないのに戻している"


def test_undo_times_is_recorded():
    """**戻したことも残す。**"""
    proj = _logged()
    snap = [(1, proj.segments[1].start, proj.segments[1].end, False, False)]
    proj.edit_time(1, 101.0, 111.0, reviewed=True)
    assert proj.restore_times(snap) == 1
    assert proj.edit_log[-1]["op"] == "undo_times"
    assert proj.segments[1].time_edited is False


def test_clearing_a_removed_speaker_is_recorded():
    """名簿から人を消して割当が外れた経緯も残す。

    あとから読む人にとって「誰の発言か分からなくなった理由」がこれ。
    """
    proj = _logged()
    sid = proj.speakers[0].id
    proj.assign_speaker(1, sid)
    assert proj.clear_speakers([sid]) == 1
    rec = proj.edit_log[-1]
    assert rec["op"] == "clear_speakers" and rec["removed"] == [sid]
    assert proj.segments[1].speaker_id is None
    assert proj.segments[1].reviewed is False


def test_added_speaker_is_marked_heard():
    """足した発話に選んだ話者は ✓（いま聴いた直後に人が選んだ）。"""
    proj = _logged()
    seg = proj.add_utterance(103.0, 103.6, "はい", "g:D")
    proj.set_added_speaker(seg.index, proj.speakers[1].id)
    assert seg.reviewed is True
    assert proj.edit_log[-1]["op"] == "assign"
    assert proj.edit_log[-1]["reviewed"] is True


def test_log_counts_and_last_time():
    proj = _logged()
    proj.assign_speaker(1, proj.speakers[0].id)
    proj.edit_text(1, "直しました")
    counts = proj.log_counts()
    assert counts["assign"] == 1 and counts["edit_text"] == 1
    assert proj.log_last_at().startswith("20"), proj.log_last_at()
    assert Project(audio_path="x.m4a").log_last_at() == ""


def test_log_survives_save_and_load(tmp_path=None):
    """履歴が保存され、読み直しても残る（成果物に出す前提）。"""
    import tempfile
    proj = _logged()
    proj.assign_speaker(1, proj.speakers[0].id)
    with tempfile.TemporaryDirectory() as d:
        path = Path(d) / "a.speakers.json"
        proj.json_path = str(path)
        proj.save()
        again = Project.load(str(path))
    assert [r["op"] for r in again.edit_log] == ["assign"]
    assert again.edit_log[0]["before"] is None


def test_legend_only_when_something_was_inserted():
    """凡例は使われたときだけ。無関係な凡例は記録を汚す。"""
    proj = _base_for_insert()
    assert has_inserted_utterances(proj) is False
    added = _add_at(proj, proj.segments[1], 3, 103.0, "はい", "sp02")
    assert has_inserted_utterances(proj) is True
    proj.remove_added_utterance(added.index)
    assert has_inserted_utterances(proj) is False, "消したのに凡例が出る"


def test_legend_without_a_cut_is_not_shown():
    """位置を持たない足し方は本文に差し込まれないので、凡例も要らない。"""
    proj = _base_for_insert()
    proj.add_utterance(103.0, 103.6, "うん", "g:D")
    assert has_inserted_utterances(proj) is False


def test_legend_explains_each_style():
    """記号だけ出しても受け取った人が読めない。両方に説明がある。"""
    assert CONTINUE_MARK in INSERT_LEGEND[INSERT_STYLE_LINE]
    assert "BTSJ" in INSERT_LEGEND[INSERT_STYLE_LINE]
    assert "BTSJ" in INSERT_LEGEND[INSERT_STYLE_INLINE]
    # 検証要約に「凡例」という項目があるので、見出しをそれと重ねない
    assert not any(v.startswith("凡例") for v in INSERT_LEGEND.values())


def test_write_text_carries_the_style(tmp_path=None):
    """テキスト出力にも同じ選択が効く(凡例・,,・後半の時刻なし)。"""
    import tempfile
    proj = _base_for_insert()
    _add_at(proj, proj.segments[1], 3, 103.0, "はい", "sp02")
    with tempfile.TemporaryDirectory() as d:
        out = write_text(proj, Path(d) / "a.txt", insert_style=INSERT_STYLE_LINE)
        body = out.read_text(encoding="utf-8")
        assert INSERT_LEGEND[INSERT_STYLE_LINE] in body
        assert "ながら" + CONTINUE_MARK in body
        # 後半の行は時刻の桁を空ける（[ で始まらない）
        tail = [ln for ln in body.splitlines() if "けれどもそれで" in ln][0]
        assert not tail.lstrip().startswith("["), tail
        assert tail.startswith(" "), tail

        out2 = write_text(proj, Path(d) / "b.txt", insert_style=INSERT_STYLE_INLINE)
        body2 = out2.read_text(encoding="utf-8")
        assert "(山本：はい)" in body2, body2
        assert CONTINUE_MARK not in body2, body2


def test_remove_added_utterance_tells_the_two_apart():
    """**足した区間と、転写が出した区間を、記録の上で区別する。**

    どちらも人が明示的に消せる（自動削除をしない原則は「自動」に対する
    もの・2026-08-22 に実機の指摘で見直した）。ただし意味が違う——
    前者は自分の入力の取り消し、後者は機械の出力を人の判断で落とすこと。
    """
    proj = _splittable()
    seg = proj.add_utterance(105.0, 105.8, "はい")
    proj.remove_added_utterance(0)                   # ASR 由来の区間
    assert proj.edit_log[-1]["op"] == "remove_segment"
    assert "text" in proj.edit_log[-1], "消した中身が残っていない"
    seg = next(s for s in proj.segments if proj.is_added_utterance(s))
    proj.remove_added_utterance(seg.index)
    assert proj.edit_log[-1]["op"] == "remove_added_utterance"
    assert [s.index for s in proj.segments] == list(range(len(proj.segments)))
    # 消したら鍵も落ちる(再実行の引き継ぎが古い鍵を掴まないように)
    assert proj.added_utterance_keys() == set()


def test_added_keys_survive_save_and_load(tmp_path=None):
    """保存して読み直しても、人が足した区間だと分かる。

    edit_log は本体 JSON に入り、再実行でも引き継がれる。
    """
    import tempfile
    proj = _splittable()
    seg = proj.add_utterance(105.0, 105.8, "はい")
    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "x.speakers.json"
        proj.json_path = str(p)
        proj.save()
        again = Project.load(p)
    assert again.added_utterance_keys() == {round(seg.orig_start, 3)}
    got = [s for s in again.segments if s.text == "はい"]
    assert len(got) == 1 and again.is_added_utterance(got[0]) is True


def test_split_segment_basics():
    proj = _splittable()
    head, tail = proj.split_segment(1, 105.0, len("そうですね"))
    assert len(proj.segments) == 4
    assert [s.index for s in proj.segments] == [0, 1, 2, 3]     # 振り直される
    assert (head.start, head.end) == (100.0, 105.0)
    assert (tail.start, tail.end) == (105.0, 110.0)
    assert (head.text, tail.text) == ("そうですね", "いいと思います")
    # 前半は元の声のまとまりと話者を維持、後半は擬似不明で未確定
    assert head.cluster == "0:B" and head.speaker_id == "sp01"
    assert tail.cluster == "0:?" and tail.speaker_id is None
    assert tail.is_pseudo_cluster is True
    # 範囲が変わったので、どちらも聴き直し対象
    assert head.reviewed is False and tail.reviewed is False
    assert head.time_edited is True and tail.time_edited is True
    # 再実行時の系譜キーは親の値を両方が共有する
    assert (head.orig_start, head.orig_end) == (100.0, 110.0)
    assert (tail.orig_start, tail.orig_end) == (100.0, 110.0)
    # 隣の区間には触らない
    assert (proj.segments[0].start, proj.segments[3].start) == (90.0, 110.0)


def test_split_segment_clamps_boundary_and_cut():
    proj = _splittable()
    # 区間の外を指されても内側に収める(前後に最短の長さを残す)
    head, tail = proj.split_segment(1, 999.0, 999)
    assert head.end == 109.9 and tail.start == 109.9
    assert tail.text == ""                     # cut は本文の長さで頭打ち

    proj = _splittable()
    head, tail = proj.split_segment(1, 0.0, -5)
    assert head.end == 100.1 and head.text == ""


def test_split_segment_rejects_too_short():
    proj = _splittable()
    proj.segments[1].end = proj.segments[1].start + 0.15    # 2 つに割れない
    try:
        proj.split_segment(1, 100.05, 3)
    except ValueError:
        return
    raise AssertionError("短すぎる区間が分割できてしまった")


def test_merge_segments_basics():
    proj = _splittable()
    proj.segments[1].speaker_id = "sp01"
    proj.segments[2].speaker_id = "sp01"
    proj.segments[2].note = "あとのメモ"
    merged = proj.merge_segments(1)
    assert len(proj.segments) == 2
    assert [s.index for s in proj.segments] == [0, 1]
    assert (merged.start, merged.end) == (100.0, 120.0)
    assert merged.text == "そうですねいいと思います次の発言"   # 空白を挟まない
    assert merged.cluster == "0:B"                            # 前側を採用
    assert merged.speaker_id == "sp01"                        # 同じ話者なら維持
    assert merged.reviewed is False
    assert merged.time_edited is True
    assert merged.note == "めも / あとのメモ"
    # 系譜は前側の始まりと後側の終わり(再実行時の吸収に要る)
    assert (merged.orig_start, merged.orig_end) == (100.0, 120.0)


def test_merge_segments_drops_conflicting_speaker():
    proj = _splittable()
    proj.segments[1].speaker_id = "sp01"
    proj.segments[2].speaker_id = "sp02"
    merged = proj.merge_segments(1)
    assert merged.speaker_id is None      # どちらが正しいかは機械には決められない


def test_merge_segments_rejects_last():
    proj = _splittable()
    try:
        proj.merge_segments(len(proj.segments) - 1)
    except ValueError:
        return
    raise AssertionError("次の区間が無いのに結合できてしまった")


def test_split_then_merge_restores_shape():
    """分割の取り消しは結合で行う(Ctrl+Z の対象外)。"""
    proj = _splittable()
    before = proj.segments[1]
    span, text = (before.start, before.end), before.text
    proj.split_segment(1, 105.0, 5)
    proj.merge_segments(1)
    after = proj.segments[1]
    assert len(proj.segments) == 3
    assert (after.start, after.end) == span
    assert after.text == text
    assert (after.orig_start, after.orig_end) == span


# ======================================================================
# 時刻を耳で確かめたか(time_reviewed。schema 4)
# ======================================================================

def test_schema3_file_treats_edited_as_reviewed():
    """v3 までは『時刻を直した = 耳で合わせた』しかなかった。"""
    old = {
        "schema": 3,
        "audio_path": "a.m4a",
        "segments": [
            {"index": 0, "start": 16.0, "end": 26.0, "text": "あ", "cluster": "0:A",
             "orig_start": 10.0, "orig_end": 20.0, "time_edited": True},
            {"index": 1, "start": 30.0, "end": 40.0, "text": "い", "cluster": "0:A",
             "orig_start": 30.0, "orig_end": 40.0},
        ],
    }
    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "old.speakers.json"
        p.write_text(json.dumps(old, ensure_ascii=False), encoding="utf-8")
        proj = Project.load(p)
    assert proj.segments[0].time_reviewed is True      # 機械の推定ではなく手動だった
    assert proj.segments[1].time_reviewed is False


def test_unreviewed_time_survives_save_and_load():
    """『時刻は入れたが未確認』(✎△)の状態が保存で消えない。"""
    proj = Project(audio_path="a.m4a", duration=100.0)
    proj.segments = [
        Segment(index=0, start=16.0, end=26.0, text="あ", cluster="0:A",
                time_edited=True, time_reviewed=False,
                orig_start=10.0, orig_end=20.0),
        Segment(index=1, start=36.0, end=46.0, text="い", cluster="0:A",
                time_edited=True, time_reviewed=True,
                orig_start=30.0, orig_end=40.0),
    ]
    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "a.speakers.json"
        proj.save(p)
        data = json.loads(p.read_text(encoding="utf-8"))
        assert data["schema"] == SCHEMA_VERSION == 5
        again = Project.load(p)
    assert again.segments[0].time_edited is True
    assert again.segments[0].time_reviewed is False    # 既定値で上書きされない
    assert again.segments[1].time_reviewed is True


# ======================================================================
# 検証履歴の器(schema 5)
# ======================================================================

def test_v4_file_reads_with_empty_history_fields():
    """v4 以前のファイルは既定値だけで読め、保存すると v5 になる。"""
    old = {
        "schema": 4,
        "audio_path": "a.m4a",
        "duration": 100.0,
        "segments": [
            {"index": 0, "start": 0.0, "end": 5.0, "text": "あ", "cluster": "0:A"},
        ],
    }
    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "old.speakers.json"
        p.write_text(json.dumps(old, ensure_ascii=False), encoding="utf-8")
        proj = Project.load(p)
        assert proj.source_sha256 == ""
        assert proj.engine == {}
        assert proj.doc_revision == 0
        assert proj.edit_log == []
        # 開いて保存しただけでは版を進めない(開いた事実は版ではない)
        proj.save(p)
        data = json.loads(p.read_text(encoding="utf-8"))
    assert data["schema"] == 5
    assert data["doc_revision"] == 0


def test_write_docx_verification_summary():
    """Word の末尾に検証要約(SHA-256・処理経路・版・確認状態)が付く。"""
    from docx import Document
    from src.segments import write_docx

    proj = _make_project()
    sid = proj.speakers[0].id
    proj.segments[0].speaker_id = sid
    proj.segments[0].reviewed = True
    proj.segments[1].speaker_id = sid          # まとめて適用(未確認)扱い
    proj.segments[0].time_edited = True
    proj.segments[0].time_reviewed = True
    proj.source_sha256 = "cd" * 32
    proj.engine = {"mode": "cloud", "model": "gemini-2.5-flash",
                   "at": "2026-08-14T05:00:00Z"}
    with tempfile.TemporaryDirectory() as d:
        out = Path(d) / "out.docx"
        write_docx(proj, out, revision=3)
        text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "検証要約" in text
    assert "cd" * 32 in text
    assert "クラウド / gemini-2.5-flash" in text
    assert "revision 3" in text
    assert "聴いて確定 1 区間" in text and "まとめて適用 1 区間" in text
    assert "聴いて確認 1 区間" in text              # 時刻の ✎
    assert "凡例" in text and "区間の数です" in text
    assert "保証するものではありません" in text


def test_verification_never_reads_as_a_fraction():
    """内訳を「/」で区切らない。分数に見えて読み違えられる。

    実出力で「時刻の修正: 聴いて確認 41 / 適用のみ 40」となり、41 が 40 を
    超える分数に見えた(事業計画 v29 §5)。検証要約は確認の履歴そのものなので、
    読み違えられる表示はそれ自体が信用を損なう。分母(全区間数)を明示する。
    """
    from src.segments import build_verification

    proj = _make_project()
    for i, s in enumerate(proj.segments):
        s.time_edited = True
        s.time_reviewed = i % 2 == 0          # 半々にして両方 0 でなくする
    rows = dict(build_verification(proj, 1))

    total = proj.total_count
    for key in ("話者の確認", "時刻の修正"):
        value = rows[key]
        # 数と数の間に「/」が来ないこと(処理経路の「/」は別の行なので無関係)
        assert " / " not in value, f"{key} が分数に見える: {value}"
        assert f"全 {total} 区間" in value, f"{key} に分母がない: {value}"


def test_verification_counts_add_up():
    """内訳の合計が分母を超えない(超えたら数え方が壊れている)。"""
    from src.segments import build_verification

    proj = _make_project()
    proj.segments[0].speaker_id = proj.speakers[0].id
    proj.segments[0].reviewed = True
    proj.segments[1].speaker_id = proj.speakers[0].id
    for s in proj.segments[:2]:
        s.time_edited = True
    proj.segments[0].time_reviewed = True

    rows = dict(build_verification(proj, 1))
    nums = [int(x) for x in re.findall(r"(\d+) 区間", rows["話者の確認"])]
    total, breakdown = nums[0], nums[1:]
    assert sum(breakdown) == total
    t_nums = [int(x) for x in re.findall(r"(\d+) 区間", rows["時刻の修正"])]
    assert t_nums[0] == proj.total_count          # 分母は全区間
    assert t_nums[1] == t_nums[2] + t_nums[3]     # 直した数 = 確認 + 適用のみ
    assert t_nums[1] <= t_nums[0]                 # 分子は分母を超えない


def test_write_docx_verification_can_be_omitted():
    from docx import Document
    from src.segments import write_docx

    proj = _make_project()
    with tempfile.TemporaryDirectory() as d:
        out = Path(d) / "out.docx"
        write_docx(proj, out, include_verification=False)
        text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "検証要約" not in text


def test_history_fields_round_trip():
    """v5 の 4 フィールドが保存・読込で欠けない。"""
    proj = Project(audio_path="a.m4a", duration=10.0)
    proj.segments = [Segment(index=0, start=0.0, end=5.0, text="あ",
                             cluster="0:A")]
    proj.source_sha256 = "ab" * 32
    proj.engine = {"mode": "cloud", "model": "gemini-2.5-flash",
                   "app_version": "2.1.0", "at": "2026-08-14T05:00:00Z"}
    proj.doc_revision = 3
    proj.edit_log = [{"at": "2026-08-14T05:01:00Z", "actor": "user",
                      "kind": "time", "target": 0.0,
                      "before": [0.0, 5.0], "after": [1.0, 6.0]}]
    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "a.speakers.json"
        proj.save(p)
        again = Project.load(p)
    assert again.source_sha256 == "ab" * 32
    assert again.engine["model"] == "gemini-2.5-flash"
    assert again.doc_revision == 3
    assert again.edit_log[0]["kind"] == "time"


def test_split_inherits_time_reviewed():
    """境界は人が決めるが、外側の端は親のまま。分割で確認済みには昇格させない。"""
    proj = _splittable()
    proj.segments[1].time_edited = True
    proj.segments[1].time_reviewed = True
    head, tail = proj.split_segment(1, 105.0, 5)
    assert head.time_reviewed is True and tail.time_reviewed is True

    proj = _splittable()                        # 親が未確認(✎△ や AI の時刻のまま)
    head, tail = proj.split_segment(1, 105.0, 5)
    assert head.time_edited is True and tail.time_edited is True
    assert head.time_reviewed is False and tail.time_reviewed is False


def test_merge_needs_both_sides_reviewed():
    """片側でも未確認なら、結合した区間も未確認。"""
    for a_ok, b_ok, want in ((True, True, True), (True, False, False),
                             (False, True, False), (False, False, False)):
        proj = _splittable()
        for seg, ok in ((proj.segments[1], a_ok), (proj.segments[2], b_ok)):
            seg.time_edited = True
            seg.time_reviewed = ok
        assert proj.merge_segments(1).time_reviewed is want


# ======================================================================
# 再実行時の引き継ぎ(carry-over)
# ======================================================================

def _carry(old_segments, new_segments):
    """_carry_over_assignments を単体で回す。戻り値は (新しい区間リスト, ログ)"""
    from src.pipeline import _carry_over_assignments      # google-genai を引くので局所 import

    old = Project(audio_path="a.m4a")
    old.segments = old_segments
    logs: list[str] = []
    return _carry_over_assignments(old, new_segments, logs.append), logs


def test_carry_over_does_not_lose_a_neighbour_to_a_greedy_match():
    """**発言が黙って消えないこと。**実データで 2 件失われていた。

    突き合わせが「新しい区間の順に、許容 2 秒以内なら当てる」貪欲だったため、
    **本当の相手より手前の区間に先に当たっていた。**当たった区間は旧区間で
    置き換えられて消え、本当の相手は「使用済み」で当たらず残る——
    **欠落と重複が同じ 1 つの原因から出ていた**(2026-08-19・§10.3.5)。

    実データの形をそのまま写してある(ずれ 1.70 秒)。
    """
    old = [
        Segment(index=0, start=2110.50, end=2132.30, text="そこが県さんが",
                cluster="0:A", speaker_id="sp01", reviewed=True,
                time_edited=True, orig_start=2110.51, orig_end=2132.29),
    ]
    new = [
        Segment(index=0, start=2108.81, end=2110.51,
                text="そういうのがあるんですよ", cluster="0:A"),
        Segment(index=1, start=2110.51, end=2132.29, text="そこが県さんが",
                cluster="0:A"),
    ]
    result, _ = _carry(old, new)
    texts = [s.text for s in result]
    assert "そういうのがあるんですよ" in texts, texts
    assert len(result) == 2, texts
    assert result[1].speaker_id == "sp01", "人の割当が別の区間に付いた"


def test_carry_over_matches_the_closest_one():
    """**いちばん近いもの同士で当てる。**手前から順に当てない。"""
    old = [
        Segment(index=0, start=105.0, end=115.0, text="あ", cluster="0:A",
                speaker_id="sp01", reviewed=True, time_edited=True,
                orig_start=105.0, orig_end=115.0),
    ]
    new = [
        Segment(index=0, start=103.5, end=105.0, text="手前", cluster="0:A"),
        Segment(index=1, start=105.0, end=115.0, text="あ", cluster="0:A"),
        Segment(index=2, start=115.0, end=120.0, text="次", cluster="0:B"),
    ]
    result, _ = _carry(old, new)
    assert [s.text for s in result] == ["手前", "あ", "次"], [
        s.text for s in result]
    assert result[1].speaker_id == "sp01"


def test_carry_over_does_not_duplicate_at_the_exact_start():
    """**区間が二重にならない（再実行の取り込みの境界）。**

    実データで、まったく同じ本文・同じ時刻の区間が 2 組できていた
    （2026-08-19 に実機の指摘から発見。設計書 §10.3.5）。原因は取り込み
    判定が `lo < start < hi` と両端を含まなかったこと。**開始が
    ちょうど lo と同じ再生成区間が取り込まれず、そのまま残っていた。**

    1 つ目は照合で使われて `used` に入るので、2 つ目はどの家族にも
    当たらず、素通りする。
    """
    old = [
        Segment(index=0, start=100.0, end=120.0, text="本文", cluster="0:A",
                speaker_id="sp01", reviewed=True, time_edited=True,
                orig_start=100.0, orig_end=120.0),
    ]
    new = [
        Segment(index=0, start=100.0, end=120.0, text="本文", cluster="0:A"),
        Segment(index=1, start=100.0, end=120.0, text="本文", cluster="0:A"),
    ]
    result, _ = _carry(old, new)
    assert len(result) == 1, [
        (s.start, s.speaker_id, s.text) for s in result]
    assert result[0].speaker_id == "sp01", "人の割当が残っていない"


def test_carry_over_does_not_swallow_the_next_segment():
    """**終わりの側は含めない。**含めると隣の区間まで飲み込む。

    区間の終了と次の区間の開始が一致するのが普通なので、`hi` を
    「以下」にすると次の発言が消える。
    """
    old = [
        Segment(index=0, start=100.0, end=110.0, text="本文", cluster="0:A",
                speaker_id="sp01", reviewed=True, time_edited=True,
                orig_start=100.0, orig_end=110.0),
    ]
    new = [
        Segment(index=0, start=100.0, end=110.0, text="本文", cluster="0:A"),
        Segment(index=1, start=110.0, end=120.0, text="次の発言", cluster="0:B"),
    ]
    result, _ = _carry(old, new)
    assert [s.text for s in result] == ["本文", "次の発言"], [
        s.text for s in result]


def test_carry_over_keeps_the_next_chunk_head_at_a_split_boundary():
    """**開始が一致するだけの別区間を消さないこと。**

    2026-08-28 に test_pipeline_integration で発覚。再実行で発言が 1 件、
    黙って消えていた(12 区間 → 11 区間)。

    `lo <= start < hi` で二重化を直した副作用(§10.3.5・3 回目)。
    分割した子は**親の系譜(orig)を継ぐ**ので範囲が狭い。その範囲の lo が
    チャンク境界と重なると、次のチャンクの先頭区間は開始だけが一致し、
    範囲はまったく違うのに取り込まれて消える。

    実データの形: チャンク 0 の最終区間 orig(60.03, 60.33) を分割 →
    子 2 つが同じ系譜を持つ。チャンク 1 の先頭は orig(60.03, 80.03)。
    """
    old = [
        # 分割した前半・後半。どちらも親の系譜 (60.03, 60.33) を継ぐ
        Segment(index=0, start=60.03, end=60.20, text="一点よ", cluster="0:C",
                time_edited=True, orig_start=60.03, orig_end=60.33),
        Segment(index=1, start=60.20, end=60.33, text="落ちていた発言",
                cluster="0:?", time_edited=True, text_edited=True,
                orig_start=60.03, orig_end=60.33),
    ]
    new = [
        Segment(index=0, start=60.03, end=60.33, text="一点よろしいですか。",
                cluster="0:C"),
        # チャンク 1 の先頭。開始は同じだが 20 秒の別発言
        Segment(index=1, start=60.03, end=80.03, text="議事を始めます。",
                cluster="1:A"),
    ]
    result, _ = _carry(old, new)
    texts = [x.text for x in result]
    assert "議事を始めます。" in texts, f"**次のチャンクの発言が消えている**: {texts}"


def test_carry_over_absorbs_pieces_inside_a_merged_segment():
    """結合した区間の内側は取り込む(再実行で 2 つに戻らない)。

    中点判定でも、結合区間の中にすっぽり入る断片は取り込まれること。
    """
    old = [
        Segment(index=0, start=80.03, end=120.00, text="まとめた本文",
                cluster="0:B", speaker_id="sp01", reviewed=True,
                time_edited=True, orig_start=80.03, orig_end=120.00),
    ]
    new = [
        Segment(index=0, start=80.03, end=105.03, text="前半", cluster="0:B"),
        Segment(index=1, start=105.03, end=120.00, text="後半", cluster="0:B"),
    ]
    result, _ = _carry(old, new)
    assert [x.text for x in result] == ["まとめた本文"], [x.text for x in result]


def test_carry_over_absorbs_a_tail_piece_that_overruns_the_merged_end():
    """**結合区間の末尾を数十ミリ秒はみ出す断片も取り込む。**

    単純な包含(seg.end <= hi)にすると、ここが取り込まれずに残り、
    §10.3.5 で直したはずの二重化が戻る。中点判定を選んだのはこのため。
    """
    old = [
        Segment(index=0, start=80.03, end=120.00, text="まとめた本文",
                cluster="0:B", speaker_id="sp01", reviewed=True,
                time_edited=True, orig_start=80.03, orig_end=120.00),
    ]
    new = [
        Segment(index=0, start=80.03, end=105.03, text="前半", cluster="0:B"),
        # 終わりが 0.06 秒だけ範囲を超える(チャンク末尾の丸めで実際に起きる)
        Segment(index=1, start=105.03, end=120.06, text="後半", cluster="0:B"),
    ]
    result, _ = _carry(old, new)
    assert [x.text for x in result] == ["まとめた本文"], [x.text for x in result]


def test_carry_over_warns_about_leftover_duplicates():
    """**重複が残ったら知らせる。ただし勝手に消さない。**

    本当に同じことを 2 回言った可能性もあり、逐語の記録から機械が発言を
    削るのは筋が違う（CLAUDE.md の「自動削除をしない」と同じ線）。
    人が見て決められるよう、時刻を添えて知らせるだけにする。
    """
    old = [
        Segment(index=0, start=100.0, end=120.0, text="本文", cluster="0:A",
                speaker_id="sp01", reviewed=True,
                orig_start=100.0, orig_end=120.0),   # time_edited は無し
    ]
    new = [
        Segment(index=0, start=100.0, end=120.0, text="本文", cluster="0:A"),
        Segment(index=1, start=100.0, end=120.0, text="本文", cluster="0:A"),
    ]
    result, logs = _carry(old, new)
    assert len(result) == 2, "**勝手に消している**"
    assert any("同じ時刻・同じ本文" in m for m in logs), logs
    assert any("消していない" in m for m in logs), logs


def test_carry_over_says_nothing_when_there_is_no_duplicate():
    old = [
        Segment(index=0, start=100.0, end=110.0, text="本文", cluster="0:A",
                speaker_id="sp01", orig_start=100.0, orig_end=110.0),
    ]
    new = [
        Segment(index=0, start=100.0, end=110.0, text="本文", cluster="0:A"),
        Segment(index=1, start=110.0, end=120.0, text="次", cluster="0:B"),
    ]
    _result, logs = _carry(old, new)
    assert not any("同じ時刻" in m for m in logs), logs


def test_carry_over_keeps_edited_times():
    """時刻を直した区間は orig_start で照合され、直した時刻のまま残る。"""
    old = [
        Segment(index=0, start=106.0, end=116.0, text="あ", cluster="0:A",
                speaker_id="sp01", reviewed=True, time_edited=True,
                orig_start=100.0, orig_end=110.0),
    ]
    new = [
        Segment(index=0, start=100.0, end=110.0, text="あ", cluster="0:A"),
        Segment(index=1, start=110.0, end=120.0, text="い", cluster="0:B"),
    ]
    result, _ = _carry(old, new)
    assert len(result) == 2
    assert (result[0].start, result[0].end) == (106.0, 116.0)
    assert result[0].time_edited is True
    assert result[0].speaker_id == "sp01" and result[0].reviewed is True
    assert result[1].start == 110.0        # 隣は再生成側のまま
    assert [s.index for s in result] == [0, 1]


def test_carry_over_replaces_with_split_family():
    """分割で 2 つになった区間が、再生成の 1 区間を丸ごと置き換える。"""
    old = [
        Segment(index=0, start=100.0, end=105.0, text="そうですね", cluster="0:A",
                speaker_id="sp01", reviewed=True, time_edited=True,
                orig_start=100.0, orig_end=110.0),
        Segment(index=1, start=105.0, end=110.0, text="いいと思います", cluster="0:?",
                time_edited=True, orig_start=100.0, orig_end=110.0),
    ]
    new = [
        Segment(index=0, start=100.0, end=110.0, text="そうですねいいと思います", cluster="0:A"),
        Segment(index=1, start=110.0, end=120.0, text="次の議題です", cluster="0:B"),
    ]
    result, _ = _carry(old, new)
    assert [s.text for s in result] == ["そうですね", "いいと思います", "次の議題です"]
    assert result[1].cluster == "0:?"      # 擬似クラスタが保たれる
    assert [s.index for s in result] == [0, 1, 2]


def test_carry_over_absorbs_merged_range():
    """結合した区間は、再生成で分かれて出てきたぶんを取り込む(区間が戻らない)。"""
    old = [
        Segment(index=0, start=100.0, end=110.0, text="そうですねいいと思います",
                cluster="0:A", speaker_id="sp01", time_edited=True,
                orig_start=100.0, orig_end=110.0),
    ]
    new = [
        Segment(index=0, start=100.0, end=105.0, text="そうですね", cluster="0:A"),
        Segment(index=1, start=105.0, end=110.0, text="いいと思います", cluster="0:A"),
        Segment(index=2, start=110.0, end=120.0, text="次の議題です", cluster="0:B"),
    ]
    result, logs = _carry(old, new)
    assert [s.text for s in result] == ["そうですねいいと思います", "次の議題です"]
    assert [s.index for s in result] == [0, 1]
    # 消したことをユーザーが気づけるようログに出す
    assert any("取り込みました" in m for m in logs)


def test_carry_over_legacy_behaviour_unchanged():
    """orig_* を持たない旧ファイル(= orig が start と同値)は従来どおり動く。"""
    old = [
        Segment(index=0, start=100.0, end=110.0, text="あ(手直し)", cluster="0:A",
                speaker_id="sp01", reviewed=True, note="めも", text_edited=True),
        Segment(index=1, start=110.0, end=120.0, text="い", cluster="0:B",
                speaker_id="sp02"),
    ]
    # 再実行で時刻がわずかに動いた(照合の許容 ±2 秒の内側)
    new = [
        Segment(index=0, start=101.0, end=111.0, text="あ", cluster="0:A"),
        Segment(index=1, start=111.0, end=121.0, text="い", cluster="0:B"),
    ]
    result, _ = _carry(old, new)
    assert len(result) == 2
    # 時刻には手が入っていないので、再生成側の時刻をそのまま使う
    assert (result[0].start, result[0].end) == (101.0, 111.0)
    # 人が入れた情報だけが移る
    assert result[0].speaker_id == "sp01" and result[0].reviewed is True
    assert result[0].text == "あ(手直し)" and result[0].text_edited is True
    assert result[0].note == "めも"
    assert result[1].speaker_id == "sp02"


def test_carry_over_absorbs_only_matched_families():
    """照合できなかった旧区間の範囲では吸収しない(穴が開くのを防ぐ)。"""
    old = [
        # 再生成側に対応する区間が無い(±2 秒に何も無い)結合済み区間
        Segment(index=0, start=300.0, end=330.0, text="消えた区間", cluster="0:A",
                speaker_id="sp01", time_edited=True, orig_start=300.0, orig_end=330.0),
    ]
    new = [
        Segment(index=0, start=310.0, end=320.0, text="残すべき区間", cluster="0:B"),
    ]
    result, _ = _carry(old, new)
    assert [s.text for s in result] == ["残すべき区間"]


# ======================================================================
# 話者分離(diarize.py)
#
# モデルもライブラリも要らない部分だけを見る。実モデルでの確認は
# tests/test_align_integration.py が担う(CI 対象外)。
# ======================================================================

def test_speaker_turn_round_trip():
    t = diarize.SpeakerTurn(start=1.5, end=3.25, speaker=2)
    assert diarize.SpeakerTurn.from_dict(t.to_dict()) == t


def test_turns_cache_key_includes_the_speaker_count():
    """話者数が変われば結果が変わる。同じキャッシュを共有してはいけない。"""
    a = diarize.turns_cache_path("w", "ff00", 6)
    b = diarize.turns_cache_path("w", "ff00", 9)
    assert a != b
    assert a == diarize.turns_cache_path("w", "ff00", 6)     # 同じなら同じ
    # 手動配置のモデルを差し替えたら別物
    assert diarize.turns_cache_path("w", "ff00", 6, model_dir="C:/m/a") != a
    # 指紋が取れない音声はキャッシュしない(古い結果の使い回しを防ぐ)
    assert diarize.turns_cache_path("w", "", 6) is None


def test_turns_cache_round_trip_and_version():
    with tempfile.TemporaryDirectory() as d:
        path = diarize.turns_cache_path(d, "ff00", 6)
        turns = [diarize.SpeakerTurn(0.0, 1.0, 0),
                 diarize.SpeakerTurn(0.8, 2.0, 1)]      # 重なっていてよい
        diarize.save_turns(path, turns, num_speakers=6,
                           fingerprint="ff00", duration=2.0)
        assert diarize.load_turns(path) == turns

        data = json.loads(path.read_text(encoding="utf-8"))
        data["diarize_ver"] = diarize.DIARIZE_VER + 1
        path.write_text(json.dumps(data), encoding="utf-8")
        assert diarize.load_turns(path) is None          # 実装版が違えば読まない

        path.write_text("これは JSON ではない", encoding="utf-8")
        assert diarize.load_turns(path) is None
        assert diarize.load_turns(Path(d) / "nope.json") is None


def test_missing_models_say_where_they_were_looked_for():
    """「モデルがありません」だけでは、何を置けばよいか分からない。"""
    with tempfile.TemporaryDirectory() as d:
        try:
            diarize.resolve_models(d)
        except diarize.DiarizeUnavailable as e:
            msg = str(e)
            assert d in msg                       # 探した場所
            assert "model.onnx" in msg            # 必要なファイル
            assert "DIARIZE_MODELS" in msg        # 逃げ道
        else:
            raise AssertionError("モデルが無いのに通ってしまった")


def test_default_speaker_count_is_not_automatic():
    """自動(-1)は選ばない。実測で 246 人・148 まとまりに割れた(設計書 §7)。"""
    assert diarize.DEFAULT_NUM_SPEAKERS > 0


# ======================================================================
# 話者分離の結果を区間へ落とす(segments.py の純粋関数)
# ======================================================================

def _turns(*rows):
    return [diarize.SpeakerTurn(s, e, spk) for s, e, spk in rows]


def _segs(*rows):
    return [Segment(index=i, start=s, end=e, text=t, cluster="0:?", chunk=0)
            for i, (s, e, t) in enumerate(rows)]


def test_speaker_letters_follow_the_order_of_appearance():
    """話者番号は連番とは限らない。先に話した人から A/B/C を振る。"""
    segs = _segs((0.0, 2.0, "あ"), (3.0, 5.0, "い"), (6.0, 8.0, "う"))
    turns = _turns((0.0, 2.0, 22), (3.0, 5.0, 4), (6.0, 8.0, 22))
    got = assign_speaker_clusters(segs, turns)
    assert got == ["g:A", "g:B", "g:A"]      # 22 が A、4 が B


def test_assign_takes_the_most_overlapping_speaker():
    """重なりが最も大きい話者を採る(話者区間は重なりうる)。"""
    segs = _segs((10.0, 20.0, "あ"))
    turns = _turns((0.0, 12.0, 0), (11.0, 25.0, 1))   # 2 秒 対 9 秒
    assert assign_speaker_clusters(segs, turns) == ["g:B"]


def test_assign_marks_unknown_when_overlap_is_thin():
    """重なりが足りなければ判別不能。幽霊話者はここで消える。"""
    segs = _segs((10.0, 20.0, "あ"))
    turns = _turns((19.0, 21.0, 0))          # 10 秒の区間に 1 秒しか重ならない
    assert assign_speaker_clusters(segs, turns) == ["g:?"]
    assert assign_speaker_clusters(segs, []) == ["g:?"]


def test_cluster_label_shows_global_clusters_without_chunk_numbers():
    seg = Segment(index=0, start=0.0, end=1.0, text="あ", cluster="g:A")
    assert seg.cluster_label == "声A"
    assert seg.is_pseudo_cluster is False
    ghost = Segment(index=0, start=0.0, end=1.0, text="あ", cluster="g:?")
    assert ghost.is_pseudo_cluster is True
    old = Segment(index=0, start=0.0, end=1.0, text="あ", cluster="0:A")
    assert old.cluster_label == "C1-A"        # クラウドは従来どおり


def test_merge_joins_a_sentence_split_across_segments():
    """同じ話者・文の途中・間隔が短い、の 3 つが揃ったときだけ連結する。"""
    segs = _segs((0.0, 1.0, "それでは"), (1.1, 2.0, "始めます。"),
                 (2.05, 3.0, "よろしく"))
    for s in segs:
        s.cluster = "g:A"
    out = merge_same_speaker(segs)
    assert [s.text for s in out] == ["それでは始めます。", "よろしく"]
    assert out[0].start == 0.0 and out[0].end == 2.0
    assert [s.index for s in out] == [0, 1]        # 通し番号を振り直す


def test_merge_keeps_the_original_times():
    """orig_start / orig_end は再実行時の突き合わせの鍵。連結で失わない。"""
    segs = _segs((0.0, 1.0, "それでは"), (1.1, 2.0, "始めます。"))
    for s in segs:
        s.cluster = "g:A"
    out = merge_same_speaker(segs)
    assert out[0].orig_start == 0.0 and out[0].orig_end == 2.0


def test_merge_refuses_when_the_conditions_are_not_met():
    def joined(rows, cluster="g:A", gap_ok=True):
        segs = _segs(*rows)
        for s in segs:
            s.cluster = cluster
        return [s.text for s in merge_same_speaker(segs)]

    # 文が終わっている
    assert joined([(0.0, 1.0, "終わりです。"), (1.1, 2.0, "次です")]) == \
        ["終わりです。", "次です"]
    # 間隔が空いている
    assert joined([(0.0, 1.0, "それでは"), (1.5, 2.0, "始めます")]) == \
        ["それでは", "始めます"]
    # 擬似クラスタ(判別不能)はまとめない。中身がばらばらのため
    assert joined([(0.0, 1.0, "それでは"), (1.1, 2.0, "始めます")],
                  cluster="g:?") == ["それでは", "始めます"]


def test_merge_does_not_touch_what_a_person_edited():
    """人が手を付けた区間は連結しない。その作業が消えるため。"""
    segs = _segs((0.0, 1.0, "それでは"), (1.1, 2.0, "始めます"))
    for s in segs:
        s.cluster = "g:A"
    segs[0].speaker_id = "sp01"
    assert len(merge_same_speaker(segs)) == 2

    segs2 = _segs((0.0, 1.0, "それでは"), (1.1, 2.0, "始めます"))
    for s in segs2:
        s.cluster = "g:A"
    segs2[1].text_edited = True
    assert len(merge_same_speaker(segs2)) == 2


def test_merge_joins_different_speakers_never():
    segs = _segs((0.0, 1.0, "それでは"), (1.1, 2.0, "始めます"))
    segs[0].cluster, segs[1].cluster = "g:A", "g:B"
    assert len(merge_same_speaker(segs)) == 2


# ======================================================================
# 正解ラベルの評価(evaluate.py)
#
# 測る前に手順を固定するための道具。標本の選び方が実行のたびに変わると、
# 「作りやすい区間を選んだ」という疑いを自分で否定できなくなる。
# ======================================================================

def _eval_segs(n=600):
    """長さがばらばらの区間を作る(3 層に散らばるように)。"""
    out = []
    for i in range(n):
        dur = 0.4 + (i % 6) * 0.4      # 0.4/0.8/1.2/1.6/2.0/2.4 秒
        out.append(Segment(index=i, start=i * 10.0, end=i * 10.0 + dur,
                           text=f"発言 {i}", cluster="0:?"))
    return out


def test_cer_normalizes_punctuation_but_keeps_fillers():
    """句読点の違いは誤りに数えない。フィラーは数える(測りたい対象)。"""
    assert evaluate.cer("はい、そうです。", "はいそうです") == 0.0
    # フィラーが落ちれば誤りになる
    got = evaluate.cer("あの、そうですね", "そうですね")
    assert got > 0.0


def test_cer_is_not_capped_at_one():
    """ひどく間違えた場合と、まあまあ間違えた場合を区別できること。"""
    mild = evaluate.cer("あいうえお", "あいうえか")
    awful = evaluate.cer("あい", "まったくちがうながいぶんしょう")
    assert mild < 1.0 < awful


def test_edit_distance_basics():
    assert evaluate.edit_distance("", "") == 0
    assert evaluate.edit_distance("あい", "あい") == 0
    assert evaluate.edit_distance("あい", "あいう") == 1
    assert evaluate.edit_distance("あい", "") == 2


def test_edit_ops_breakdown():
    """内訳の 3 つが、それぞれ正しく数えられているか。"""
    same = evaluate.edit_ops("あいう", "あいう")
    assert (same["sub"], same["dele"], same["ins"]) == (0, 0, 0)
    assert same["len"] == 3

    sub = evaluate.edit_ops("あいう", "あかう")          # 1 文字が別の字に
    assert (sub["sub"], sub["dele"], sub["ins"]) == (1, 0, 0)

    dele = evaluate.edit_ops("あいう", "あう")           # 1 文字が落ちた
    assert (dele["sub"], dele["dele"], dele["ins"]) == (0, 1, 0)

    ins = evaluate.edit_ops("あいう", "あいえう")        # 1 文字が増えた
    assert (ins["sub"], ins["dele"], ins["ins"]) == (0, 0, 1)

    # cer() と同じ正規化を通ること（句読点は落ちる）
    assert evaluate.edit_ops("はい、そうです。", "はいそうです")["sub"] == 0


def test_edit_ops_sums_to_cer():
    """**内訳の合計は cer() と一致する。**

    別実装にすると、内訳と CER が食い違っても気づけない。**「同じ整列で
    数えている」ことをここで縛る。**画面に出ている「誤字」に再現できる
    定義が無いのは、まさにこの縛りが無かったためである
    (HANDOFF「画面に出ている『(実測)』の 1 つは、誰も再現できない」)。
    """
    pairs = [
        ("あいうえお", "あいうえか"),
        ("あの、そうですね", "そうですね"),
        ("はい", "はいはいはい"),
        ("議事を始めます", "議事をはじめマス"),
        ("あい", "まったくちがうながいぶんしょう"),
        ("", "なにか"),
    ]
    for t, h in pairs:
        o = evaluate.edit_ops(t, h)
        if not o["len"]:
            continue
        total = (o["sub"] + o["dele"] + o["ins"]) / o["len"]
        assert abs(total - evaluate.cer(t, h)) < 1e-9, (
            f"内訳の合計が cer() と一致しません: {t!r} / {h!r} "
            f"→ 内訳 {total} vs cer {evaluate.cer(t, h)}")


def test_retention_rate_reports_both_counts():
    """保持率は 1.0 を超えうる。頭打ちにせず、実数も返す。"""
    rate, got, want = evaluate.retention_rate(
        "あの、あの、えー", "あの", evaluate.FILLER_TERMS)
    assert (got, want) == (1, 3)
    assert abs(rate - 1 / 3) < 1e-9
    rate2, got2, want2 = evaluate.retention_rate(
        "あの", "あの、あの", evaluate.FILLER_TERMS)
    assert rate2 > 1.0 and (got2, want2) == (2, 1)
    # 正解に 1 つも無ければ 0(割れない)
    assert evaluate.retention_rate("こんにちは", "あの", evaluate.FILLER_TERMS)[0] == 0.0


def test_sample_is_the_same_every_time():
    """同じ種なら必ず同じ標本になる(測定前に固定できる)。"""
    segs = _eval_segs()
    a = evaluate.select_sample(segs, per_stratum=20)
    b = evaluate.select_sample(segs, per_stratum=20)
    assert [(x.index, x.round) for x in a] == [(x.index, x.round) for x in b]
    c = evaluate.select_sample(segs, per_stratum=20, seed=999)
    assert [x.index for x in a] != [x.index for x in c]


def test_sample_takes_the_asked_number_from_each_stratum_and_round():
    segs = _eval_segs()
    picked = evaluate.select_sample(segs, per_stratum=20, rounds=2)
    per = Counter((x.stratum, x.round) for x in picked)
    for stratum in evaluate.STRATA:
        assert per[(stratum, 1)] == 20
        assert per[(stratum, 2)] == 20
    assert len({x.index for x in picked}) == 120     # 重複しない
    # 巡ごとに時間順(人が会話を追えるように)
    for r in (1, 2):
        starts = [x.start for x in picked if x.round == r]
        assert starts == sorted(starts)


def test_adding_a_round_does_not_change_the_earlier_one():
    """巡を増やしても 1 巡目の中身は動かない(固定したまま足せる)。"""
    segs = _eval_segs()
    one = evaluate.select_sample(segs, per_stratum=20, rounds=1)
    two = evaluate.select_sample(segs, per_stratum=20, rounds=2)
    first_of_two = [x for x in two if x.round == 1]
    assert [x.index for x in one] == [x.index for x in first_of_two]


def test_each_round_spreads_over_the_whole_recording():
    """1 巡だけでも全長に散らばる(途中でやめても前半に偏らない)。"""
    segs = _eval_segs()
    picked = [x for x in evaluate.select_sample(segs, per_stratum=20, rounds=2)
              if x.round == 1]
    last_start = segs[-1].start
    halves = Counter("前半" if x.start < last_start / 2 else "後半" for x in picked)
    # 母集団が一様なので、どちらの半分にも 3 割以上は入るはず
    assert halves["前半"] > len(picked) * 0.3
    assert halves["後半"] > len(picked) * 0.3


def test_sample_does_not_pad_a_small_stratum():
    """層の件数が足りないときは全件。水増ししない。"""
    segs = [Segment(index=i, start=i * 10.0, end=i * 10.0 + 5.0,
                    text="長い", cluster="0:?") for i in range(4)]
    segs.append(Segment(index=99, start=999.0, end=999.5, text="ごく短い",
                        cluster="0:?"))
    picked = evaluate.select_sample(segs, per_stratum=10, rounds=1)
    per = Counter(x.stratum for x in picked)
    assert per[evaluate.STRATUM_LONG] == 4
    assert per[evaluate.STRATUM_VERY_SHORT] == 1
    assert per[evaluate.STRATUM_SHORT] == 0


def test_strata_boundaries():
    assert evaluate.stratum_of(0.99) == evaluate.STRATUM_VERY_SHORT
    assert evaluate.stratum_of(1.0) == evaluate.STRATUM_SHORT
    assert evaluate.stratum_of(1.99) == evaluate.STRATUM_SHORT
    assert evaluate.stratum_of(2.0) == evaluate.STRATUM_LONG


def test_cluster_purity_uses_the_majority_of_each_group():
    """まとまりごとに最も多い話者へ対応づけ、一致した割合を返す。"""
    pairs = [
        ("g1", "sp01"), ("g1", "sp01"), ("g1", "sp02"),   # 2/3
        ("g2", "sp03"), ("g2", "sp03"),                    # 2/2
    ]
    purity, detail = evaluate.cluster_purity(pairs)
    assert abs(purity - 4 / 5) < 1e-9
    assert detail["g1"] == ("sp01", 2, 3)
    assert detail["g2"] == ("sp03", 2, 2)


def test_purity_has_a_floor_above_chance():
    """でたらめな正解でも純度は 1/話者数 より高く出る(下駄)。

    まとまりごとに多数派を採る性質による。併記しないと実力を読み違える。
    """
    pairs = []
    for c in range(5):                       # 5 つのまとまり × 各 20 件
        for i in range(20):
            pairs.append((f"g{c}", f"sp{i % 4:02d}"))   # 話者は 4 名
    chance = evaluate.purity_chance_level(pairs, trials=30)
    assert chance > 1 / 4          # 当てずっぽう(25%)より高い
    assert chance < 0.60           # かといって当たっているわけでもない


def test_cluster_purity_skips_unknown_truth():
    """正解が「分からない」区間は分母に入れない。"""
    purity, _ = evaluate.cluster_purity(
        [("g1", "sp01"), ("g1", None), ("g1", "sp01")])
    assert abs(purity - 1.0) < 1e-9


def test_weighted_rate_follows_the_population():
    """層ごとに同じ件数を採るので、全体値は母集団の構成で重み付けする。"""
    rates = {evaluate.STRATUM_SHORT: 0.50, evaluate.STRATUM_LONG: 0.90}
    sizes = {evaluate.STRATUM_SHORT: 670, evaluate.STRATUM_LONG: 549}
    got = evaluate.weighted_rate(rates, sizes)
    assert abs(got - (0.50 * 670 + 0.90 * 549) / 1219) < 1e-9
    # 単純平均(0.70)とは違う
    assert abs(got - 0.70) > 0.01


def test_verdict_says_undecidable_near_the_threshold():
    """70% の近くでは「判定不能」と言う。出た数字を見てから基準を動かさない。"""
    half = 0.03
    assert evaluate.verdict(0.85, half) == "達成"
    assert evaluate.verdict(0.50, half) == "未達"
    assert evaluate.verdict(0.72, half) == "判定不能"


def test_halfwidth_shrinks_with_more_samples():
    wide = evaluate.proportion_halfwidth(0.7, 200, population=1219)
    narrow = evaluate.proportion_halfwidth(0.7, 400, population=1219)
    assert narrow < wide


def test_stratified_halfwidth_uses_the_population_weights():
    """層別の精度は、層の重みの二乗で効く(単純な n では決まらない)。"""
    sizes = {evaluate.STRATUM_VERY_SHORT: 285,
             evaluate.STRATUM_SHORT: 385,
             evaluate.STRATUM_LONG: 549}
    rates = {k: 0.7 for k in sizes}
    one = evaluate.stratified_halfwidth(rates, sizes, {k: 100 for k in sizes})
    two = evaluate.stratified_halfwidth(rates, sizes, {k: 200 for k in sizes})
    assert two < one
    # 設計書 §2.3 の見込み(1 巡 ±4〜5 / 2 巡 ±3 ポイント前後)に合うこと
    assert 0.035 < one < 0.055
    assert 0.020 < two < 0.035


def test_stratified_halfwidth_is_zero_when_everything_is_measured():
    """全件を測れば標本誤差は無い(有限母集団の補正が効く)。"""
    sizes = {evaluate.STRATUM_LONG: 50}
    half = evaluate.stratified_halfwidth(
        {evaluate.STRATUM_LONG: 0.7}, sizes, {evaluate.STRATUM_LONG: 50})
    assert half == 0.0


# ======================================================================
# アダプタ境界(Utterance → Segment)
#
# クラウドもローカルも、チャンク単位で Utterance を作るところまでが
# 経路ごとの仕事。その先は共通の後段が引き受ける(設計書 §4.2)。
# ======================================================================

def test_parse_utterances_keeps_relative_times():
    """相対秒のまま返し、クラスタにチャンク番号を付けない。"""
    us = parse_utterances(
        "[00:10] 【A】 ひとつめ。\n[00:20] 【B】 ふたつめ。\n", chunk_seconds=60.0)
    assert [u.rel_start for u in us] == [10.0, 20.0]
    assert [u.cluster for u in us] == ["A", "B"]     # "0:A" ではない


def test_utterances_to_segments_adds_offset_and_namespace():
    """オフセットの足し込み・チャンク名前空間・通し番号は後段の仕事。"""
    us = [
        Utterance(rel_start=0.0, rel_end=5.0, text="あ", cluster="A"),
        Utterance(rel_start=5.0, rel_end=9.0, text="い", cluster="B"),
    ]
    segs = utterances_to_segments(
        us, chunk_index=2, offset_seconds=1800.0, start_index=40)
    assert [s.index for s in segs] == [40, 41]
    assert [s.start for s in segs] == [1800.0, 1805.0]
    assert [s.end for s in segs] == [1805.0, 1809.0]
    assert [s.cluster for s in segs] == ["2:A", "2:B"]
    assert all(s.chunk == 2 for s in segs)


def test_utterances_to_segments_does_not_redistribute():
    """按分を通さない。渡した時刻がそのまま(オフセットぶんだけずれて)出る。

    按分(redistribute_times)は Gemini のタイムスタンプがドリフトする既知バグ
    への対策であって、実測時刻にかけるものではない(設計書 §4.1)。
    ここに按分が紛れ込むと、ローカル経路で測った時刻が本文の長さで
    作り直されてしまう。
    """
    # 本文の長さがばらばらで、間も詰まっている(按分が働けば必ず動く形)
    us = [
        Utterance(rel_start=0.0, rel_end=1.0, text="あ" * 200, cluster="?"),
        Utterance(rel_start=1.0, rel_end=30.0, text="い", cluster="?"),
    ]
    segs = utterances_to_segments(us)
    assert [(s.start, s.end) for s in segs] == [(0.0, 1.0), (1.0, 30.0)]


def test_utterances_to_segments_gives_zero_length_a_minimum():
    """長さ 0 の区間は聴き直せないので、最低限の長さを与える。"""
    segs = utterances_to_segments(
        [Utterance(rel_start=12.0, rel_end=12.0, text="ん", cluster="?")])
    assert (segs[0].start, segs[0].end) == (12.0, 13.0)


# ======================================================================
# ローカル転写(local_asr.py)
#
# faster-whisper の代わりに偽のモデルを差し込む。実モデルを使う確認は
# tests/test_align_integration.py が担う(CI 対象外)。
# ======================================================================

class _FakeWord:
    def __init__(self, word, start, end):
        self.word, self.start, self.end = word, start, end


class _FakeSegment:
    def __init__(self, start, end, text, words=None):
        self.start, self.end, self.text = start, end, text
        self.words = words or []


class _FakeInfo:
    def __init__(self, duration):
        self.duration = duration


class _FakeWhisper:
    """WhisperModel の代わり。モデルもライブラリも要らない。"""

    def __init__(self, segments, duration):
        self._segments, self._duration = segments, duration
        self.calls = []

    def transcribe(self, path, **kwargs):
        self.calls.append(kwargs)
        return iter(self._segments), _FakeInfo(self._duration)


def _local(segments, duration=60.0):
    """偽モデルを差し込んだ LocalTranscriber を作る。"""
    t = LocalTranscriber(model="small")
    fake = _FakeWhisper(segments, duration)
    t._whisper = fake            # _load を通さずに差し替える
    return t, fake


def test_local_asr_keeps_short_backchannel():
    """短い相づちは残す。中身が空の区間だけ落とす。"""
    t, _ = _local([
        _FakeSegment(0.0, 3.2, " 本日はお忙しい中ありがとうございます。"),
        _FakeSegment(3.4, 3.9, " はい"),          # 0.5 秒の相づち
        _FakeSegment(4.0, 4.2, "   "),            # 中身が無い
        _FakeSegment(4.5, 6.0, " そうですね。"),
    ])
    result = t.transcribe(__file__)
    assert [u.text for u in result.utterances] == [
        "本日はお忙しい中ありがとうございます。", "はい", "そうですね。"
    ]


def test_local_asr_marks_every_cluster_unknown():
    """話者分離が入るまでは全区間が判別不能。誰の声かを決める者がいない。"""
    t, _ = _local([
        _FakeSegment(0.0, 1.0, "あ"),
        _FakeSegment(1.0, 2.0, "い"),
    ])
    result = t.transcribe(__file__)
    assert {u.cluster for u in result.utterances} == {PSEUDO_UNKNOWN}


def test_local_asr_returns_words_from_the_same_pass():
    """本文と単語時刻が 1 回の転写から取れる(別々に転写しない)。"""
    t, _ = _local([
        _FakeSegment(0.0, 1.5, " はい、そうです。", words=[
            _FakeWord(" はい", 0.0, 0.4),
            _FakeWord("、", 0.4, 0.5),
            _FakeWord("  ", 0.5, 0.5),        # 空白だけの語は捨てる
            _FakeWord("そうです", 0.6, 1.5),
        ]),
    ])
    result = t.transcribe(__file__)
    assert [w.text for w in result.words] == ["はい", "、", "そうです"]
    assert result.words[0].start == 0.0 and result.words[-1].end == 1.5
    # 時刻はチャンク先頭からの相対秒のまま(絶対秒にするのは後段の仕事)
    assert result.utterances[0].rel_start == 0.0


def test_local_asr_never_makes_negative_duration():
    """まれに終わりが始まりを下回る。負の長さは作らない。"""
    t, _ = _local([_FakeSegment(10.0, 9.5, "逆転している")])
    result = t.transcribe(__file__)
    u = result.utterances[0]
    assert u.rel_start == 10.0 and u.rel_end == 10.0


def test_local_asr_does_not_use_vad():
    """VAD は使わない。本文を作る経路なので、落ちた発言は記録から消える。"""
    t, fake = _local([_FakeSegment(0.0, 1.0, "あ")])
    t.transcribe(__file__)
    assert fake.calls[0]["vad_filter"] is False
    assert fake.calls[0]["word_timestamps"] is True
    assert fake.calls[0]["language"] == "ja"


def test_local_asr_takes_word_timestamps_by_default():
    """単語時刻は既定で取る(自動点検の実測値になる)。製品経路は常にこちら。

    切れるようにしたのは測定用——CT2 変換が不完全なモデルは単語時刻の
    位置合わせでネイティブ層ごと落ちる(kotoba-whisper-v2.0-faster で実測)。
    """
    t, fake = _local([_FakeSegment(0.0, 1.0, "あ")])
    t.transcribe(__file__)
    assert fake.calls[0]["word_timestamps"] is True
    t.transcribe(__file__, word_timestamps=False)
    assert fake.calls[1]["word_timestamps"] is False


def test_local_asr_keeps_conditioning_on_by_default():
    """直前の本文の引き継ぎは既定で入り。faster-whisper の既定と同じ。

    ここが黙って False になると、同じキャッシュキー(音声指紋 + チャンク長 +
    逐語フラグ)のまま別物の転写ができてしまう。切るのは測定専用。
    """
    t, fake = _local([_FakeSegment(0.0, 1.0, "あ")])
    t.transcribe(__file__)
    assert fake.calls[0]["condition_on_previous_text"] is True


def test_local_asr_can_turn_conditioning_off():
    """測定のために切れる(段階 2 の比較で使う)。"""
    t, fake = _local([_FakeSegment(0.0, 1.0, "あ")])
    t.condition_on_previous_text = False
    t.transcribe(__file__)
    assert fake.calls[0]["condition_on_previous_text"] is False


def test_local_asr_cancel_leaves_nothing_behind():
    """中断したら None。途中までの区間を返さない。"""
    t, _ = _local([
        _FakeSegment(0.0, 1.0, "一つめ"),
        _FakeSegment(1.0, 2.0, "二つめ"),
    ])
    seen = {"n": 0}

    def cancelled():
        seen["n"] += 1
        return seen["n"] > 1          # 2 区間目に入ったところで中断

    assert t.transcribe(__file__, is_cancelled=cancelled) is None


def test_local_asr_progress_ends_at_total():
    """進捗は (処理済み秒, 全体秒) で、最後は必ず全体に届く。"""
    t, _ = _local([_FakeSegment(0.0, 30.0, "あ")], duration=90.0)
    seen = []
    t.transcribe(__file__, on_progress=lambda done, total: seen.append((done, total)))
    assert seen[0] == (30.0, 90.0)
    assert seen[-1] == (90.0, 90.0)


def test_local_asr_bad_model_dir_is_reported_before_transcribing():
    """置き場の間違いは、音声を分割し終わる前に分かる。

    例外は align.py と同じ AlignUnavailable 系(LocalAsrUnavailable はその
    サブクラス)。呼び出し側は 1 つ catch すれば両方の経路を拾える。
    """
    with tempfile.TemporaryDirectory() as d:
        missing = Path(d) / "no-such-model"
        try:
            LocalTranscriber(model="small", model_dir=missing)
        except AlignUnavailable as e:
            assert "モデルフォルダが見つかりません" in str(e)
        else:
            raise AssertionError("存在しないモデルフォルダが通ってしまった")


# ======================================================================
# ライセンス表記(話者分離設計書 §9)
# ======================================================================

def test_credits_file_exists_and_is_reachable():
    """表記のファイルが実在し、読める。

    **同梱している TitaNet は CC-BY-4.0 で、表示が配布の条件。**
    ファイルが消えても他のテストは通ってしまうので、ここで固定する。
    """
    from src.config import credits_path, credits_text
    assert credits_path().exists(), f"表記がありません: {credits_path()}"
    assert len(credits_text()) > 200


def test_credits_names_the_cc_by_model():
    """CC-BY-4.0 のモデルと、その条件が書いてある。

    表記の義務があるのは TitaNet。名前とライセンス名の両方が要る
    (「第三者の部品を使っています」だけでは表示にならない)。
    """
    from src.config import credits_text
    t = credits_text()
    for want in ("TitaNet", "CC BY 4.0", "NVIDIA"):
        assert want in t, f"表記に {want} がありません"


def test_credits_names_the_other_bundled_models():
    """同梱している残りのモデルも挙げる。"""
    from src.config import credits_text
    t = credits_text()
    for want in ("pyannote", "MIT", "Whisper"):
        assert want in t, f"表記に {want} がありません"


def test_credits_survives_a_missing_file():
    """読めなくても落とさない。表記が出ないのは問題だが、それで転写が
    止まるのは筋が違う。代わりに要点だけでも出す。"""
    import src.config as cfg
    real = cfg.credits_path
    cfg.credits_path = lambda: Path("存在しない場所") / "CREDITS.txt"
    try:
        t = cfg.credits_text()
    finally:
        cfg.credits_path = real
    assert "TitaNet" in t and "CC BY 4.0" in t


def test_transcribed_segment_is_not_mistaken_for_an_added_one():
    """**転写の区間を「人が足した」と誤判定しない。**

    実データで起きた（42:15.9・2026-08-20）。足した発話と開始時刻が
    たまたま一致した転写の区間に ＋ 印が付き、［この区間を消す］まで
    押せる状態になっていた。「音声認識が出した区間には削除の入口を
    作らない」（CLAUDE.md）に反する。

    身元は (orig_start, 終わり, まとまり) の 3 つ組で見る。
    """
    proj = _base_for_insert()
    added = proj.add_utterance(103.0, 103.6, "はい", "g:D")
    # 開始だけ同じ、終わりもまとまりも違う転写の区間
    ghost = Segment(index=0, start=103.0, end=108.0,
                    text="機械が出した本文", cluster="g:B")
    proj.segments.append(ghost)
    proj.renumber()

    assert proj.is_added_utterance(added) is True
    assert proj.is_added_utterance(ghost) is False, "転写の区間を足したものと誤認"
    # 消すこと自体はできる（人が明示的に消すのは許す）。**記録の上で
    # 取り違えないこと**がここの要点——消したときの op が別になる。
    proj.remove_added_utterance(ghost.index)
    assert proj.edit_log[-1]["op"] == "remove_segment"


def test_two_added_at_the_same_time_both_reach_the_output():
    """**同じ時刻に 2 件足しても、両方 Word に出る。**

    身元を開始時刻だけで持っていたため、辞書で 2 件目が 1 件目を上書きし、
    上書きされたほうが差し込み先を失って**出力から丸ごと落ちていた**
    （実データ 42:12.7 で発生・2026-08-20）。
    """
    proj = _base_for_insert()
    base = proj.segments[1]
    a = _add_at(proj, base, 3, 103.0, "はい", "sp02")
    b = _add_at(proj, base, 7, 103.0, "ええ", "sp02")
    assert a.start == b.start, "この検査は同時刻でなければ意味がない"

    bodies = [text for _s, _sid, text, _c in _merge_runs(proj, True)]
    assert "はい" in bodies, bodies
    assert "ええ" in bodies, bodies
    # 割り込み位置と本文の対応が入れ替わっていないこと
    assert bodies.index("はい") < bodies.index("ええ"), bodies


def test_removing_one_of_two_leaves_the_other_removable():
    """**2 件のうち 1 件を消しても、残りは消せる。**

    鍵を集合で持っていたため、1 件消すと鍵ごと消え、残った 1 件が
    「人が足したもの」と認識されなくなって**二度と消せなくなっていた**。
    画面側が ValueError を握りつぶしていたので誰にも見えず、消し残しが
    作業ファイルに溜まった（実データ 42:12.7・01:04:57.2）。
    """
    proj = _base_for_insert()
    base = proj.segments[1]
    a = _add_at(proj, base, 3, 103.0, "はい", "sp02")
    b = _add_at(proj, base, 7, 103.0, "ええ", "sp02")

    proj.remove_added_utterance(b.index)
    left = [s for s in proj.segments if s.text == "はい"]
    assert len(left) == 1
    assert proj.is_added_utterance(left[0]) is True, "残った 1 件が消せない"
    proj.remove_added_utterance(left[0].index)
    assert not [s for s in proj.segments if s.text in ("はい", "ええ")]
    assert a.text == "はい"      # 消したのは指したものだけ


def _for_replace() -> Project:
    """実データの形をなぞる。**「資格」は本物と誤りが混ざっている。**"""
    proj = Project(audio_path="a.m4a", duration=300.0)
    proj.speakers = parse_roster("西村" + chr(10) + "山本")
    proj.segments = [
        Segment(index=0, start=10.0, end=20.0, cluster="g:A",
                text="前のPTAの会長さんは防災士の資格も取ってらっしゃって"),
        Segment(index=1, start=20.0, end=30.0, cluster="g:B",
                text="県は資格のことですからと言うだけで"),
        Segment(index=2, start=30.0, end=40.0, cluster="g:C",
                text="同層会ですね同層会があのー"),
    ]
    return proj


def test_find_text_returns_one_hit_per_occurrence():
    """**同じ区間に 2 回出たら 2 件返す。**

    実データ「同層会ですね同層会が」。区間単位にまとめると、片方だけ
    直したい場合に手が出せない。
    """
    proj = _for_replace()
    hits = proj.find_text("同層会")
    assert len(hits) == 2, [h.before + h.term + h.after for h in hits]
    assert [h.nth for h in hits] == [0, 1]
    assert all(h.key == segment_key(proj.segments[2]) for h in hits)


def test_find_text_carries_context():
    """**前後が無いと○×を付けられない。**「資格」は本物と誤りが混ざる。"""
    proj = _for_replace()
    hits = proj.find_text("資格")
    assert len(hits) == 2
    whole = [h.before + h.term + h.after for h in hits]
    assert "防災士の資格" in whole[0], whole
    assert "資格のことですから" in whole[1], whole


def test_replace_text_only_touches_the_chosen_places():
    """**選んだ箇所だけ直す。**全部置き換えると本物の「資格」が壊れる。"""
    proj = _for_replace()
    hits = proj.find_text("資格")
    wrong = [h for h in hits if "ことですから" in h.after]
    assert len(wrong) == 1
    n = proj.replace_text("資格", "私学", [h.target for h in wrong])
    assert n == 1
    assert proj.segments[0].text.endswith("防災士の資格も取ってらっしゃって")
    assert proj.segments[1].text == "県は私学のことですからと言うだけで"


def test_replace_text_can_pick_one_of_two_in_a_segment():
    """同じ区間の 2 回のうち、後ろだけ直す。"""
    proj = _for_replace()
    hits = proj.find_text("同層会")
    assert proj.replace_text("同層会", "同窓会", [hits[1].target]) == 1
    assert proj.segments[2].text == "同層会ですね同窓会があのー"


def test_replace_text_records_one_judgement():
    """**14 箇所直しても、人の判断は 1 回。**記録も 1 件（編集履歴 §1.1）。"""
    proj = _for_replace()
    hits = proj.find_text("同層会")
    proj.replace_text("同層会", "同窓会", [h.target for h in hits])
    recs = [r for r in proj.edit_log if r.get("op") == "replace_text_bulk"]
    assert len(recs) == 1, proj.edit_log
    assert recs[0]["before"] == "同層会" and recs[0]["after"] == "同窓会"
    assert recs[0]["count"] == 2 and recs[0]["segments"] == 1
    assert "replace_text_bulk" in LOG_LABELS, "検証要約に op 名のまま出てしまう"


def test_replace_text_never_marks_reviewed():
    """**機械が ✓ を立てる経路は作らない。**製品価値そのもの（CLAUDE.md）。"""
    proj = _for_replace()
    hits = proj.find_text("同層会")
    proj.replace_text("同層会", "同窓会", [h.target for h in hits])
    seg = proj.segments[2]
    assert seg.reviewed is False, "聴いていないのに確定になった"
    assert seg.text_edited is True, "再実行で機械の出力に戻される"


def test_replace_text_does_nothing_without_a_target():
    """空・同じ語・対象なしでは、記録も残さない（編集履歴 §1.5）。"""
    proj = _for_replace()
    assert proj.replace_text("", "私学", []) == 0
    assert proj.replace_text("資格", "資格", []) == 0
    assert proj.replace_text("資格", "私学", []) == 0
    assert proj.find_text("") == []
    assert not [r for r in proj.edit_log if r.get("op") == "replace_text_bulk"]


def _for_leave() -> Project:
    """途中退席の形。32:17 に本人が「中座します」と言って帰る。"""
    proj = Project(audio_path="a.m4a", duration=4000.0)
    proj.speakers = parse_roster("吉沢忠一" + chr(10) + "西村香介")
    yo, ni = proj.speakers[0].id, proj.speakers[1].id
    proj.segments = [
        Segment(index=0, start=100.0, end=110.0, text="冒頭のあいさつ",
                cluster="g:A", speaker_id=yo, reviewed=True),
        Segment(index=1, start=1937.0, end=1945.0, text="ちょっと中座させて",
                cluster="g:A", speaker_id=yo, reviewed=True),
        Segment(index=2, start=2000.0, end=2010.0, text="その後の発言",
                cluster="g:A", speaker_id=yo, reviewed=False),
        Segment(index=3, start=3874.0, end=3880.0, text="申し訳ありませんでした。",
                cluster="g:A", speaker_id=yo, reviewed=True),
    ]
    return proj


def test_replace_speaker_only_touches_the_chosen_segments():
    """**退席の直前は本人の発言。**一律に変えると壊れる。"""
    proj = _for_leave()
    yo, ni = proj.speakers[0].id, proj.speakers[1].id
    after = [s for s in proj.segments if s.start > 1950]
    n = proj.replace_speaker(yo, ni, [segment_key(s) for s in after])
    assert n == 2
    assert proj.segments[0].speaker_id == yo   # 冒頭
    assert proj.segments[1].speaker_id == yo   # 「中座させて」は本人
    assert proj.segments[2].speaker_id == ni
    assert proj.segments[3].speaker_id == ni


def test_replace_speaker_never_keeps_the_heard_mark():
    """**✓ のまま残さない。**確かめたのは「居なかった」ことであって声ではない。

    ✓ を残すと「1 区間ずつ聴いて確かめた」という意味になり、`reviewed` の
    意味論が壊れる（CLAUDE.md）。
    """
    proj = _for_leave()
    yo, ni = proj.speakers[0].id, proj.speakers[1].id
    heard = proj.segments[3]
    assert heard.reviewed is True              # この検査の前提
    proj.replace_speaker(yo, ni, [segment_key(heard)])
    assert heard.speaker_id == ni
    assert heard.reviewed is False, "機械が ✓ を立てる経路ができている"


def test_replace_speaker_records_one_judgement_with_before():
    """記録は 1 件。**`before` も残す**（編集履歴 §1.1・§1.2）。"""
    proj = _for_leave()
    yo, ni = proj.speakers[0].id, proj.speakers[1].id
    after = [s for s in proj.segments if s.start > 1950]
    proj.replace_speaker(yo, ni, [segment_key(s) for s in after])
    recs = [r for r in proj.edit_log if r.get("op") == "replace_speaker_bulk"]
    assert len(recs) == 1, proj.edit_log
    assert recs[0]["before"] == yo and recs[0]["after"] == ni
    assert recs[0]["count"] == 2
    # ✓ から △ に落ちたものが分かる（あとで聴き直す手がかり）
    assert recs[0]["was_reviewed"] == [list(segment_key(proj.segments[3]))]
    assert "replace_speaker_bulk" in LOG_LABELS


def test_replace_speaker_ignores_segments_of_other_people():
    """指定に他人の区間が混ざっても巻き込まない。"""
    proj = _for_leave()
    yo, ni = proj.speakers[0].id, proj.speakers[1].id
    proj.segments[2].speaker_id = ni           # もう西村になっている
    n = proj.replace_speaker(yo, ni, [segment_key(s) for s in proj.segments])
    assert n == 3                               # 吉沢のままの 3 区間だけ
    assert proj.replace_speaker(yo, yo, [segment_key(proj.segments[0])]) == 0
    assert proj.replace_speaker(yo, ni, []) == 0


def test_device_choice_pairs_with_the_model():
    """**装置とモデルは組で決まる。**CPU で large-v3 は実用外(実時間比が数倍)。

    CPU の既定は small。GPU があれば large-v3 を**選べる**ようにしてある。

    **かつてここには「速くなって正確」と書いてあったが、撤回した。**
    根拠にしていた「24 分 → 15.5 分」は **small/CPU と large-v3/GPU を
    比べた値**で、機械が違う。同じ GPU で測ると large-v3 のほうが遅い
    （2026-08-31・7 回: 合計 13〜17 分 → 26〜32 分）。併記されていた
    「誤字」の値も撤回した（定義が復元できない。設計書 §9.5.1）。
    **large-v3 を勧める根拠は固有名詞だけで、速さではない。**

    **GPU 側の既定は `test_default_model_only_picks_what_is_there` が見る。**
    ここで `default_model(GPU_DEVICE)` を検査していたが、あれは large-v3 が
    ディスクにあるかどうかで変わる(align.py:193)ので、手元では通り CI では
    落ちる——機械の状態が漏れていただけだった。向こうは一時フォルダで
    両方の分岐を作るので、どの機械でも同じ結果になる。
    """
    from src import align
    assert align.default_model("cpu") == align.DEFAULT_MODEL == "small"
    assert "large-v3" in align.AVAILABLE_MODELS
    # GPU を使わない指定は、GPU のある機械でも CPU の組を返す
    assert align.pick_device(prefer_gpu=False) == (align.DEVICE,
                                                   align.COMPUTE_TYPE)


def test_gpu_uses_int8_not_plain_float16():
    """**float16 ではなく int8_float16。**同じ品質で 3.5 倍速い（実測）。

    large-v3 で 67 分の音声が 53.7 分 → 15.5 分（2026-08-20）。
    脱落 11.6% → 11.0%、落とした発言の回収は同じ 17/34。
    **併記されていた「誤字」の値は撤回した**（設計書 §9.5.1）。

    **この 15.5 分を「終わるまでの時間」と読まないこと。**転写だけの値で、
    話者分離を含まない。2026-08-31 に測り直すと転写だけで 17〜20 分あり、
    話者分離まで含めた合計は 26〜32 分だった。
    """
    from src import align
    assert align.GPU_COMPUTE_TYPE == "int8_float16"



def test_model_source_does_not_leak_the_install_path():
    """**ログに絶対パスを出さない。**

    `installer.iss` の `AppId` は互換のため変えられない（変えると 1 台に
    2 つ入る）ので、**旧版から更新した人のインストール先は旧名のまま**残る。
    v2.0.6 の `DefaultDirName` は `{autopf}\\GeminiTranscriber` だった。

    その結果、ログに

        ローカル転写は端末内で完結します。録音も本文も外へ出しません。
        ...
        モデルの場所: C:\\...\\GeminiTranscriber\\_internal\\models\\asr\\small

    と出ていた。**主張の数行下に Gemini と出る。**

    **アプリの名前も出さない。**新名（WhosaidEditor）であっても、パスを
    出せば同じ形の問題を繰り返す余地が残る。「設定フォルダ」で足りる。
    """
    import os
    import tempfile
    from pathlib import Path
    from src import align

    NG = ("GeminiTranscriber", "WhosaidEditor", "AppData", "Program Files",
          "_internal", "_MEIPASS")

    def clean(s, why):
        for bad in NG:
            assert bad not in s, f"{why}: 「{bad}」が出ている → {s}"
        # ドライブレターから始まる絶対パスも出さない
        assert ":\\" not in s, f"{why}: 絶対パスが出ている → {s}"

    # --- 同梱 / 取得済み / 見つからない、いずれもパスを出さない -------
    real_find = align.find_bundled_model
    real_root = align.downloaded_model_root
    try:
        with tempfile.TemporaryDirectory() as d:
            dl = Path(d) / "dl"
            bundled = Path(d) / "app" / "_internal" / "models" / "asr" / "small"
            (dl / "small").mkdir(parents=True)
            bundled.mkdir(parents=True)
            align.downloaded_model_root = lambda: dl

            align.find_bundled_model = lambda m: dl / m
            got = align.model_source("small")
            clean(got, "あとから取得したもの")
            assert "取得" in got and "設定フォルダ" in got, got

            align.find_bundled_model = lambda m: bundled
            got = align.model_source("small")
            clean(got, "同梱")
            assert "同梱" in got, got

            align.find_bundled_model = lambda m: None
            got = align.model_source("small")
            clean(got, "見つからない")

            # --- 手動指定だけはパスを出す（利用者が入れた値）---------
            manual = Path(d) / "私が置いた場所"
            got = align.model_source("small", manual)
            assert str(manual) in got, "手動指定はパスを返して見せる"
            assert "手動" in got, got
    finally:
        align.find_bundled_model = real_find
        align.downloaded_model_root = real_root
        os.environ.pop("WHOSAID_ASR_MODELS", None)


def test_local_asr_logs_the_source_not_the_path():
    """**ログ本文にも絶対パスが出ないこと。**関数だけ直しても、
    呼び出し側が `self.target` を出していたら意味が無い。
    """
    from pathlib import Path
    src = (Path(__file__).resolve().parent.parent / "src" / "local_asr.py"
           ).read_text(encoding="utf-8")
    assert 'モデルの場所: {self.target}' not in src, \
        "絶対パスを出す書き方に戻っている"
    assert "model_source(" in src, "種別を出す関数を呼んでいない"
    # 診断の行そのものは残っていること（2026-08-22 に意図して足された）
    assert "モデルの場所" in src, "**行ごと消してはいけない。**置き換えるだけ"

def test_cache_key_separates_device_and_prompt():
    """**装置の精度とプロンプトの版が違えば、別のキャッシュになる。**

    どちらも転写結果を変える。キーに入れないと古い転写を使い回す
    （CLAUDE.md「どれかを欠くと古い転写の使い回し事故になる」）。
    プロンプトを変えたときは句読点の密度が 1 万字あたり 452 → 1055 に
    動いた（実測）。
    """
    from src import local_asr
    def key(**kw):
        return local_asr.chunk_cache_path(
            "/tmp", "chunk_0000", fingerprint="abc", chunk_seconds=420,
            model="large-v3", **kw).name

    base = key()
    assert base != key(compute_type="int8_float16"), "精度で分かれていない"
    assert base != key(prompt_ver=1), "プロンプトの版で分かれていない"
    assert key(compute_type="int8") != key(compute_type="int8_float16")
    assert key(prompt_ver=1) != key(prompt_ver=2)
    # モデルで分かれるのは従来どおり（壊していないこと）
    assert local_asr.chunk_cache_path(
        "/tmp", "c", fingerprint="a", chunk_seconds=420, model="small").name         != local_asr.chunk_cache_path(
            "/tmp", "c", fingerprint="a", chunk_seconds=420,
            model="large-v3").name
    # 指紋が無ければ従来どおりキャッシュしない
    assert local_asr.chunk_cache_path(
        "/tmp", "c", fingerprint="", chunk_seconds=420, model="small") is None


def test_style_prompt_is_versioned():
    """プロンプトを変えたら版を上げる、という約束を検査で縛る。

    版を上げ忘れると、古い転写を使い回して「直したのに変わらない」になる。
    """
    from src import local_asr
    assert local_asr.PROMPT_VER >= 1
    assert "、" in local_asr.STYLE_PROMPT, "句読点のための例文でなくなっている"
    # 用語を入れない（句読点の効果と用語の効果を混ぜないため・実測の経緯）
    assert "耐震" not in local_asr.STYLE_PROMPT


def test_carry_speakers_never_keeps_the_heard_mark():
    """**別の転写から写した割当は、必ず △。**

    モデルを変えると区間の切れ目が変わる（実データで 725 → 772）。
    切れ目が違う以上「その区間を聴いて確定した」とは言えないので、
    元が ✓ でも △ に落とす（CLAUDE.md の ✓/△ の意味論）。
    """
    proj = _for_leave()
    yo = proj.speakers[0].id
    for s in proj.segments:
        s.speaker_id, s.reviewed = None, False
    n = proj.carry_speakers(
        [(segment_key(proj.segments[0]), yo),
         (segment_key(proj.segments[2]), yo)], source="前.json")
    assert n == 2
    assert proj.segments[0].speaker_id == yo
    assert proj.segments[0].reviewed is False, "機械が ✓ を立てている"
    assert proj.segments[1].speaker_id is None, "指定していない区間が埋まった"


def test_carry_speakers_records_one_judgement():
    """記録は 1 件。どこから写したかを残す（編集履歴 §1.1・§1.2）。"""
    proj = _for_leave()
    yo = proj.speakers[0].id
    for s in proj.segments:
        s.speaker_id = None
    proj.carry_speakers([(segment_key(s), yo) for s in proj.segments],
                        source="前.json")
    recs = [r for r in proj.edit_log if r.get("op") == "carry_speakers"]
    assert len(recs) == 1, proj.edit_log
    assert recs[0]["count"] == 4 and recs[0]["source"] == "前.json"
    assert "carry_speakers" in LOG_LABELS


def test_carry_speakers_does_nothing_without_targets():
    """空・話者なし・すでに同じ、では記録も残さない（編集履歴 §1.5）。"""
    proj = _for_leave()
    yo = proj.speakers[0].id
    assert proj.carry_speakers([], source="x") == 0
    assert proj.carry_speakers(
        [(segment_key(proj.segments[0]), None)], source="x") == 0
    # すでに同じ人なら変化なし
    assert proj.carry_speakers(
        [(segment_key(proj.segments[0]), yo)], source="x") == 0
    assert not [r for r in proj.edit_log if r.get("op") == "carry_speakers"]


def _utts(*texts):
    from src.segments import Utterance
    return [Utterance(rel_start=i * 2.0, rel_end=i * 2.0 + 1.5, text=t)
            for i, t in enumerate(texts)]


def test_loop_detection_counts_across_segments():
    """**区間ごとに数える。**本文を 1 本に繋いで正規表現で探すと拾えない。

    実データで large-v3 が同じ文を 20 区間繰り返し、4 分ぶんの会話が
    丸ごと消えた（2026-08-20）。1 本に繋いだ文字列で調べていたため、
    最初の測定では 0 件と出てしまった。
    """
    from src.local_asr import max_repeat_run, LOOP_RUN_THRESHOLD
    assert max_repeat_run(_utts("あ", "い", "う")) == 1
    assert max_repeat_run(_utts("あ", "同じ", "同じ", "同じ", "い")) == 3
    assert max_repeat_run(_utts(*(["同じ"] * 20))) == 20
    assert max_repeat_run([]) == 0
    # 空の本文は数えない（繋がっているように見えてしまう）
    assert max_repeat_run(_utts("同じ", "", "同じ")) == 1
    # **数え方そのもの**の検査。暴走かどうかの判断は find_runaway が持つ
    assert LOOP_RUN_THRESHOLD >= 3


def _run_of(text, durations):
    """同じ本文を、指定した長さで並べた区間列。"""
    from src.segments import Utterance
    out, t = [], 0.0
    for d in durations:
        out.append(Utterance(rel_start=t, rel_end=t + d, text=text))
        t += d + 0.3
    return out


def test_runaway_is_judged_by_count_and_speaking_rate():
    """**「同じ本文が続く」だけでは暴走と決めない。**

    当初は「3 区間以上」としたが、15 分チャンクで誤発動だらけになった
    （実機で 5 チャンク中 4 つ・処理時間がほぼ倍。2026-08-21）。1 チャンクに
    260 区間あると、本物の相づちが 3〜4 回続くのは普通。

    見本は**実データそのもの**（同じ音声から人が見分けた 8 件）。
    """
    from src.local_asr import find_runaway
    # --- 本物（発動してはいけない）---
    # 「吉澤さん」4 回。人が実際にそう言っている
    assert find_runaway(_run_of("吉澤さん", [1.0, 1.0, 0.7, 0.5])) is None
    # 「もちろん。」3 回・中央値 14 字/秒
    assert find_runaway(_run_of("もちろん。", [0.36, 0.22, 0.5])) is None
    # 「はい。」4 回。1 つだけ潰れているが本物
    assert find_runaway(_run_of("はい。", [0.36, 0.36, 0.36, 0.05])) is None

    # --- 暴走（発動してほしい）---
    # 「ありがとうございます。」8 回。0.1 秒に 11 字は入らない
    got = find_runaway(_run_of("ありがとうございます。",
                               [0.4, 0.6, 0.2, 0.6, 0.5, 0.2, 0.1, 0.6]))
    assert got is not None and got[0] == 8
    # 「nick」8 回。話速は正常だが回数で拾う
    assert find_runaway(_run_of("nick", [1.0] * 8)) is not None
    # 「はい。」6 回・中央値 44 字/秒
    assert find_runaway(_run_of("はい。", [0.05] * 6)) is not None

    # 2 回では見ない（相づちの応酬を毎回起こし直すことになる）
    assert find_runaway(_run_of("はい。", [0.05, 0.05])) is None



def test_runaway_log_says_which_one_was_kept():
    """**3 つの分岐すべてで、字数と「どちらを残したか」を出す。**

    「直りませんでした」の分岐だけ字数を出さず、元を残したことも
    言っていなかった（2026-08-31 に実測で気づいた）。読み手は
    **本物の相づちか本文が消えたか**を判断できず、しかも
    **どちらの結果が残ったのかも分からなかった。**

    あわせて「どの分岐でも元を返す」ことも縛る。起こし直しの結果で
    上書きされることは無い——実データでも確かめた（最終チャンクの字数は、
    暴走が出た回が最小ではなかった）。
    """
    import tempfile
    from src.local_asr import ChunkResult, LocalTranscriber

    def utts(text, n, dur=0.36):
        return _run_of(text, [dur] * n)

    def run(retry_result):
        """元は暴走あり。起こし直しの結果を差し替えて 3 分岐を作る。"""
        t = object.__new__(LocalTranscriber)
        t.condition_on_previous_text = True
        orig = ChunkResult(utterances=utts("ありがとうございます。", 5))
        calls = {"n": 0}

        def fake_once(*a, **k):
            calls["n"] += 1
            return orig if calls["n"] == 1 else retry_result

        t._once = fake_once
        logs = []
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
            p = Path(f.name)
        try:
            got = LocalTranscriber.transcribe(t, p, on_log=logs.append)
        finally:
            p.unlink(missing_ok=True)
        return got, orig, "\n".join(logs)

    def varied(n):
        """**本文を散らす。**同じ文を並べると、それ自体が暴走判定に当たる
        （最初この検査を書いたとき、分岐1 のダミーが分岐3 に落ちた）。"""
        from src.segments import Utterance
        out, s = [], 0.0
        for i in range(n):
            out.append(Utterance(rel_start=s, rel_end=s + 1.0,
                                 text=f"ちゃんと起きた {i} 番目の本文です。"))
            s += 1.3
        return out

    # --- 分岐1: 直った（暴走が消え、本文も減っていない）---
    longer = ChunkResult(utterances=varied(9))
    got, orig, log = run(longer)
    assert got is longer, "直ったら起こし直しを採る"
    assert "直りました" in log
    assert "字 → " in log, "分岐1 に字数が無い"

    # --- 分岐2: 暴走は消えたが本文が減った → 元を残す ---
    shorter = ChunkResult(utterances=_run_of("短い。", [1.0] * 2))
    got, orig, log = run(shorter)
    assert got is orig, "本文が減るなら元を残す"
    assert "元のままにします" in log
    assert "字 → " in log, "分岐2 に字数が無い"

    # --- 分岐3: 起こし直しても暴走が残った → 元を残す ---
    still = ChunkResult(utterances=utts("ありがとうございます。", 4))
    got, orig, log = run(still)
    assert got is orig, "**直らなくても元を返す。**起こし直しで上書きしない"
    assert "直りませんでした" in log
    # ここが今回の本題。**他の 2 分岐と揃っていること**
    assert "字 → " in log, "分岐3 に字数が無い（他の 2 分岐にはある）"
    assert "元のままにします" in log, "分岐3 がどちらを残したか言っていない"
    # 断定していないこと。本物の相づちの可能性を落とさない
    assert "本物の相づち" in log
    assert "失われている可能性" in log

def test_speaking_rate_line_sits_between_the_real_cases():
    """線の位置を実データで縛る。**本物 13〜14、暴走 23〜44 の間。**"""
    from src.local_asr import MAX_CHARS_PER_SECOND, LOOP_RUN_THRESHOLD
    assert 14 < MAX_CHARS_PER_SECOND < 23, "本物と暴走の間から外れている"
    # 本物の繰り返しは実データで最大 4 回、暴走は 6 回以上
    assert 4 < LOOP_RUN_THRESHOLD <= 6


def test_the_real_runaway_greeting_is_caught():
    """**実データそのもの。**01:05:48「失礼いたします。」3 区間（暴走）。

    2026-08-22 に利用者が音声を聴いて確かめた。実際の音声は
    「失礼いたします。これで失礼いたします」の 2 発言だけで、3 区間は
    同じ音声を 3 回書いたもの。長さは実際の値（0.36 / 0.44 / 0.48 秒・8 字）
    で、中央値 18.2 字/秒。

    **一度これを「本物」と誤解して線を 25 へ上げかけた。**聴いた人の判定が
    入って初めて分かったので、耳で確かめた事実だけを根拠にすること。
    """
    from src.local_asr import find_runaway
    got = find_runaway(_run_of("失礼いたします。", [0.36, 0.44, 0.48]))
    assert got is not None and got[0] == 3


def test_retry_replaces_the_run_and_says_so():
    """暴走したら引き継ぎを切って起こし直し、**直ったことを利用者に伝える**。"""
    from src import local_asr
    tr = local_asr.LocalTranscriber.__new__(local_asr.LocalTranscriber)
    tr.condition_on_previous_text = True
    calls, logs = [], []
    # **起こし直した側は本文が増える。**実際に起きた形（1,288 字 → 1,723 字）。
    # 減るなら元を採るべきで、それは別の検査で見る。
    good = local_asr.ChunkResult(utterances=_utts(
        "本日はお集まりいただきありがとうございます",
        "それでは議事に入ります", "資料をご覧ください", "以上です"))
    bad = local_asr.ChunkResult(utterances=_utts(*(["同じ"] * 8)))

    def fake_once(path, condition, **kw):
        calls.append(condition)
        return bad if condition else good

    tr._once = fake_once
    got = local_asr.LocalTranscriber.transcribe(
        tr, __file__, on_log=logs.append)
    assert calls == [True, False], "引き継ぎを切って起こし直していない"
    assert got is good, "直った結果に差し替えていない"
    assert any("暴走" in x for x in logs), logs
    assert any("直りました" in x for x in logs), logs


def test_retry_is_rejected_when_the_text_shrinks():
    """**繰り返しが減っただけでは採らない。**本文が減っていないことも見る。

    引き継ぎを切ると、壊れていない箇所では脱落が増える（実測 2026-08-21:
    逐語正解 4 帯で脱落 10.4% → 16.6%、落とした発言の回収 18 → 10 件）。
    閾値ちょうど（3 区間）で発動した chunk_0001 は、起こし直すと繰り返しは
    3 → 2 に減ったが本文が 1,758 字 → 1,669 字に**減った**。
    本当に壊れていた chunk_0007 は 1,288 字 → 1,723 字と**増えている**。
    """
    from src import local_asr
    tr = local_asr.LocalTranscriber.__new__(local_asr.LocalTranscriber)
    tr.condition_on_previous_text = True
    logs = []
    # **暴走と判定される形**（6 回以上）だが、起こし直すと本文が減る
    bad = local_asr.ChunkResult(
        utterances=_utts(*(["はい。"] * 6), "長い発言をしています"))
    thin = local_asr.ChunkResult(utterances=_utts("はい。", "はい。"))
    tr._once = lambda path, condition, **kw: bad if condition else thin
    got = local_asr.LocalTranscriber.transcribe(
        tr, __file__, on_log=logs.append)
    assert got is bad, "本文が減るほうを採ってしまっている"
    assert any("元のままにします" in x for x in logs), logs
    # 本文が増えるなら採る（本当に壊れていた場合）
    fat = local_asr.ChunkResult(utterances=_utts(
        "そういう経過もありましたね", "確かに経営が悪くなってきているんでしょうね",
        "だから全然続ける気がないのかなと思っちゃうんですよ"))
    tr._once = lambda path, condition, **kw: bad if condition else fat
    assert local_asr.LocalTranscriber.transcribe(tr, __file__) is fat


def test_retry_that_fails_is_reported_not_hidden():
    """**直らなかったことを黙らない。**本文が消える事故なので、
    気づかないまま納品されるのがいちばん悪い。"""
    from src import local_asr
    tr = local_asr.LocalTranscriber.__new__(local_asr.LocalTranscriber)
    tr.condition_on_previous_text = True
    logs = []
    bad = local_asr.ChunkResult(utterances=_utts(*(["同じ"] * 8)))
    tr._once = lambda path, condition, **kw: bad
    got = local_asr.LocalTranscriber.transcribe(
        tr, __file__, on_log=logs.append)
    assert got is bad
    assert any("直りませんでした" in x for x in logs), logs
    assert any("聴いて確かめ" in x for x in logs), logs


def test_no_retry_when_nothing_looks_wrong():
    """暴走していなければ、起こし直さない（時間が倍になる）。"""
    from src import local_asr
    tr = local_asr.LocalTranscriber.__new__(local_asr.LocalTranscriber)
    tr.condition_on_previous_text = True
    calls = []
    ok = local_asr.ChunkResult(utterances=_utts("あ", "い", "う"))

    def fake_once(path, condition, **kw):
        calls.append(condition)
        return ok

    tr._once = fake_once
    assert local_asr.LocalTranscriber.transcribe(tr, __file__) is ok
    assert calls == [True], "起こし直してしまっている"


def test_config_is_carried_over_from_the_old_name():
    """**改名で利用者の設定を失わせない。**

    `%APPDATA%` の置き場が変わるので、そのままだと旧版の利用者が
    API キーも名簿も失う（2026-08-21 の改名: GeminiTranscriber →
    WhosaidEditor）。
    """
    import os
    import tempfile
    from unittest import mock
    import src.config as cfg

    with tempfile.TemporaryDirectory() as d:
        base = Path(d)
        old = base / cfg.LEGACY_APP_NAMES[0]
        old.mkdir()
        (old / "config.json").write_text(
            '{"api_key": "abc", "local_model": "large-v3"}', encoding="utf-8")
        with mock.patch.dict(os.environ, {"APPDATA": str(base)}):
            got = cfg.load_config()
        assert got.get("api_key") == "abc", "引き継げていない"
        assert got.get("local_model") == "large-v3"
        new_dir = base / cfg.APP_NAME
        assert (new_dir / "config.json").exists()
        assert (new_dir / "migrated-from.txt").exists(), "経緯が残っていない"
        # **旧フォルダは消さない**（旧版に戻したい人がいるかもしれない）
        assert (old / "config.json").exists()


def test_config_migration_never_overwrites_the_new_one():
    """新しい側に設定があれば、旧名から上書きしない。"""
    import os
    import tempfile
    from unittest import mock
    import src.config as cfg

    with tempfile.TemporaryDirectory() as d:
        base = Path(d)
        old = base / cfg.LEGACY_APP_NAMES[0]
        old.mkdir()
        (old / "config.json").write_text('{"api_key": "古い"}', encoding="utf-8")
        new_dir = base / cfg.APP_NAME
        new_dir.mkdir()
        (new_dir / "config.json").write_text('{"api_key": "新しい"}',
                                             encoding="utf-8")
        with mock.patch.dict(os.environ, {"APPDATA": str(base)}):
            got = cfg.load_config()
        assert got.get("api_key") == "新しい", "旧名の設定で上書きされた"


def test_app_name_no_longer_claims_gemini():
    """**名前が事実と食い違わないこと。**既定はローカルで、録音は外に出ない。

    「Gemini」は他社の商標でもある。旧名は引き継ぎのためだけに残す。
    """
    import src.config as cfg
    assert cfg.APP_NAME == "WhosaidEditor"
    assert "Gemini" not in cfg.APP_NAME
    assert "GeminiTranscriber" in cfg.LEGACY_APP_NAMES, "引き継ぎ先が消えている"


def test_engine_descriptions_match_reality():
    """**画面の説明が、実際の動きと矛盾していないこと。**

    話者分離を入れたあとも「声のまとまりは作られず」と書いたままで、
    同じ画面のチェックボックス（既定 ON）と矛盾していた（2026-08-21 に発覚）。
    無料公開では「ローカルでは話者が分からない」と誤解され、いちばんの売りが
    伝わらない。**機能を足したら説明も直す**、を検査で縛る。
    """
    from src import gui, pipeline
    spec = pipeline.EngineSpec(mode=pipeline.ENGINE_LOCAL)
    assert spec.wants_diarize is True, "この検査の前提（既定で話者分離が入る）"
    for bad in ("作られず", "作られない"):
        assert bad not in gui.ENGINE_LOCAL_DESC, (
            f"話者分離が既定で入るのに「{bad}」と書いてある")
    assert "話者分離" in gui.ENGINE_LOCAL_DESC, "作れることが書かれていない"
    # 逐語モードはローカルでは使えない（こちらは今も本当）
    assert "逐語" in gui.ENGINE_LOCAL_DESC

    # **説明は 1 か所ではない。**この検査は ENGINE_LOCAL_DESC しか見ておらず、
    # 「話者の決め方」側の注記(lbl_diar_note)は
    # 「※ ローカルでは声のまとまりを作らないため、全区間が未判別になります。」
    # のまま残っていた。**同じ画面の上と下で逆のことを言っていた**のを、
    # 配布版の画面を見て見つけた（2026-08-29）。検査の網が狭かった。
    import inspect
    src = inspect.getsource(gui.App._update_local_diar_note)
    on_branch = src.split("else:")[0]      # 話者分離が入っている側の文面
    for bad in ("作らない", "作られない", "作られず"):
        assert bad not in on_branch, (
            f"話者分離が入っているのに「{bad}」と書いてある: {on_branch}")
    assert "話者分離" in on_branch, "どこで作るのかが書かれていない"


def test_licence_names_a_holder():
    """**公開する前に著作権者を埋める。**プレースホルダのままにしない。"""
    text = (Path(__file__).resolve().parent.parent / "LICENSE").read_text(
        encoding="utf-8")
    assert "MIT License" in text
    assert "[Your Name]" not in text, "LICENSE がプレースホルダのまま"
    assert "Copyright (c)" in text


def test_bundled_model_is_used_before_the_network():
    """**同梱したモデルがあれば、それを使う。**

    faster-whisper はモデル名を渡すと Hugging Face へ取りに行く。同梱を
    見に行かないと、新規インストール直後に通信が要り、画面の
    「ネットワークを遮断したままでも動きます」が嘘になる（設計書 §9）。
    """
    import os
    import tempfile
    from unittest import mock
    from src import align

    with tempfile.TemporaryDirectory() as d:
        root = Path(d)
        (root / "small").mkdir()
        (root / "small" / "model.bin").write_bytes(b"x")
        # **探し場所ごと差し替える。**環境変数だけだと、この機械に実際に
        # 取得済みのモデルがあると拾ってしまい、検査が環境に依存する
        # （実際にそうなった・2026-08-21）。
        with mock.patch.object(align, "asr_model_dirs", lambda m: [root / m]):
            assert align.find_bundled_model("small") == root / "small"
            assert align.resolve_model("small") == str(root / "small")
            # 同梱していないものは名前のまま（faster-whisper が取りに行く）
            assert align.resolve_model("large-v3") == "large-v3"


def test_model_dir_still_wins_over_the_bundled_one():
    """手動で置いたフォルダが最優先、という約束を壊さない（設計書 §9）。"""
    import os
    import tempfile
    from unittest import mock
    from src import align

    with tempfile.TemporaryDirectory() as d:
        root = Path(d)
        (root / "small").mkdir()
        (root / "small" / "model.bin").write_bytes(b"x")
        hand = root / "手で置いた"
        hand.mkdir()
        (hand / "model.bin").write_bytes(b"y")
        with mock.patch.object(align, "asr_model_dirs", lambda m: [root / m]):
            assert align.resolve_model("small", model_dir=hand) == str(hand)


def test_bundled_model_needs_the_weights():
    """**フォルダだけあって中身が無いものを掴まない。**

    探す場所は複数ある（環境変数 → 同梱 → 手元）。空のフォルダで
    止まってしまうと、本物が置いてあっても見つけられない。
    """
    import os
    import tempfile
    from unittest import mock
    from src import align

    with tempfile.TemporaryDirectory() as d:
        root = Path(d)
        (root / "small").mkdir()          # model.bin が無い
        with mock.patch.object(
                align, "asr_model_dirs",
                lambda m: [root / m, Path("models/asr") / m]):
            got = align.find_bundled_model("small")
            assert got != root / "small", "中身の無いフォルダを掴んでいる"
        # どこにも無ければ None（名前のまま返り、faster-whisper が取りに行く）
        with mock.patch.object(align, "asr_model_dirs",
                               lambda m: [root / "small"]):
            assert align.find_bundled_model("small") is None
            assert align.resolve_model("small") == "small"


def test_default_model_only_picks_what_is_there():
    """**無いモデルを既定にしない。**初回の転写でいきなり 3GB は落とさない。

    large-v3 は GPU があれば速くて正確だが（§9.5.1）、3GB あり同梱して
    いない。取得済みのときだけ既定にする。勧めるのは画面の仕事。
    """
    import os
    import tempfile
    from unittest import mock
    from src import align

    with tempfile.TemporaryDirectory() as d:
        root = Path(d)
        with mock.patch.dict(os.environ, {"WHOSAID_ASR_MODELS": str(root)}),              mock.patch.object(align, "asr_model_dirs",
                               lambda m: [root / m]):
            # まだ無い → small が既定で、large-v3 を勧める
            assert align.default_model(align.GPU_DEVICE) == align.DEFAULT_MODEL
            with mock.patch.object(align, "pick_device",
                                   lambda *a, **k: (align.GPU_DEVICE, "x")):
                assert align.suggest_gpu_model() == align.GPU_DEFAULT_MODEL
            # 取得したら既定が上がり、もう勧めない
            (root / align.GPU_DEFAULT_MODEL).mkdir()
            (root / align.GPU_DEFAULT_MODEL / "model.bin").write_bytes(b"x")
            assert align.default_model(align.GPU_DEVICE) == align.GPU_DEFAULT_MODEL
            with mock.patch.object(align, "pick_device",
                                   lambda *a, **k: (align.GPU_DEVICE, "x")):
                assert align.suggest_gpu_model() is None
        # CPU の機械には勧めない
        with mock.patch.object(align, "pick_device",
                               lambda *a, **k: (align.DEVICE, align.COMPUTE_TYPE)):
            assert align.suggest_gpu_model() is None


def test_fetch_leaves_nothing_behind_when_it_fails():
    """**途中で切れた取得を残さない。**残すと「あるのに動かない」になる。"""
    import tempfile
    from unittest import mock
    from src import asr_fetch

    with tempfile.TemporaryDirectory() as d:
        root = Path(d)
        with mock.patch("huggingface_hub.snapshot_download",
                        side_effect=OSError("通信できません")):
            try:
                asr_fetch.fetch("small", root)
            except asr_fetch.AsrFetchError as e:
                assert "通信が遮断された環境" in str(e), "対処法が書かれていない"
            else:
                raise AssertionError("失敗を知らせていない")
        assert not (root / "small").exists(), "壊れたフォルダが残っている"
        assert not (root / ".small.part").exists(), "作業用が残っている"


def test_fetch_rejects_an_incomplete_download():
    """核のファイルが欠けたものを置かない（利用者の端末で初めて分かるのは遅い）。"""
    import tempfile
    from unittest import mock
    from src import asr_fetch

    with tempfile.TemporaryDirectory() as d:
        root = Path(d)

        def fake(**kw):
            out = Path(kw["local_dir"])
            out.mkdir(parents=True, exist_ok=True)
            (out / "config.json").write_text("{}")   # model.bin が無い
            return str(out)

        with mock.patch("huggingface_hub.snapshot_download", side_effect=fake):
            try:
                asr_fetch.fetch("small", root)
            except asr_fetch.AsrFetchError as e:
                assert "model.bin" in str(e)
            else:
                raise AssertionError("欠けたものを受け入れてしまった")
        assert not (root / "small").exists()


def test_fetch_sizes_are_stated_before_downloading():
    """**大きさを先に伝える。**断りも入れずに GB 級を落とさない。"""
    from src import asr_fetch
    for m in ("small", "medium", "large-v3"):
        assert asr_fetch.SIZES_MB.get(m, 0) > 0, f"{m} の目安が無い"
    assert asr_fetch.SIZES_MB["large-v3"] > 2000, "3GB 級だと伝わらない"


def test_cuda_parts_are_pinned_and_verified():
    """**版を留め、中身を確かめてから使う。**

    `nvidia-cublas-cu12` は proprietary で、配布元の資産が差し替わることが
    ありうる。黙って別物を動かすのではなく止まってほしい。
    """
    from src import cuda_fetch
    assert cuda_fetch.CUBLAS_VERSION, "版が留まっていない"
    assert len(cuda_fetch.CUBLAS_SHA256) == 64, "ハッシュが無い"
    assert cuda_fetch.CUBLAS_VERSION in cuda_fetch.CUBLAS_WHEEL
    # **要るのは 2 つだけ**（実測。cuDNN 1.0GB と nvrtc 178MB は不要）
    assert set(cuda_fetch.WANT) == {"cublas64_12.dll", "cublasLt64_12.dll"}
    assert cuda_fetch.WHEEL_SIZE_MB > 100, "大きさを伝えられない"


def test_cuda_fetch_stops_on_a_different_file():
    """中身が想定と違えば、使わずに止める。"""
    import tempfile
    from unittest import mock
    from src import cuda_fetch

    with tempfile.TemporaryDirectory() as d:
        out = Path(d) / "cuda"

        def fake_dl(url, dest, on_progress):
            Path(dest).write_bytes("別物".encode("utf-8"))

        with mock.patch.object(cuda_fetch, "_wheel_url", lambda: "http://x"), \
                mock.patch.object(cuda_fetch, "_download", fake_dl):
            try:
                cuda_fetch.fetch(out)
            except cuda_fetch.CudaFetchError as e:
                assert "想定と違います" in str(e)
            else:
                raise AssertionError("別物を受け入れてしまった")
        assert not any(out.glob("*.dll")) if out.exists() else True


def test_cuda_fetch_failure_says_cpu_still_works():
    """**取得できなくても止めない。**CPU では今までどおり動く。"""
    import tempfile
    from unittest import mock
    from src import cuda_fetch

    with tempfile.TemporaryDirectory() as d:
        with mock.patch.object(cuda_fetch, "_wheel_url", lambda: "http://x"), \
                mock.patch.object(
                 cuda_fetch, "_download",
                 side_effect=cuda_fetch.CudaFetchError(
                     "GPU で動かすための部品を取得できませんでした。\n"
                     "取得できなくても、CPU では今までどおり動きます。")):
            try:
                cuda_fetch.fetch(Path(d) / "cuda")
            except cuda_fetch.CudaFetchError as e:
                assert "CPU では今までどおり動きます" in str(e), \
                    "取得できなくても使えることが伝わらない"
            else:
                raise AssertionError("失敗を知らせていない")


def test_cuda_is_not_bundled():
    """**同梱しない。**proprietary を MIT の配布物に混ぜない。

    受け取った人が「MIT だから自由に再配布できる」と誤解する。
    利用者の PC が配布元から直接取ってくる形にしてある。
    """
    spec = (Path(__file__).resolve().parent.parent / "build.spec").read_text(
        encoding="utf-8")
    for bad in ("cublas", "cudnn", "nvidia"):
        assert bad not in spec.lower(), f"build.spec に {bad} が入っている"
    req = (Path(__file__).resolve().parent.parent
           / "requirements.txt").read_text(encoding="utf-8")
    assert "nvidia-" not in req, "requirements に nvidia が入っている"


def test_every_bundled_nvidia_dll_is_named_in_credits():
    """**同梱物と CREDITS を突き合わせる。**

    build.spec に書いていなくても、PyInstaller は依存パッケージの中身を
    そのまま拾う。実際 CTranslate2 の配布物に cudnn64_9.dll が入っていて、
    CREDITS が「NVIDIA の部品は同梱していない」と読める状態のまま
    配布物に混ざっていた（2026-08-22 に発覚）。**書いてあることと
    実物が違う**のは、この製品でいちばん避けたい種類の欠陥。
    """
    import ctranslate2
    pkg = Path(ctranslate2.__file__).parent
    credits = (Path(__file__).resolve().parent.parent
               / "resources" / "CREDITS.txt").read_text(encoding="utf-8")
    found = [f.name for f in pkg.glob("*.dll")
             if f.name.lower().startswith(("cudnn", "cublas", "cuda", "nvrtc"))]
    assert found, "NVIDIA の dll が 1 つも無い（配布物が変わった可能性）"
    for dll in found:
        assert dll in credits, f"{dll} を同梱しているのに CREDITS に無い"


def test_bundled_ffmpeg_is_lgpl():
    """**同梱する ffmpeg は LGPL 版でなければならない。**

    本体は MIT。GPL の FFmpeg を同梱すると「同梱物全体を GPL にせよ」と
    衝突する。別プロセスで起動する使い方なので、LGPL なら MIT のまま配れる。

    2026-08-21 に、手元の ffmpeg が `--enable-gpl` 版だったことが分かった。
    CI も gyan.dev の release-essentials（GPL 版）を優先しており、コメントには
    「どちらも LGPL 版」と書いてあった。**誰も確かめていなかった。**
    """
    import subprocess
    root = Path(__file__).resolve().parent.parent
    exe = root / "ffmpeg" / "ffmpeg.exe"
    if not exe.is_file():
        return          # 未取得の環境では飛ばす（CI 側でも検査している）
    out = subprocess.run([str(exe), "-version"], capture_output=True,
                         text=True, encoding="utf-8", errors="replace")
    cfg = (out.stdout or "") + (out.stderr or "")
    for bad in ("--enable-gpl", "--enable-nonfree"):
        assert bad not in cfg, (
            f"同梱する ffmpeg が {bad} です。CREDITS.txt は LGPL と表示して"
            "おり、MIT の本体に同梱できません。BtbN の *-lgpl.zip を使うこと")
    # 表示と実態が合っていること
    from src.config import credits_text
    assert "LGPL" in credits_text()


def test_bundled_ffplay_is_present():
    """**区間再生に要る。**無いと簡易再生に落ち、聴き比べができない。

    割当作業は区間を聴き比べるのが中心なので、これが欠けると製品の
    中心機能が損なわれる。ビルド時に気づけるようにする。
    """
    root = Path(__file__).resolve().parent.parent
    if not (root / "ffmpeg" / "ffmpeg.exe").is_file():
        return          # ffmpeg ごと未取得の環境では飛ばす
    assert (root / "ffmpeg" / "ffplay.exe").is_file(), (
        "ffplay.exe がありません。区間再生が簡易再生に落ちます。"
        "README のビルド手順を参照してください")


def test_cuda_dlls_are_loaded_by_full_path():
    """**フルパスで先に読み込む。**PATH も add_dll_directory も凍結した exe
    では効かなかった（実機で「cublas64_12.dll is not found」・2026-08-22）。

    開発環境では動いていたが、それは**シェルの PATH に入っていたため**で、
    仕組みが効いていたのではない。手元で試すだけでは気づけない種類。

    一度プロセスに読み込ませれば、あとから名前で LoadLibrary されたときに
    既読のものが使われる。
    """
    import inspect
    import re
    from src import align
    src = inspect.getsource(align.add_cuda_dll_path)
    assert "ctypes" in src and "WinDLL" in src, (
        "フルパスでの読み込みが無い。PATH だけでは凍結した exe で落ちる")
    # **依存の順**（cublas は cublasLt を必要とする）。docstring にも
    # 名前が出るので、位置の比較ではなく**読み込む並びそのもの**を見る。
    order = re.search(r'for name in \(([^)]*)\)', src)
    assert order, "読み込む並びが見つからない"
    names = re.findall(r'"([^"]+)"', order.group(1))
    assert names == ["cublasLt64_12.dll", "cublas64_12.dll"], names
    # 読めなくても止めない（CPU に落ちるだけ）
    assert "except OSError" in src


class _Ran:
    """実際に走った転写器のふり（record に渡すもの）。"""

    def __init__(self, model, device, compute_type):
        self.model, self.device, self.compute_type = model, device, compute_type
        self.loaded = True


def test_engine_records_what_actually_ran():
    """**記録は「実際に走った値」でなければならない。**

    以前は `pick_device()`（＝使えそうかの判定）を書いていた。GPU の DLL が
    無くて CPU + small に落ちた実行が、記録には「large-v3 / cuda /
    int8_float16」と残っていた（実機で発生・2026-08-22）。**キャッシュの
    ファイル名は `.small.int8.` と正しく書いていたのに、engine だけが嘘を
    ついていた。**この製品は検証済みの記録を作るものなので、根幹に関わる。
    """
    from src.pipeline import EngineSpec, ENGINE_LOCAL
    spec = EngineSpec(mode=ENGINE_LOCAL, model="large-v3")
    rec = spec.record(_Ran("small", "cpu", "int8"))
    assert rec["model"] == "small", "頼んだモデルのほうを書いている"
    assert rec["device"] == "cpu"
    assert rec["compute_type"] == "int8"
    # **落ちたことも残す。**あとから「なぜ small なのか」が分かるように
    assert rec["requested_model"] == "large-v3"
    assert "fell_back" in rec


def test_model_advice_never_recommends_a_smaller_model():
    """**「small なら誤字が 4 割減る」は嘘。**逆さまの案内を出さない。

    large-v3 を選んでいるのに large-v3 が見つからないと `default_model()`
    が small を返し、案内文が「いまの設定は large-v3 ですが、small なら
    誤字が約 4 割減り…」になっていた（実機で発生・2026-08-22）。
    勧めてよいのは large-v3 へ上げるときだけ。

    **目印を移した(2026-08-31)。**以前は案内文の「誤字が約 4 割減り」を
    目印にしていたが、**その主張を撤回した**ため文言が変わり、検査が
    落ちた(設計書 §9.5.1 の注記)。**勧める理由は今後も変わりうる**ので、
    理由ではなく「GPU の案内であること」を指す文言に目印を移す。
    """
    import inspect
    from src import pipeline
    src = inspect.getsource(pipeline.run_segment_pipeline)
    anchor = "この機械には GPU があります"
    assert anchor in src, "GPU の案内が見つからない(目印を確かめること)"
    i = src.index(anchor)
    cond = src[max(0, i - 700):i]
    assert "best == align.GPU_DEFAULT_MODEL" in cond, \
        "小さいモデルを勧める経路が残っている"


def test_gpu_failure_says_where_it_looked():
    """**GPU が使えないとき、どこを見て何があったかを出す。**

    「cublas64_12.dll is not found」だけでは、開発機で再現しない不具合を
    推測で追うことになる（実機で 2 度外した・2026-08-22）。
    """
    from src import align
    report = align.add_cuda_dll_path()
    assert isinstance(report, list)
    import inspect
    from src import local_asr
    src = inspect.getsource(local_asr.LocalTranscriber._load)
    assert "add_cuda_dll_path()" in src, "失敗時に探索結果を出していない"
    assert "config_dir()" in src, "設定の置き場を出していない"


def test_engine_does_not_claim_a_device_it_never_used():
    """**キャッシュが全部当たった実行で、装置を名乗らない。**

    `_load` が一度も走らないと `device` は「使えそうか」の見込みのまま。
    それを「cuda で走った」と書くのは、直したはずの嘘の別経路。
    本文の素性（モデル・精度）はキャッシュの鍵が持っているので名乗ってよい。
    """
    from src.pipeline import EngineSpec, ENGINE_LOCAL
    spec = EngineSpec(mode=ENGINE_LOCAL, model="large-v3")
    ran = _Ran("large-v3", "cuda", "int8_float16")
    ran.loaded = False
    rec = spec.record(ran)
    assert rec["model"] == "large-v3", "本文の素性は鍵から分かる"
    assert rec["compute_type"] == "int8_float16"
    assert "cuda" not in rec["device"], "使っていない装置を名乗っている"
    assert "キャッシュ" in rec["device"]


def test_engine_record_stays_quiet_when_nothing_changed():
    """落ちていなければ、余計な印を付けない（読みにくくなる）。"""
    from src.pipeline import EngineSpec, ENGINE_LOCAL
    spec = EngineSpec(mode=ENGINE_LOCAL, model="large-v3")
    rec = spec.record(_Ran("large-v3", "cuda", "int8_float16"))
    assert rec["model"] == "large-v3" and rec["device"] == "cuda"
    assert "requested_model" not in rec and "fell_back" not in rec


def test_cpu_fallback_says_the_model_changed():
    """**モデルを落とすことも伝える。**

    「CPU に切り替えます」だけ出して黙って large-v3 → small にしていたため、
    20 分待ったあとに別物ができていた（実機で発生・2026-08-22）。
    """
    import inspect
    from src import local_asr
    src = inspect.getsource(local_asr.LocalTranscriber._load)
    assert "に切り替えます" in src
    # 落ちた先のモデル名を出しているか（変数で出していること）
    assert "self.model != was" in src, "モデルが変わったかを見ていない"
    # 2026-08-31 まで「誤字」という語で見ていたが、**その指標は撤回した**
    # (定義が復元できず、測り直すと向きが逆だった。設計書 §9.5.1 の注記)。
    # 落ちる品質として言えるのは固有名詞なので、そちらで見る。
    assert "固有名詞" in src, "品質が落ちることを伝えていない"


def test_diarize_runs_in_a_child_process():
    """**話者分離は子プロセスで回す。**

    sherpa-onnx は計算中に GIL を握るので、同じプロセスで回すと画面が
    まったく描けない。実測(2026-08-22・音声 7 分):

        同居       主スレッドが動けた割合 15% / 最長の空白 17.9 秒
        子プロセス 同 98% / 最長の空白 23 ミリ秒

    67 分では空白が 3 分以上になり、Windows が「応答なし」を付ける。
    利用者には落ちたようにしか見えない(実機で報告を受けた)。
    """
    argv = diarize._child_argv(Path("a.wav"), Path("o.json"), 9, None)
    assert "--diarize" in argv and "a.wav" in argv and "o.json" in argv
    assert "--speakers" in argv and "9" in argv
    assert "--model-dir" not in argv, "指定していないのに渡している"
    with_dir = diarize._child_argv(Path("a.wav"), Path("o.json"), 6, "C:/m")
    assert "--model-dir" in with_dir and "C:/m" in with_dir


def test_child_entry_is_checked_before_the_window():
    """**画面を作る前に見る。**あとに置くと、子プロセスまで窓を出す。"""
    import inspect
    from src import main as main_mod
    src = inspect.getsource(main_mod.main)
    assert src.index("--diarize") < src.index("from src.gui import App")


def test_progress_file_is_read_loosely():
    """進捗が読めなくても止めない(書きかけの瞬間に当たることがある)。"""
    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "x.progress"
        assert diarize._read_progress(p) is None        # まだ無い
        p.write_text("", encoding="utf-8")
        assert diarize._read_progress(p) is None        # 書きかけ
        p.write_text("12.5 420.0", encoding="utf-8")
        assert diarize._read_progress(p) == (12.5, 420.0)


def test_child_failure_leaves_a_message():
    """子が失敗したら理由をファイルに残す。**標準出力は使えない。**

    画面付きで固めた実行ファイルでは `sys.stdout` が None になり、
    書いた時点で落ちる。
    """
    with tempfile.TemporaryDirectory() as d:
        out = Path(d) / "turns.json"
        rc = diarize.child_main([str(Path(d) / "無い.wav"), str(out),
                                 "--speakers", "4"])
        assert rc != 0, "失敗したのに 0 を返している"
        assert not out.exists(), "失敗したのに結果を書いている"
        assert out.with_suffix(".err").exists(), "理由が残っていない"
        assert out.with_suffix(".err").read_text(encoding="utf-8").strip()


def test_inproc_escape_hatch_exists():
    """子プロセスで動かせない環境のために、同居で走らせる道を残す。

    **機能ごと失うよりは、固まっても動くほうがまし。**
    """
    import inspect
    src = inspect.getsource(diarize._run)
    assert "WHOSAID_DIARIZE_INPROC" in src
    assert "_diarize_here" in src


def test_a_person_can_remove_a_transcribed_segment():
    """**聴いて確かめた人は、転写が出した区間も消せる。**

    禁じているのは**自動**削除であって、音声を聴いて機械の重複だと判断した
    人が消すことは別のこと。実機で「同じ音声が 3 区間に書かれた」場面に
    当たり、結合を 2 回してから分割する遠回りを強いられた（2026-08-22）。
    """
    proj = _make_project()
    n = len(proj.segments)
    seg = proj.segments[1]
    assert not proj.is_added_utterance(seg), "前提: 転写が出した区間"
    proj.remove_added_utterance(1)
    assert len(proj.segments) == n - 1


def test_removing_a_transcribed_segment_keeps_what_was_removed():
    """**消した中身を履歴に残す。**

    足した区間を消すのは「自分の入力の取り消し」だが、こちらは機械の出力を
    人の判断で落とすこと。第三者が「ここで何が消されたか」を追えなければ、
    検証済みの記録にならない。
    """
    proj = _make_project()
    seg = proj.segments[1]
    text, start = seg.text, seg.start
    proj.remove_added_utterance(1)
    rec = proj.edit_log[-1]
    assert rec["op"] == "remove_segment", "足した区間と同じ扱いになっている"
    assert rec["actor"] == "user", "自動処理として残っている"
    assert rec["text"] == text, "消した本文が残っていない"
    assert rec["start"] == start
    from src.segments import LOG_LABELS
    assert "remove_segment" in LOG_LABELS, "履歴の表示名が無い"


def test_removing_an_added_utterance_stays_the_same():
    """足した区間を消したときの記録は、これまでどおり。"""
    proj = _make_project()
    seg = proj.segments[0]
    proj.add_utterance(start=seg.end + 0.1, end=seg.end + 0.5,
                       text="はい", cluster=seg.cluster)
    added = next(s for s in proj.segments if proj.is_added_utterance(s))
    proj.remove_added_utterance(added.index)
    rec = proj.edit_log[-1]
    assert rec["op"] == "remove_added_utterance"
    assert "text" not in rec, "取り消しなのに中身まで残している"


def test_uploaded_audio_is_deleted_even_on_failure():
    """**失敗しても、クラウドへ上げた音声を必ず消す。**

    正常終了のときだけ消していたので、途中でエラーになると Gemini 側に
    残っていた（48 時間後に自動削除されるが、録音は機微情報なので待たない）。
    公開前チェック v01 の指摘（2026-08-23 に対応）。
    """
    import inspect
    from src import transcribe
    src = inspect.getsource(transcribe)
    i = src.index("files.delete")
    head = src[max(0, i - 600):i]
    assert "finally:" in head, "finally の中で消していない"
    assert "uploaded = None" in src, "上げる前に初期化していない"


def test_installer_names_a_real_publisher():
    """**インストーラの発行者を空欄のままにしない。**

    LICENSE は実名に直したのに installer.iss は「Your Name」のままで、
    インストール画面にそう表示される状態だった（2026-08-23 に対応）。
    """
    iss = (Path(__file__).resolve().parent.parent
           / "installer.iss").read_text(encoding="utf-8")
    assert "Your Name" not in iss, "発行者が差し込み文字のまま"
    m = re.search(r'#define\s+MyAppPublisher\s+"([^"]+)"', iss)
    assert m and m.group(1).strip(), "発行者が空"


def test_api_key_is_dropped_when_it_cannot_be_protected():
    """**包めないなら書かない。**平文へ静かに落ちない。

    以前は DPAPI が失敗すると平文のまま書いていた。PRIVACY.md も画面も
    「DPAPI で暗号化する」と断言しているので、**成立しないことがある
    安全性を利用者に約束している**状態だった（2026-08-29）。

    秘密でない項目は書く。鍵が包めないからといって名簿まで失わせない。
    """
    from unittest import mock
    from src import config as cfg
    key = "AIzaSyTESTKEY-1234567890abcdefghij"
    with tempfile.TemporaryDirectory() as d:
        real = os.environ.get("APPDATA")
        os.environ["APPDATA"] = d
        try:
            # DPAPI が使えない機械を作る（包めずに平文が返る）
            with mock.patch.object(cfg, "_dpapi", lambda *a, **k: None):
                dropped = cfg.save_config({"api_key": key, "roster": "佐藤"})
            assert dropped == ["api_key"], f"知らせていない: {dropped}"
            raw = cfg.config_path().read_text(encoding="utf-8")
            assert key not in raw, "**包めないのに平文で書いている**"
            assert "api_key" not in raw, "空でも項目を残している"
            assert "佐藤" in raw, "秘密でない項目まで落としている"
            assert cfg.load_config().get("api_key") is None
        finally:
            if real is not None:
                os.environ["APPDATA"] = real


def test_saving_a_protected_key_reports_nothing_dropped():
    """包めたときは何も落とさない（上の検査が常に真にならないこと）。"""
    from src import config as cfg
    if sys.platform != "win32":
        print("       (DPAPI が無いので飛ばします)")
        return
    with tempfile.TemporaryDirectory() as d:
        real = os.environ.get("APPDATA")
        os.environ["APPDATA"] = d
        try:
            dropped = cfg.save_config({"api_key": "AIzaSyOK-0123456789",
                                       "roster": "佐藤"})
            assert dropped == [], f"包めたのに落としている: {dropped}"
            assert cfg.load_config().get("api_key") == "AIzaSyOK-0123456789"
        finally:
            if real is not None:
                os.environ["APPDATA"] = real


def test_api_key_is_not_stored_in_the_clear():
    """**API キーを平文で置かない。**

    設定ファイルは %APPDATA% にあり、そのままだと中を開けば読める。
    Windows の DPAPI で包み、そのパソコンのその利用者でしか復号できない
    形にする（公開前チェック v01 の指摘・2026-08-23 に対応）。
    """
    from src import config as cfg
    key = "AIzaSyTESTKEY-1234567890abcdefghij"
    with tempfile.TemporaryDirectory() as d:
        real = os.environ.get("APPDATA")
        os.environ["APPDATA"] = d
        try:
            cfg.save_config({"api_key": key, "roster": "佐藤"})
            raw = cfg.config_path().read_text(encoding="utf-8")
            if sys.platform == "win32":
                assert key not in raw, "設定ファイルに平文で書かれている"
                assert cfg._ENC_PREFIX in raw, "包んだ印が無い"
            # 包んでも、読み直せば元に戻る
            assert cfg.load_config().get("api_key") == key
            # 秘密でない項目は包まない（読めなくなると困る）
            assert "佐藤" in raw
        finally:
            if real is not None:
                os.environ["APPDATA"] = real


def test_old_plaintext_key_still_loads():
    """**旧版の平文の設定を読めなくしない。**移行で鍵を失わせない。"""
    from src import config as cfg
    key = "AIzaSyOLDPLAINTEXTKEY-0987654321"
    with tempfile.TemporaryDirectory() as d:
        real = os.environ.get("APPDATA")
        os.environ["APPDATA"] = d
        try:
            p = cfg.config_path()
            p.write_text(json.dumps({"api_key": key}), encoding="utf-8")
            assert cfg.load_config().get("api_key") == key
        finally:
            if real is not None:
                os.environ["APPDATA"] = real


def test_secret_round_trip_and_idempotence():
    """包む・ほどくが往復する。二重に包まない。"""
    from src.config import protect_secret, unprotect_secret, is_protected
    v = "sk-test-abcdefghijklmnopqrstuvwxyz"
    enc = protect_secret(v)
    assert unprotect_secret(enc) == v
    assert protect_secret(enc) == enc, "二重に包んでいる"
    assert unprotect_secret(v) == v, "平文を壊している"
    if sys.platform == "win32":
        assert is_protected(enc) and not is_protected(v)



def test_unreadable_secret_is_not_destroyed_by_an_unrelated_save():
    """**解けない鍵を、鍵と無関係な保存が空文字で上書きしていた。**

    筋道（2026-08-31 に再現した）:
      1. `unprotect_secret` は DPAPI で解けないとき "" を返す
      2. `load_config` が api_key = "" を dict に入れる
      3. **出力先を変えただけの保存**が走ると、save_config が "" を素通しし、
         包まれた元の値を "" で上書きする（save_config の呼び出しは 8 箇所）
      4. `dropped` は空なので**警告も出ない**

    **上書き前なら、正しい PC・正しいアカウントへ戻れば復号できた。**
    上書き後は永久に失われる。設定を別 PC からコピーした・Windows の
    プロファイルを作り直した・バックアップから別機に戻した——どれも普通に起きる。

    直したあとの規則は「無い＝消す / 空＝残す / 平文＝包んで書く」。
    """
    import base64
    import json
    import os
    import tempfile
    from src import config as cfg

    keep = os.environ.get("APPDATA")
    try:
        with tempfile.TemporaryDirectory() as d:
            os.environ["APPDATA"] = d
            p = cfg.config_path()
            # 別の PC / 別アカウントで包まれた鍵（このマシンでは解けない）
            bad = cfg._ENC_PREFIX + base64.b64encode(b"not-a-dpapi-blob").decode()
            p.write_text(json.dumps({"api_key": bad, "roster": "佐藤",
                                     "out_dir": "D:/x"}, ensure_ascii=False),
                         encoding="utf-8")

            c = cfg.load_config()
            assert c["api_key"] == "", "解けない鍵は画面に出さない"
            # **「解けなかった」と「元から空」を分けていること**
            assert cfg.UNREADABLE_MARK in c, "解けなかった印が無い"
            assert c[cfg.UNREADABLE_MARK] == ["api_key"]

            # 鍵と無関係な保存
            c["out_dir"] = "D:/y"
            assert cfg.save_config(c) == [], "落としたものは無いはず"
            after = json.loads(p.read_text(encoding="utf-8"))
            assert after["api_key"] == bad, "**包まれた鍵が消えた**"
            assert after["out_dir"] == "D:/y", "他の設定は保存される"
            assert cfg.UNREADABLE_MARK not in after, "印がファイルに漏れている"

            # --- 「無い＝消す」は残っていること（画面の「消す」が使う）---
            c2 = cfg.load_config()
            c2.pop("api_key", None)
            cfg.save_config(c2)
            assert "api_key" not in json.loads(p.read_text(encoding="utf-8")), \
                "pop したら消えること（消す入口が壊れる）"

            # --- 元から鍵が無いなら、空を渡しても項目を作らない ---
            cfg.save_config({"api_key": "", "roster": "佐藤"})
            assert "api_key" not in json.loads(p.read_text(encoding="utf-8"))

            # --- 平文を渡せば、ちゃんと包んで上書きする ---
            cfg.save_config({"api_key": "AIzaSyREAL-0123456789", "roster": "佐藤"})
            got = json.loads(p.read_text(encoding="utf-8"))["api_key"]
            if cfg.is_protected(cfg.protect_secret("x")):     # DPAPI がある環境
                assert cfg.is_protected(got), "包まずに書いている"
                assert got != bad, "古い値が残ってしまっている"

            # --- 印はファイルへ書かない（`_` で始まる項目すべて）---
            cfg.save_config({"roster": "佐藤", "_なにか": 1, "普通": 2})
            got = json.loads(p.read_text(encoding="utf-8"))
            assert "_なにか" not in got and got.get("普通") == 2
    finally:
        if keep is None:
            os.environ.pop("APPDATA", None)
        else:
            os.environ["APPDATA"] = keep


def test_empty_secret_does_not_mean_delete():
    """**空は「消してよい」ではない。**この読み違いが上の穴の正体だった。

    別々に縛る——上のテストが長いので、規則そのものを短く残しておく。
    消す入口は `pop` を使うこと（`""` を代入して消したつもりにしない）。
    """
    import json
    import os
    import tempfile
    from src import config as cfg

    keep = os.environ.get("APPDATA")
    try:
        with tempfile.TemporaryDirectory() as d:
            os.environ["APPDATA"] = d
            cfg.save_config({"api_key": "AIzaSyKEEP-0123456789"})
            first = json.loads(cfg.config_path().read_text(encoding="utf-8"))
            assert first.get("api_key"), "前提: 鍵が入っていること"
            # 空を渡す＝渡すものが無い。**消さない**
            cfg.save_config({"api_key": "", "roster": "田中"})
            after = json.loads(cfg.config_path().read_text(encoding="utf-8"))
            assert after.get("api_key") == first["api_key"], "空で消してはいけない"
            assert after.get("roster") == "田中"
    finally:
        if keep is None:
            os.environ.pop("APPDATA", None)
        else:
            os.environ["APPDATA"] = keep


def test_plaintext_secret_is_marked_on_load_but_not_rewritten():
    """**平文のまま置かれている鍵を、読み込み時に見分ける。**

    旧版（v2.0.x 以前）が書いた設定や、旧名フォルダから引き継いだ設定には
    鍵が平文で入っている（2026-09-03 に一号機・二号機で確認）。読み込みは
    旧版互換でそのまま通すので、**画面は暗号文と同じ顔をしていた。**

    印を付けるだけで、**値は書き換えない。**印はファイルに書かない。
    「解けなかった」「元から空」「平文」は三つとも別の状態。
    """
    import json
    import os
    import tempfile
    from src import config as cfg

    keep = os.environ.get("APPDATA")
    try:
        with tempfile.TemporaryDirectory() as d:
            os.environ["APPDATA"] = d
            p = cfg.config_path()

            # --- 平文の鍵 → 印が立ち、値はそのまま読める ---
            p.write_text(json.dumps({"api_key": "AIzaSyPLAIN-0123456789",
                                     "roster": "佐藤"}, ensure_ascii=False),
                         encoding="utf-8")
            c = cfg.load_config()
            assert c["api_key"] == "AIzaSyPLAIN-0123456789", "旧版互換で読めること"
            assert c.get(cfg.PLAINTEXT_MARK) == ["api_key"], "平文の印が無い"
            assert cfg.UNREADABLE_MARK not in c, "平文を「解けない」と混同している"
            # **読んだだけでは書き換えない**（黙って包み直さない）
            assert json.loads(p.read_text(encoding="utf-8"))["api_key"] \
                == "AIzaSyPLAIN-0123456789", "読み込みがファイルを書き換えた"

            # --- 元から空 / 無い → 印は立たない ---
            p.write_text(json.dumps({"api_key": "", "roster": "佐藤"}),
                         encoding="utf-8")
            c = cfg.load_config()
            assert cfg.PLAINTEXT_MARK not in c, "空に平文の印が立っている"
            p.write_text(json.dumps({"roster": "佐藤"}), encoding="utf-8")
            c = cfg.load_config()
            assert cfg.PLAINTEXT_MARK not in c, "無いのに平文の印が立っている"

            # --- 解けない暗号文 → 平文の印ではなく、解けない印 ---
            import base64
            bad = cfg._ENC_PREFIX + base64.b64encode(b"not-a-dpapi-blob").decode()
            p.write_text(json.dumps({"api_key": bad}), encoding="utf-8")
            c = cfg.load_config()
            assert cfg.PLAINTEXT_MARK not in c, "解けない鍵に平文の印が立っている"
            assert c.get(cfg.UNREADABLE_MARK) == ["api_key"]

            # --- 包まれた鍵 → どちらの印も立たない ---
            if cfg.is_protected(cfg.protect_secret("x")):     # DPAPI がある環境
                cfg.save_config({"api_key": "AIzaSyREAL-0123456789"})
                c = cfg.load_config()
                assert cfg.PLAINTEXT_MARK not in c, "包まれた鍵に平文の印が立っている"
                assert cfg.UNREADABLE_MARK not in c

            # --- 印はファイルへ漏れない ---
            p.write_text(json.dumps({"api_key": "AIzaSyPLAIN-0123456789"}),
                         encoding="utf-8")
            c = cfg.load_config()
            c["roster"] = "鈴木"
            cfg.save_config(c)
            after = json.loads(p.read_text(encoding="utf-8"))
            assert cfg.PLAINTEXT_MARK not in after, "印がファイルに漏れている"
    finally:
        if keep is None:
            os.environ.pop("APPDATA", None)
        else:
            os.environ["APPDATA"] = keep


def test_plaintext_mark_does_not_accumulate_across_round_trips():
    """**印のリストが、読む・保存する・読むの往復で増えないこと。**

    `setdefault(...).append(k)` は「印がファイルに書かれない」前提に立って
    いる。もし何かの経路で `_plaintext` がファイルに書かれると、次の
    読み込みで既存のリストに追記され、起動のたびに要素が増える。
    前提に頼らず、往復そのものを縛る（2026-09-03 のレビューで指摘）。
    """
    import json
    import os
    import tempfile
    from src import config as cfg

    keep = os.environ.get("APPDATA")
    try:
        with tempfile.TemporaryDirectory() as d:
            os.environ["APPDATA"] = d
            p = cfg.config_path()
            p.write_text(json.dumps({"api_key": "AIzaSyPLAIN-0123456789",
                                     "roster": "佐藤"}, ensure_ascii=False),
                         encoding="utf-8")

            # --- 読むだけを繰り返しても 1 件のまま ---
            for _ in range(3):
                c = cfg.load_config()
                assert c.get(cfg.PLAINTEXT_MARK) == ["api_key"], \
                    f"読み込みで印が増減した: {c.get(cfg.PLAINTEXT_MARK)}"

            # --- 印の入った dict をそのまま保存 → ファイルに印は無い ---
            c["roster"] = "鈴木"
            cfg.save_config(c)
            raw = json.loads(p.read_text(encoding="utf-8"))
            assert cfg.PLAINTEXT_MARK not in raw, "印がファイルに書かれた"
            assert not [k for k in raw if str(k).startswith("_")], \
                f"`_` 始まりの項目がファイルに漏れた: {sorted(raw)}"
            assert raw.get("roster") == "鈴木"

            # --- 読み直す → 印は無い（鍵が包まれた）か、あっても 1 件 ---
            c2 = cfg.load_config()
            marks = c2.get(cfg.PLAINTEXT_MARK)
            assert marks is None or marks == ["api_key"], f"印が増えた: {marks}"
            if cfg.is_protected(cfg.protect_secret("x")):     # DPAPI がある環境
                assert marks is None, "保存で包まれたはずの鍵に平文の印が残っている"
                assert cfg.is_protected(raw["api_key"]), "保存を通っても包まれていない"

            # --- ファイルに古い印が紛れていても、拾わずに立て直す ---
            p.write_text(json.dumps({"api_key": "AIzaSyPLAIN-0123456789",
                                     cfg.PLAINTEXT_MARK: ["api_key", "api_key"],
                                     cfg.UNREADABLE_MARK: ["api_key"]}),
                         encoding="utf-8")
            c3 = cfg.load_config()
            assert c3.get(cfg.PLAINTEXT_MARK) == ["api_key"], \
                f"ファイルの古い印に追記している: {c3.get(cfg.PLAINTEXT_MARK)}"
            assert cfg.UNREADABLE_MARK not in c3, "ファイルの古い印を拾っている"
    finally:
        if keep is None:
            os.environ.pop("APPDATA", None)
        else:
            os.environ["APPDATA"] = keep


def test_public_documents_exist_and_say_the_truth():
    """**公開に必要な文書を、実態と揃えたまま保つ。**

    無料配布にあたって用意した（2026-08-23）。消したり、実装と食い違ったまま
    放置したりしないよう、機械に見張らせる。
    """
    root = Path(__file__).resolve().parent.parent
    for name in ("PRIVACY.md", "TERMS_OF_USE.md", "SECURITY.md", "LICENSE"):
        f = root / name
        assert f.is_file(), f"{name} が無い"
        assert f.stat().st_size > 200, f"{name} が中身のない状態"

    lic = (root / "LICENSE").read_text(encoding="utf-8")
    assert "[Your Name]" not in lic, "著作権表示が差し込み文字のまま"

    # **PRIVACY が「暗号化する」と書くなら、実装がそうなっていること。**
    priv = (root / "PRIVACY.md").read_text(encoding="utf-8")
    if "DPAPI" in priv:
        from src import config as cfg
        assert hasattr(cfg, "protect_secret"), "書いてあるのに実装が無い"
        assert "api_key" in cfg.SECRET_KEYS

    # **「平文では保存しない」と書くなら、包めないとき本当に書かないこと。**
    # ここは文書のほうが実装より強いことを約束しがちな箇所で、実際に
    # 食い違っていた（2026-08-29 に発覚。DPAPI 失敗時は平文で書いていた）。
    if "平文で保存することはありません" in priv:
        import tempfile as _tf
        from unittest import mock
        from src import config as cfg
        with _tf.TemporaryDirectory() as _d:
            _real = os.environ.get("APPDATA")
            os.environ["APPDATA"] = _d
            try:
                with mock.patch.object(cfg, "_dpapi", lambda *a, **k: None):
                    cfg.save_config({"api_key": "AIzaSyPLAIN-0123456789"})
                _raw = cfg.config_path().read_text(encoding="utf-8")
                assert "AIzaSyPLAIN-0123456789" not in _raw, (
                    "PRIVACY.md は「平文で保存することはありません」と書いて"
                    "いるのに、平文で書いている")
            finally:
                if _real is not None:
                    os.environ["APPDATA"] = _real

    # **README が「既定では端末から出ない」と書くなら、既定がローカルであること。**
    readme = (root / "README.md").read_text(encoding="utf-8")
    assert "端末から出ません" in readme
    from src.segments import ENGINE_LOCAL
    import inspect
    from src import gui
    src = inspect.getsource(gui.App._build_ui)
    assert "既定はローカル" in src, "README と既定が食い違う"


def test_diarize_start_is_logged_by_the_parent():
    """**話者分離を始めたことを、親のログに出す。**

    この行は `_diarize_here` の中にあり、子プロセスへ移したときに親の
    ログから消えた。結果、転写が終わってから 10 分以上まったく無反応に
    見えた（別 PC の実機で発覚・2026-08-23）。**「落ちたと思われる」問題を
    直したはずが、別の形で再発していた。**
    """
    import inspect
    from src import diarize
    outer = inspect.getsource(diarize.diarize)
    inner = inspect.getsource(diarize._diarize_here)
    assert "話者を分けています" in outer, "親が始まりを告げていない"
    assert "話者を分けています" not in inner, "子と二重に出る"
    assert "2 割" in outer, "どれくらいかかるかを伝えていない"


def test_cpu_run_says_why_gpu_is_not_used():
    """**CPU で回すときは、その理由を言う。**

    「GPU が無い」のか「あるが部品が足りない」のかで、利用者がすることが
    変わる。別 PC で動かしたときに、遅い理由が分からなかった（2026-08-23）。
    """
    import inspect
    from src import local_asr
    src = inspect.getsource(local_asr.LocalTranscriber._load)
    assert "GPU は使いません" in src
    why = inspect.getsource(local_asr.LocalTranscriber._why_not_gpu)
    for phrase in ("見つかりません", "部品がありません"):
        assert phrase in why, f"理由の場合分けに「{phrase}」が無い"


# ======================================================================
# pytest が無い環境向けの簡易ランナー
# ======================================================================

if __name__ == "__main__":
    failures = 0
    for name, fn in sorted(globals().items()):
        if name.startswith("test_") and callable(fn):
            try:
                fn()
                print(f"  ok   {name}")
            except Exception as e:
                failures += 1
                import traceback
                print(f"  FAIL {name}: {e}")
                traceback.print_exc()
    print(f"\n{'FAILED' if failures else 'ALL PASSED'} ({failures} failures)")
    sys.exit(1 if failures else 0)
